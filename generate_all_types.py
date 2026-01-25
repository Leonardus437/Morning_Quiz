#!/usr/bin/env python3
"""Generate C Programming Questions - DOC and PDF with ALL Question Types"""

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import os

def create_doc():
    doc = Document()
    
    # Question 1 - Multiple Choice
    doc.add_paragraph("1. What is the correct syntax to declare an integer variable in C?")
    doc.add_paragraph("A) int x;")
    doc.add_paragraph("B) integer x;")
    doc.add_paragraph("C) var x;")
    doc.add_paragraph("D) x int;")
    doc.add_paragraph("Answer: A")
    doc.add_paragraph("")
    
    # Question 2 - True/False
    doc.add_paragraph("2. The main() function is mandatory in every C program. True or False?")
    doc.add_paragraph("Answer: True")
    doc.add_paragraph("")
    
    # Question 3 - Short Answer
    doc.add_paragraph("3. What does the #include directive do in C? Explain briefly.")
    doc.add_paragraph("Answer: Includes header files containing function declarations and definitions")
    doc.add_paragraph("")
    
    # Question 4 - Essay
    doc.add_paragraph("4. Explain in detail the difference between arrays and pointers in C programming. Discuss their memory allocation, syntax, and use cases with examples.")
    doc.add_paragraph("Answer: Arrays are collections of elements stored in contiguous memory locations with fixed size, while pointers are variables that store memory addresses. Arrays use static memory allocation, pointers use dynamic allocation. Example: int arr[5] vs int *ptr.")
    doc.add_paragraph("")
    
    # Question 5 - Multiple Select
    doc.add_paragraph("5. Select all correct answers: Which of the following are valid C data types?")
    doc.add_paragraph("A) int")
    doc.add_paragraph("B) float")
    doc.add_paragraph("C) string")
    doc.add_paragraph("D) char")
    doc.add_paragraph("E) boolean")
    doc.add_paragraph("Answer: A,B,D")
    doc.add_paragraph("")
    
    # Question 6 - Dropdown Select
    doc.add_paragraph("6. Select from dropdown: Which operator is used for address-of in C?")
    doc.add_paragraph("A) *")
    doc.add_paragraph("B) &")
    doc.add_paragraph("C) ->")
    doc.add_paragraph("D) .")
    doc.add_paragraph("Answer: B")
    doc.add_paragraph("")
    
    # Question 7 - Fill in Blanks
    doc.add_paragraph("7. Fill in the blanks: The _____ function is used to allocate memory dynamically and _____ function is used to free allocated memory.")
    doc.add_paragraph("Answer: malloc|||free")
    doc.add_paragraph("")
    
    # Question 8 - Matching Pairs
    doc.add_paragraph("8. Match each C keyword with its purpose:")
    doc.add_paragraph("Left: if, for, struct")
    doc.add_paragraph("Right: Conditional statement, Loop statement, User-defined data type")
    doc.add_paragraph("Answer: if:Conditional statement|||for:Loop statement|||struct:User-defined data type")
    doc.add_paragraph("")
    
    # Question 9 - Drag & Drop Ordering
    doc.add_paragraph("9. Arrange in correct order: Put the C program compilation steps in the right sequence.")
    doc.add_paragraph("Items: Linking, Preprocessing, Assembly, Compilation")
    doc.add_paragraph("Answer: Preprocessing,Compilation,Assembly,Linking")
    doc.add_paragraph("")
    
    # Question 10 - Linear Scale
    doc.add_paragraph("10. Rate on a scale of 1-10: How important is memory management in C programming?")
    doc.add_paragraph("Answer: 9")
    doc.add_paragraph("")
    
    # Question 11 - Code Writing
    doc.add_paragraph("11. Write your code: Create a C function to find the factorial of a number using recursion.")
    doc.add_paragraph("Answer: int factorial(int n) { if(n <= 1) return 1; else return n * factorial(n-1); }")
    doc.add_paragraph("")
    
    # Question 12 - SQL Query (adapted for C file operations)
    doc.add_paragraph("12. Write your SQL query: Write a C code snippet to open a file named 'data.txt' in read mode and check if it opened successfully.")
    doc.add_paragraph("Answer: FILE *fp = fopen(\"data.txt\", \"r\"); if(fp == NULL) { printf(\"Error opening file\"); }")
    doc.add_paragraph("")
    
    # Question 13 - Multi-Grid
    doc.add_paragraph("13. Select one answer for each row: Rate each C concept on difficulty level.")
    doc.add_paragraph("Rows: Variables, Pointers, Functions")
    doc.add_paragraph("Columns: Easy, Medium, Hard")
    doc.add_paragraph("Answer: Variables:Easy|||Pointers:Hard|||Functions:Medium")
    
    doc.save("C_Programming_All_Types.docx")
    print("Created: C_Programming_All_Types.docx")

def create_pdf():
    doc = SimpleDocTemplate("C_Programming_All_Types.pdf", pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    questions = [
        "1. What is the correct syntax to declare an integer variable in C?\nA) int x;\nB) integer x;\nC) var x;\nD) x int;\nAnswer: A",
        "2. The main() function is mandatory in every C program. True or False?\nAnswer: True",
        "3. What does the #include directive do in C? Explain briefly.\nAnswer: Includes header files containing function declarations and definitions",
        "4. Explain in detail the difference between arrays and pointers in C programming. Discuss their memory allocation, syntax, and use cases with examples.\nAnswer: Arrays are collections of elements stored in contiguous memory locations with fixed size, while pointers are variables that store memory addresses.",
        "5. Select all correct answers: Which of the following are valid C data types?\nA) int\nB) float\nC) string\nD) char\nE) boolean\nAnswer: A,B,D",
        "6. Select from dropdown: Which operator is used for address-of in C?\nA) *\nB) &\nC) ->\nD) .\nAnswer: B",
        "7. Fill in the blanks: The _____ function is used to allocate memory dynamically and _____ function is used to free allocated memory.\nAnswer: malloc|||free",
        "8. Match each C keyword with its purpose:\nLeft: if, for, struct\nRight: Conditional statement, Loop statement, User-defined data type\nAnswer: if:Conditional statement|||for:Loop statement|||struct:User-defined data type",
        "9. Arrange in correct order: Put the C program compilation steps in the right sequence.\nItems: Linking, Preprocessing, Assembly, Compilation\nAnswer: Preprocessing,Compilation,Assembly,Linking",
        "10. Rate on a scale of 1-10: How important is memory management in C programming?\nAnswer: 9",
        "11. Write your code: Create a C function to find the factorial of a number using recursion.\nAnswer: int factorial(int n) { if(n <= 1) return 1; else return n * factorial(n-1); }",
        "12. Write your SQL query: Write a C code snippet to open a file named 'data.txt' in read mode and check if it opened successfully.\nAnswer: FILE *fp = fopen(\"data.txt\", \"r\"); if(fp == NULL) { printf(\"Error opening file\"); }",
        "13. Select one answer for each row: Rate each C concept on difficulty level.\nRows: Variables, Pointers, Functions\nColumns: Easy, Medium, Hard\nAnswer: Variables:Easy|||Pointers:Hard|||Functions:Medium"
    ]
    
    for question in questions:
        story.append(Paragraph(question, styles['Normal']))
        story.append(Spacer(1, 12))
    
    doc.build(story)
    print("Created: C_Programming_All_Types.pdf")

def main():
    print("Generating C Programming Questions with ALL 13 Types...")
    create_doc()
    create_pdf()
    print("Files created successfully!")
    print("C_Programming_All_Types.docx")
    print("C_Programming_All_Types.pdf")
    print("")
    print("Question Types Included:")
    print("1. Multiple Choice")
    print("2. True/False")
    print("3. Short Answer")
    print("4. Essay")
    print("5. Multiple Select")
    print("6. Dropdown Select")
    print("7. Fill in Blanks")
    print("8. Matching Pairs")
    print("9. Drag & Drop Ordering")
    print("10. Linear Scale")
    print("11. Code Writing")
    print("12. SQL Query")
    print("13. Multi-Grid")

if __name__ == "__main__":
    main()