"""
Professional Question Parser for TVET Quiz System
Supports 13 question types with intelligent detection and extraction
"""

import re
from typing import Dict, List, Optional, Tuple
import io

class AdvancedQuestionParser:
    """
    Professional parser that can analyze any document and extract questions
    based on their types with high accuracy.
    
    Supported Question Types:
    1. Multiple Choice (MCQ)
    2. True/False
    3. Multiple Select
    4. Fill in the Blanks
    5. Matching Pairs
    6. Drag & Drop Ordering
    7. Linear Scale
    8. Code Writing
    9. SQL Query
    10. Dropdown Select
    11. Short Answer
    12. Essay
    13. Multi-Grid
    """
    
    # Question type patterns with priority (higher priority = checked first)
    QUESTION_TYPE_PATTERNS = [
        # Specific types first (higher priority)
        ('sql_query', [
            r'\b(sql|query|database|select\s+\*|insert\s+into|create\s+table|update\s+set|delete\s+from)\b',
            r'\b(join|where|group\s+by|order\s+by|having)\b'
        ], 90),
        
        ('code_writing', [
            r'\b(write|implement|create|develop)\s+(a\s+)?(code|program|function|method|algorithm|class)\b',
            r'\b(python|java|javascript|c\+\+|programming)\b',
            r'(def\s+|function\s+|class\s+|public\s+|private\s+)'
        ], 85),
        
        ('true_false', [
            r'\b(true\s+or\s+false|t/f)\b',
            r'(state\s+whether|indicate\s+whether|is\s+it\s+true)',
            r'\?\s*$.*\b(true|false)\b'
        ], 80),
        
        ('matching_pairs', [
            r'\b(match|pair|connect|correspond|associate)\b.*\b(following|items|columns)\b',
            r'(column\s+a|column\s+b|match.*with)',
            r'(draw\s+lines|connect.*to)'
        ], 75),
        
        ('drag_drop_ordering', [
            r'\b(arrange|order|sequence|sort|rank|organize)\b.*\b(correct|proper|right)\b',
            r'(put\s+in\s+order|arrange\s+in|sequence\s+the)',
            r'(chronological|ascending|descending)\s+order'
        ], 70),
        
        ('fill_blanks', [
            r'(fill\s+in|complete|fill\s+the)\s+(blank|gap|missing)',
            r'_{3,}',  # Three or more underscores
            r'\[.*?\].*\[.*?\]',  # Multiple brackets
            r'\b(insert|supply)\s+the\s+(correct|appropriate|missing)\b'
        ], 65),
        
        ('multiple_select', [
            r'\b(select\s+all|choose\s+all|tick\s+all|mark\s+all)\b',
            r'\b(multiple\s+correct|more\s+than\s+one)\b',
            r'(all\s+that\s+apply|which\s+of.*are)'
        ], 60),
        
        ('linear_scale', [
            r'\b(rate|rating|scale)\b.*\b(1\s*to\s*10|1-10|1\s*-\s*5)\b',
            r'(on\s+a\s+scale|rank\s+from)',
            r'(strongly\s+agree|strongly\s+disagree)'
        ], 55),
        
        ('dropdown_select', [
            r'\b(dropdown|select\s+from|choose\s+from)\b.*\b(dropdown|menu|list)\b',
            r'(pick\s+one\s+from|select\s+the\s+appropriate)'
        ], 50),
        
        ('essay', [
            r'\b(essay|discuss|elaborate|explain\s+in\s+detail|describe\s+in\s+detail)\b',
            r'(write\s+an?\s+essay|write\s+a\s+detailed)',
            r'(in\s+your\s+own\s+words.*paragraph|compose\s+a)'
        ], 45),
        
        ('short_answer', [
            r'\b(short\s+answer|brief|briefly|one\s+sentence)\b',
            r'(in\s+few\s+words|concisely|summarize)',
            r'(define|what\s+is|who\s+is|when\s+did)'
        ], 40),
        
        ('multi_grid', [
            r'(grid|matrix|table).*\b(rows|columns)\b',
            r'(rate\s+each|evaluate\s+each).*\b(row|item)\b'
        ], 35),
        
        # Default to multiple choice (lowest priority)
        ('multiple_choice', [
            r'[A-D]\)',  # A) B) C) D)
            r'\b(choose|select|pick)\b.*\b(one|correct|best)\b',
            r'(which\s+of|what\s+is|who\s+is)'
        ], 10),
    ]
    
    # Question number patterns
    QUESTION_NUMBER_PATTERNS = [
        r'^Question\s+\d+',  # Question 1, Question 2
        r'^Q\.?\s*\d+',  # Q1, Q.1, Q 1
        r'^\d+\.(?!\d)',  # 1. , 2. (but not 1.5)
        r'^\d+\)',  # 1) , 2)
        r'^\(\d+\)',  # (1), (2)
        r'^\d+\s*[-–—]\s*',  # 1 - , 2 –
    ]
    
    # Answer patterns
    ANSWER_PATTERNS = [
        (r'answer\s*:\s*(.+?)(?=\n|$)', 'answer:'),
        (r'correct\s+answer\s*:\s*(.+?)(?=\n|$)', 'correct answer:'),
        (r'solution\s*:\s*(.+?)(?=\n|$)', 'solution:'),
        (r'\(answer\s*:\s*(.+?)\)', '(answer:)'),
        (r'ans\s*:\s*(.+?)(?=\n|$)', 'ans:'),
        (r'key\s*:\s*(.+?)(?=\n|$)', 'key:'),
        (r'correct\s*:\s*(.+?)(?=\n|$)', 'correct:'),
    ]
    
    # Option patterns (for MCQ, multiple select, etc.)
    OPTION_PATTERNS = [
        r'([A-Z])\)\s*([^A-Z\)\n]+)',  # A) option
        r'([A-Z])\.\s*([^A-Z\.\n]+)',  # A. option
        r'([a-z])\)\s*([^a-z\)\n]+)',  # a) option
        r'([a-z])\.\s*([^a-z\.\n]+)',  # a. option
        r'\(([A-Z])\)\s*([^\(\n]+)',  # (A) option
        r'([A-D]):\s*([^\n]+)',  # A: option
    ]
    
    # Noise patterns to skip
    NOISE_PATTERNS = [
        r'^(section|part|chapter|unit)\s+[A-Z0-9]',
        r'^(instructions?|directions?|note|important)\s*:',
        r'^(name|student|class|date|time|marks?|score|points?)\s*:',
        r'^(total|maximum|max)\s+(marks?|points?|time|score)',
        r'^(answer\s+all|choose|select)\s+questions?',
        r'^page\s+\d+',
        r'^\d+\s+(marks?|points?|minutes?|mins?)$',
        r'^end\s+of\s+(test|exam|paper|questions?|quiz)',
        r'^(exam|test|quiz)\s+(title|name)',
        r'^(department|course|subject|level)\s*:',
    ]
    
    def __init__(self):
        self.debug_mode = True
        
    def extract_text_from_file(self, content: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        filename_lower = filename.lower()
        
        try:
            if filename_lower.endswith('.pdf'):
                return self._extract_from_pdf(content)
            elif filename_lower.endswith('.docx'):
                return self._extract_from_docx(content)
            elif filename_lower.endswith('.doc'):
                return self._extract_from_doc(content)
            elif filename_lower.endswith('.txt'):
                return content.decode('utf-8', errors='ignore')
            elif filename_lower.endswith(('.rtf', '.odt')):
                # Try to extract as text
                return content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file format: {filename}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from {filename}: {str(e)}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text_parts = []
        for page in pdf_reader.pages:
            text_parts.append(page.extract_text())
        return '\n'.join(text_parts)
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return '\n'.join(text_parts)
        except ImportError:
            raise ValueError("DOCX support not available. Please install python-docx.")
    
    def _extract_from_doc(self, content: bytes) -> str:
        """Extract text from old DOC format"""
        # Try to decode as text (basic approach)
        return content.decode('utf-8', errors='ignore')
    
    def split_into_question_blocks(self, text: str) -> List[str]:
        """Split text into individual question blocks"""
        lines = text.split('\n')
        question_blocks = []
        current_block = []
        
        # Combine all question number patterns
        combined_pattern = '|'.join(f'({p})' for p in self.QUESTION_NUMBER_PATTERNS)
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if line starts a new question
            if re.match(combined_pattern, line_stripped, re.IGNORECASE):
                if current_block:
                    question_blocks.append('\n'.join(current_block))
                current_block = [line]
            elif current_block:
                current_block.append(line)
        
        # Add last block
        if current_block:
            question_blocks.append('\n'.join(current_block))
        
        return question_blocks
    
    def is_noise_block(self, text: str) -> bool:
        """Check if block is noise (headers, instructions, etc.)"""
        text_stripped = text.strip()
        
        # Too short
        if len(text_stripped) < 10:
            return True
        
        # Check noise patterns
        for pattern in self.NOISE_PATTERNS:
            if re.search(pattern, text_stripped, re.IGNORECASE):
                return True
        
        return False
    
    def detect_question_type(self, text: str) -> str:
        """Intelligently detect question type based on content"""
        text_lower = text.lower()
        
        # Sort by priority (highest first)
        sorted_patterns = sorted(self.QUESTION_TYPE_PATTERNS, key=lambda x: x[2], reverse=True)
        
        for question_type, patterns, priority in sorted_patterns:
            match_count = 0
            for pattern in patterns:
                if re.search(pattern, text_lower, re.IGNORECASE):
                    match_count += 1
            
            # If at least one pattern matches, consider it a match
            # For higher priority types, require stronger evidence
            if priority >= 70 and match_count >= 1:
                return question_type
            elif priority >= 50 and match_count >= 1:
                return question_type
            elif match_count >= 2:  # Lower priority needs multiple matches
                return question_type
        
        # Default to multiple choice
        return 'multiple_choice'
    
    def extract_answer(self, text: str) -> Tuple[Optional[str], str]:
        """Extract answer from question text"""
        for pattern, label in self.ANSWER_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                answer = match.group(1).strip()
                # Remove answer from text
                cleaned_text = re.sub(pattern, '', text, flags=re.IGNORECASE).strip()
                return answer, cleaned_text
        
        return None, text
    
    def extract_options(self, text: str) -> Tuple[List[str], str]:
        """Extract options from question text (for MCQ, multiple select, etc.)"""
        for pattern in self.OPTION_PATTERNS:
            options_found = re.findall(pattern, text)
            if len(options_found) >= 2:  # At least 2 options
                options = [opt[1].strip() for opt in options_found]
                # Remove options from text
                cleaned_text = re.sub(pattern, '', text).strip()
                return options, cleaned_text
        
        return [], text
    
    def clean_question_text(self, text: str) -> str:
        """Clean up question text"""
        # Remove question number
        combined_pattern = '|'.join(f'({p})' for p in self.QUESTION_NUMBER_PATTERNS)
        text = re.sub(combined_pattern, '', text, count=1, flags=re.IGNORECASE).strip()
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove trailing marks indicators
        text = re.sub(r'\s*\(\s*\d+\s*marks?\s*\)\s*$', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*\[\s*\d+\s*marks?\s*\]\s*$', '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def parse_question(self, block: str) -> Optional[Dict]:
        """Parse a single question block into structured data"""
        if self.is_noise_block(block):
            return None
        
        # Extract answer first
        answer, text_without_answer = self.extract_answer(block)
        
        # Detect question type
        question_type = self.detect_question_type(text_without_answer)
        
        # Extract options (for MCQ, multiple select, etc.)
        options, text_without_options = self.extract_options(text_without_answer)
        
        # Clean question text
        question_text = self.clean_question_text(text_without_options)
        
        # Validate
        if not question_text or len(question_text) < 10:
            return None
        
        # For certain types, answer is required
        if question_type in ['multiple_choice', 'true_false', 'multiple_select'] and not answer:
            # Try to infer from marked options
            marked_pattern = r'[\*✓✔]\s*([A-Za-z])[)\.]'
            marked = re.search(marked_pattern, block)
            if marked:
                answer = marked.group(1).upper()
        
        # Build result
        result = {
            'text': question_text,
            'type': question_type,
            'options': options if options else [],
            'answer': answer or ''
        }
        
        return result
    
    def parse_document(self, content: bytes, filename: str) -> Dict:
        """
        Main entry point: Parse entire document and extract all questions
        
        Returns:
            Dict with 'questions', 'count', 'skipped', 'warnings'
        """
        # Extract text
        try:
            text = self.extract_text_from_file(content, filename)
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'questions': [],
                'count': 0
            }
        
        # Split into blocks
        blocks = self.split_into_question_blocks(text)
        
        # Parse each block
        questions = []
        skipped = []
        
        for idx, block in enumerate(blocks, 1):
            parsed = self.parse_question(block)
            if parsed:
                questions.append(parsed)
            else:
                skipped.append(f"Block {idx}: {block[:50]}...")
        
        # Generate warnings
        warnings = []
        if len(questions) < len(blocks) * 0.3:
            warnings.append(
                f"Only {len(questions)} out of {len(blocks)} blocks were recognized as valid questions. "
                "Make sure each question has a clear number (1., Q1, etc.) and proper formatting."
            )
        
        return {
            'success': True,
            'questions': questions,
            'count': len(questions),
            'total_blocks': len(blocks),
            'skipped': len(skipped),
            'warnings': warnings,
            'skipped_details': skipped if self.debug_mode else []
        }


# Global parser instance
parser = AdvancedQuestionParser()
