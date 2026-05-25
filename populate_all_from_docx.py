#!/usr/bin/env python3
"""
Complete DOCX parser - extract ALL schools from ALL tables
"""

import os
import re
import json
from docx import Document
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://quiz_user:quiz_pass123@db:5432/morning_quiz")

PROVINCE_DISTRICTS = {
    'KIGALI CITY': ['GASABO', 'KICUKIRO', 'NYARUGENGE'],
    'EASTERN PROVINCE': ['BUGESERA', 'GATSIBO', 'KAYONZA', 'KIREHE', 'NGOMA', 'RWAMAGANA', 'NYAGATARE'],
    'NORTHERN PROVINCE': ['BURERA', 'GAKENKE', 'GICUMBI', 'MUSANZE', 'RULINDO'],
    'SOUTHERN PROVINCE': ['GISAGARA', 'HUYE', 'KAMONYI', 'MUHANGA', 'NYAMAGABE', 'NYANZA', 'NYARUGURU', 'RUHANGO'],
    'WESTERN PROVINCE': ['KARONGI', 'NGORORERO', 'NYABIHU', 'NYAMASHEKE', 'RUBAVU', 'RUSIZI', 'RUTSIRO']
}

def extract_level(trade_str):
    match = re.search(r'L(\d+)(?:-(\d+))?', trade_str, re.IGNORECASE)
    if match:
        return f"L{match.group(1)}-{match.group(2)}" if match.group(2) else f"L{match.group(1)}"
    return 'L3-5'

def clean_trade_name(trade_str):
    # Remove numbering
    trade_str = re.sub(r'^\s*\d+\.\s*', '', trade_str)
    # Remove level
    trade_str = re.sub(r'\s*L\d+(?:-\d+)?\s*$', '', trade_str, flags=re.IGNORECASE)
    return trade_str.strip()

def parse_docx_all_schools(docx_path):
    doc = Document(docx_path)
    all_schools = []
    
    print(f"📄 Parsing {len(doc.tables)} tables...")
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 5:
            continue
        
        # Look for header row with "District" and "School"
        header_row_idx = None
        for idx in range(min(5, len(table.rows))):
            row_text = ' '.join([cell.text for cell in table.rows[idx].cells])
            if 'District' in row_text and 'School' in row_text:
                header_row_idx = idx
                break
        
        if header_row_idx is None:
            continue
        
        # Get header cells
        header_cells = [cell.text.strip() for cell in table.rows[header_row_idx].cells]
        
        # Find column indices
        district_col = None
        school_col = None
        type_col = None
        trades_col = None
        
        for idx, header in enumerate(header_cells):
            if 'District' in header:
                district_col = idx
            elif 'School' in header:
                school_col = idx
            elif 'Type' in header:
                type_col = idx
            elif 'Accredited' in header and 'Not' not in header:
                trades_col = idx
        
        if district_col is None or school_col is None or trades_col is None:
            continue
        
        print(f"\n📋 Table {table_idx + 1}: Found data (header at row {header_row_idx})")
        
        # Parse data rows (start after header)
        for row_idx in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[row_idx]
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) <= max(district_col, school_col, trades_col):
                continue
            
            district = cells[district_col].upper().strip()
            school_name = cells[school_col].replace('\n', ' ').strip()
            school_type = cells[type_col] if type_col < len(cells) else 'TSS'
            trades_text = cells[trades_col] if trades_col < len(cells) else ''
            
            # Skip invalid rows
            if not district or not school_name or len(school_name) < 3:
                continue
            if district not in [d for districts in PROVINCE_DISTRICTS.values() for d in districts]:
                continue
            
            # Parse trades - split by | or newline
            trades = []
            if trades_text and trades_text not in ['None', '']:
                # Split by | or newline
                trade_items = re.split(r'[|\n]+', trades_text)
                for trade_item in trade_items:
                    trade_item = trade_item.strip()
                    # Remove leading numbers
                    trade_item = re.sub(r'^\d+\.\s*', '', trade_item)
                    if trade_item and len(trade_item) > 5 and trade_item not in ['None']:
                        trades.append(trade_item)
            
            if trades:
                school_data = {
                    'district': district,
                    'school_name': school_name,
                    'type': school_type,
                    'trades': trades
                }
                all_schools.append(school_data)
                print(f"  ✓ {school_name} ({district}) - {len(trades)} trades")
    
    return all_schools

def populate_all():
    engine = create_engine(DATABASE_URL)
    Session = sessionmaker(bind=engine)
    db = Session()
    
    print("\n" + "="*70)
    print("COMPLETE DATABASE POPULATION FROM DOCX")
    print("="*70)
    
    try:
        # Fix sequences
        print("\n🔧 Fixing sequences...")
        for seq in ['provinces_id_seq', 'districts_id_seq', 'schools_id_seq', 'trades_id_seq', 'school_trades_id_seq']:
            table = seq.replace('_id_seq', '')
            db.execute(text(f"SELECT setval('{seq}', (SELECT COALESCE(MAX(id), 0) FROM {table}) + 1)"))
        db.commit()
        print("✅ Sequences fixed")
        
        # Get districts
        print("\n📍 Loading districts...")
        district_ids = {}
        for province, districts in PROVINCE_DISTRICTS.items():
            for district in districts:
                result = db.execute(text("SELECT id FROM districts WHERE UPPER(name) = :name"), {"name": district})
                row = result.fetchone()
                if row:
                    district_ids[district] = row[0]
        print(f"✅ {len(district_ids)} districts loaded")
        
        # Parse DOCX
        docx_path = '/app/LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC_2024-2025_lJhq.docx'
        schools_data = parse_docx_all_schools(docx_path)
        print(f"\n✅ Extracted {len(schools_data)} schools from DOCX")
        
        # Populate
        schools_created = 0
        schools_updated = 0
        trades_created = 0
        relationships_created = 0
        
        print("\n🔄 Populating database...")
        
        debug_count = 0
        for school_data in schools_data:
            district = school_data['district']
            
            if district not in district_ids:
                continue
            
            school_name = school_data['school_name']
            school_type = school_data['type']
            
            # Get or create school
            result = db.execute(text("""
                SELECT id FROM schools 
                WHERE name = :name AND district_id = :did
            """), {"name": school_name, "did": district_ids[district]})
            
            row = result.fetchone()
            if row:
                school_id = row[0]
                schools_updated += 1
            else:
                result = db.execute(text("""
                    INSERT INTO schools (district_id, name, school_type, is_active)
                    VALUES (:did, :name, :type, true)
                    RETURNING id
                """), {"did": district_ids[district], "name": school_name, "type": school_type})
                school_id = result.fetchone()[0]
                schools_created += 1
                db.commit()
            
            # Debug first school
            if debug_count == 0:
                print(f"\n  DEBUG: Processing {school_name} (ID: {school_id})")
                print(f"  Trades to add: {school_data['trades']}")
            
            # Add trades
            for trade_full in school_data['trades']:
                level = extract_level(trade_full)
                trade_name = clean_trade_name(trade_full)
                
                if debug_count == 0:
                    print(f"    Trade: '{trade_full}' -> Name: '{trade_name}', Level: '{level}'")
                
                if len(trade_name) < 3:
                    continue
                
                # Get or create trade
                result = db.execute(text("SELECT id FROM trades WHERE name = :name"), {"name": trade_name})
                row = result.fetchone()
                
                if row:
                    trade_id = row[0]
                else:
                    try:
                        result = db.execute(text("""
                            INSERT INTO trades (name, code)
                            VALUES (:name, :code)
                            RETURNING id
                        """), {"name": trade_name, "code": trade_name[:10].upper().replace(' ', '_')})
                        trade_id = result.fetchone()[0]
                        trades_created += 1
                        db.commit()
                    except:
                        db.rollback()
                        result = db.execute(text("SELECT id FROM trades WHERE name = :name"), {"name": trade_name})
                        row = result.fetchone()
                        if not row:
                            continue
                        trade_id = row[0]
                
                # Add relationship
                try:
                    result = db.execute(text("""
                        SELECT id FROM school_trades 
                        WHERE school_id = :sid AND trade_id = :tid
                    """), {"sid": school_id, "tid": trade_id})
                    
                    if not result.fetchone():
                        db.execute(text("""
                            INSERT INTO school_trades (school_id, trade_id, levels_offered, is_active)
                            VALUES (:sid, :tid, CAST(:levels AS json), true)
                        """), {"sid": school_id, "tid": trade_id, "levels": json.dumps([level])})
                        relationships_created += 1
                        db.commit()
                        if debug_count == 0:
                            print(f"      ✅ Added relationship")
                except Exception as e:
                    if debug_count == 0:
                        print(f"      ❌ Error: {e}")
                    db.rollback()
            
            debug_count += 1
        
        # Final stats
        print("\n" + "="*70)
        print("✅ POPULATION COMPLETE!")
        print("="*70)
        print(f"Schools created: {schools_created}")
        print(f"Schools updated: {schools_updated}")
        print(f"Trades created: {trades_created}")
        print(f"Relationships created: {relationships_created}")
        
        result = db.execute(text("SELECT COUNT(*) FROM schools"))
        total_schools = result.fetchone()[0]
        
        result = db.execute(text("SELECT COUNT(*) FROM trades"))
        total_trades = result.fetchone()[0]
        
        result = db.execute(text("SELECT COUNT(*) FROM school_trades"))
        total_relationships = result.fetchone()[0]
        
        result = db.execute(text("""
            SELECT COUNT(*) FROM schools s 
            WHERE EXISTS (SELECT 1 FROM school_trades st WHERE st.school_id = s.id)
        """))
        schools_with_trades = result.fetchone()[0]
        
        print(f"\n📊 FINAL DATABASE STATS:")
        print(f"   Total schools: {total_schools}")
        print(f"   Total trades: {total_trades}")
        print(f"   Total relationships: {total_relationships}")
        print(f"   Schools WITH trades: {schools_with_trades}")
        print(f"   Schools WITHOUT trades: {total_schools - schools_with_trades}")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    populate_all()
