"""
Generate Test Quiz Document - All 13 Question Types
This creates a .doc file that teachers can upload to test all question types
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color):
    """Add colored heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_table_border(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_test_quiz():
    """Generate comprehensive test quiz document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('TVET Quiz System - Complete Test Quiz', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('All 13 Question Types - Teacher Testing Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Instructions
    instructions = doc.add_paragraph()
    instructions.add_run('Instructions: ').bold = True
    instructions.add_run('Upload this document using the AI Document Parser in the Question Types page. ')
    instructions.add_run('This will test all 13 question types in your system.')
    
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    # Question 1: Multiple Choice
    add_heading_with_color(doc, '1. Multiple Choice Question', 2, (37, 99, 235))
    doc.add_paragraph('What is the capital city of Rwanda?')
    doc.add_paragraph('A) Kigali', style='List Bullet')
    doc.add_paragraph('B) Nairobi', style='List Bullet')
    doc.add_paragraph('C) Kampala', style='List Bullet')
    doc.add_paragraph('D) Dar es Salaam', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Correct Answer: A) Kigali').bold = True
    p.add_run('\nPoints: 2')
    doc.add_paragraph()
    
    # Question 2: Multiple Select
    add_heading_with_color(doc, '2. Multiple Select Question (Checkboxes)', 2, (34, 197, 94))
    doc.add_paragraph('Which of the following are programming languages? (Select all that apply)')
    doc.add_paragraph('☐ Python', style='List Bullet')
    doc.add_paragraph('☐ Microsoft Word', style='List Bullet')
    doc.add_paragraph('☐ Java', style='List Bullet')
    doc.add_paragraph('☐ Google Chrome', style='List Bullet')
    doc.add_paragraph('☐ JavaScript', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Correct Answers: Python, Java, JavaScript').bold = True
    p.add_run('\nPoints: 3')
    doc.add_paragraph()
    
    # Question 3: Dropdown Select
    add_heading_with_color(doc, '3. Dropdown Select Question', 2, (168, 85, 247))
    doc.add_paragraph('What is the file extension for Python files?')
    doc.add_paragraph('Options: .txt, .py, .java, .html')
    p = doc.add_paragraph()
    p.add_run('Correct Answer: .py').bold = True
    p.add_run('\nPoints: 1')
    doc.add_paragraph()
    
    # Question 4: True/False
    add_heading_with_color(doc, '4. True/False Question', 2, (234, 179, 8))
    doc.add_paragraph('HTML stands for HyperText Markup Language.')
    p = doc.add_paragraph()
    p.add_run('Correct Answer: True').bold = True
    p.add_run('\nPoints: 1')
    doc.add_paragraph()
    
    # Question 5: Short Answer
    add_heading_with_color(doc, '5. Short Answer Question', 2, (99, 102, 241))
    doc.add_paragraph('What does CPU stand for?')
    p = doc.add_paragraph()
    p.add_run('Expected Answer: Central Processing Unit').bold = True
    p.add_run('\nPoints: 2')
    doc.add_paragraph()
    
    # Question 6: Essay
    add_heading_with_color(doc, '6. Essay Question (Paragraph)', 2, (236, 72, 153))
    doc.add_paragraph('Explain the importance of cybersecurity in modern businesses. (Minimum 100 words)')
    p = doc.add_paragraph()
    p.add_run('Grading Notes: Look for key concepts - data protection, privacy, threats, prevention').bold = True
    p.add_run('\nPoints: 5')
    doc.add_paragraph()
    
    # Question 7: Linear Scale
    add_heading_with_color(doc, '7. Linear Scale Question', 2, (20, 184, 166))
    doc.add_paragraph('On a scale of 1 to 10, how confident are you with database management?')
    doc.add_paragraph('Scale: 1 (Not Confident) to 10 (Very Confident)')
    p = doc.add_paragraph()
    p.add_run('Points: 1').bold = True
    doc.add_paragraph()
    
    # Question 8: Fill in the Blanks
    add_heading_with_color(doc, '8. Fill in the Blanks Question', 2, (249, 115, 22))
    doc.add_paragraph('The ___ is the brain of the computer, while ___ stores data temporarily.')
    p = doc.add_paragraph()
    p.add_run('Correct Answers: CPU, RAM').bold = True
    p.add_run('\nPoints: 2')
    doc.add_paragraph()
    
    # Question 9: Matching Pairs
    add_heading_with_color(doc, '9. Matching Pairs Question', 2, (6, 182, 212))
    doc.add_paragraph('Match the programming language with its primary use:')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    add_table_border(table)
    
    # Header
    table.rows[0].cells[0].text = 'Language'
    table.rows[0].cells[1].text = 'Primary Use'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data
    table.rows[1].cells[0].text = 'Python'
    table.rows[1].cells[1].text = 'Data Science & AI'
    table.rows[2].cells[0].text = 'JavaScript'
    table.rows[2].cells[1].text = 'Web Development'
    table.rows[3].cells[0].text = 'SQL'
    table.rows[3].cells[1].text = 'Database Management'
    
    p = doc.add_paragraph()
    p.add_run('Points: 3').bold = True
    doc.add_paragraph()
    
    # Question 10: Drag & Drop Ordering
    add_heading_with_color(doc, '10. Drag & Drop Ordering Question', 2, (139, 92, 246))
    doc.add_paragraph('Arrange the following steps of software development in correct order:')
    doc.add_paragraph('1. Testing', style='List Number')
    doc.add_paragraph('2. Design', style='List Number')
    doc.add_paragraph('3. Requirements Analysis', style='List Number')
    doc.add_paragraph('4. Implementation', style='List Number')
    doc.add_paragraph('5. Deployment', style='List Number')
    p = doc.add_paragraph()
    p.add_run('Correct Order: Requirements Analysis → Design → Implementation → Testing → Deployment').bold = True
    p.add_run('\nPoints: 3')
    doc.add_paragraph()
    
    # Question 11: Code Writing
    add_heading_with_color(doc, '11. Code Writing Question (Python)', 2, (75, 85, 99))
    doc.add_paragraph('Write a Python function that calculates the factorial of a number.')
    
    # Code block
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'def factorial(n):\n'
        '    if n == 0 or n == 1:\n'
        '        return 1\n'
        '    return n * factorial(n - 1)\n\n'
        'print(factorial(5))  # Output: 120'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    code_para.paragraph_format.left_indent = Inches(0.5)
    
    p = doc.add_paragraph()
    p.add_run('Language: Python').bold = True
    p.add_run('\nPoints: 5')
    doc.add_paragraph()
    
    # Question 12: SQL Query
    add_heading_with_color(doc, '12. SQL Query Question', 2, (59, 130, 246))
    doc.add_paragraph('Write an SQL query to select all students with grades above 80 from the "students" table.')
    
    # SQL code block
    sql_para = doc.add_paragraph()
    sql_run = sql_para.add_run(
        'SELECT * FROM students\n'
        'WHERE grade > 80\n'
        'ORDER BY grade DESC;'
    )
    sql_run.font.name = 'Courier New'
    sql_run.font.size = Pt(10)
    sql_para.paragraph_format.left_indent = Inches(0.5)
    
    p = doc.add_paragraph()
    p.add_run('Points: 4').bold = True
    doc.add_paragraph()
    
    # Question 13: Multi-Grid (Matrix)
    add_heading_with_color(doc, '13. Multi-Grid Question (Matrix)', 2, (16, 185, 129))
    doc.add_paragraph('Rate your proficiency in the following areas (1=Beginner, 5=Expert):')
    
    grid_table = doc.add_table(rows=4, cols=6)
    grid_table.style = 'Light Grid Accent 1'
    add_table_border(grid_table)
    
    # Header row
    grid_table.rows[0].cells[0].text = 'Skill'
    for i, rating in enumerate(['1', '2', '3', '4', '5'], 1):
        grid_table.rows[0].cells[i].text = rating
        grid_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Skills
    skills = ['Programming', 'Database Design', 'Networking']
    for i, skill in enumerate(skills, 1):
        grid_table.rows[i].cells[0].text = skill
        for j in range(1, 6):
            grid_table.rows[i].cells[j].text = '○'
            grid_table.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Points: 3').bold = True
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph()
    footer.add_run('End of Test Quiz - Total Points: 37').bold = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer2 = doc.add_paragraph('Generated by TVET Quiz System')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.runs[0].font.size = Pt(10)
    footer2.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    # Save document
    filename = 'TVET_Test_Quiz_All_13_Questions.docx'
    doc.save(filename)
    print(f'✅ Successfully created: {filename}')
    print(f'📄 Document contains all 13 question types')
    print(f'📊 Total points: 37')
    print(f'🎯 Ready to upload via AI Document Parser!')
    return filename

if __name__ == '__main__':
    try:
        create_test_quiz()
    except ImportError:
        print('❌ Error: python-docx library not installed')
        print('📦 Install it with: pip install python-docx')
    except Exception as e:
        print(f'❌ Error creating document: {e}')
