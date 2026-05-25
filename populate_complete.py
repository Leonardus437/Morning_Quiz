#!/usr/bin/env python3
"""
Complete population from DOCX - parse ALL tables and populate ALL schools with trades
"""

import os
import re
import json
from docx import Document
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://quiz_user:quiz_pass123@db:5432/morning_quiz")

PROVINCE_DISTRICTS = {
    'KIGALI CITY': ['GASABO', 'KICUKIRO', 'NYARUGENGE'],
    'EASTERN PROVINCE': ['BUGESERA', 'GATSIBO', 'KAYONZA', 'KIREHE', 'NGOMA', 'RWAMAGANA', 'NYAGATARE'],
    'NORTHERN PROVINCE': ['BURERA', 'GAKENKE', 'GICUMBI', 'MUSANZE', 'RULINDO'],
    'SOUTHERN PROVINCE': ['GISAGARA', 'HUYE', 'KAMONYI', 'MUHANGA', 'NYAMAGABE', 'NYANZA', 'NYARUGURU', 'RUHANGO'],
    'WESTERN PROVINCE': ['KARONGI', 'NGORORERO', 'NYABIHU', 'NYAMASHEKE', 'RUBAVU', 'RUSIZI', 'RUTSIRO']
}

def extract_level(trade_str):
    """Extract level from trade string"""
    match = re.search(r'L(\d+)(?:-(\d+))?', trade_str, re.IGNORECASE)
    if match:
        if match.group(2):
            return f"L{match.group(1)}-{match.group(2)}"
        return f"L{match.group(1)}"
    return 'L3-5'

def clean_trade_name(trade_str):
    """Remove level and numbering from trade name"""
    # Remove leading numbers like "1. " or "2. "
    trade_str = re.sub(r'^\s*\d+\.\s*', '', trade_str)
    # Remove level suffix
    trade_str = re.sub(r'\s*L\d+(?:-\d+)?\s*$', '', trade_str, flags=re.IGNORECASE)
    return trade_str.strip()

def parse_all_tables(docx_path):
    """Parse ALL tables from DOCX"""
    doc = Document(docx_path)
    all_schools = []
    
    print(f"📄 Found {len(doc.tables)} tables in DOCX")
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        
        # Check if this table has school data
        header_row = table.rows[0]
        header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
        
        # Look for tables with "District" and "School" headers
        if 'District' not in header_text and 'School' not in header_text:
            continue
        
        print(f"\n📋 Processing Table {table_idx + 1} ({len(table.rows)} rows)")
        
        # Parse each row
        for row_idx, row in enumerate(table.rows[1:], start=1):  # Skip header
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) < 7:
                continue
            
            # Extract data
            sn = cells[0]
            district = cells[1]
            school_name = cells[2]
            school_type = cells[3] if len(cells) > 3 else 'TSS'
            status = cells[4] if len(cells) > 4 else 'Public'
            boarding = cells[5] if len(cells) > 5 else 'Day'
            
            # Trades are in column 6 or later
            trades_text = ''
            for i in range(6, len(cells)):
                if cells[i] and cells[i] not in ['None', '']:
                    trades_text += ' ' + cells[i]
            
            # Skip if no valid data
            if not district or not school_name or len(school_name) < 3:
                continue
            
            # Skip header-like rows
            if 'District' in school_name or 'School Name' in school_name:
                continue
            
            # Parse trades from text
            trades = []
            if trades_text:
                # Split by newlines or numbered items
                trade_lines = re.split(r'\n|\d+\.\s+', trades_text)
                for trade_line in trade_lines:
                    trade_line = trade_line.strip()
                    if trade_line and len(trade_line) > 3 and trade_line not in ['None', 'Boarding', 'Day', 'Mixed']:
                        trades.append(trade_line)
            
            if trades:
                school_data = {
                    'sn': sn,
                    'district': district.upper().strip(),
                    'school_name': school_name.strip(),
                    'type': school_type.strip(),
                    'status': status.strip(),
                    'boarding': boarding.strip(),
                    'trades': trades
                }
                all_schools.append(school_data)
                print(f"  ✓ {school_name} ({district}) - {len(trades)} trades")
    
    return all_schools

def populate_database():
    engine = create_engine(DATABASE_URL)
    Session = sessionmaker(bind=engine)
    db = Session()
    
    print("\n" + "="*70)
    print("POPULATING DATABASE FROM DOCX")
    print("="*70)
    
    try:
        # Fix sequences first
        print("\n🔧 Fixing sequences...")
        db.execute(text("SELECT setval('provinces_id_seq', (SELECT COALESCE(MAX(id), 0) FROM provinces) + 1)"))
        db.execute(text("SELECT setval('districts_id_seq', (SELECT COALESCE(MAX(id), 0) FROM districts) + 1)"))
        db.execute(text("SELECT setval('schools_id_seq', (SELECT COALESCE(MAX(id), 0) FROM schools) + 1)"))
        db.execute(text("SELECT setval('trades_id_seq', (SELECT COALESCE(MAX(id), 0) FROM trades) + 1)"))
        db.execute(text("SELECT setval('school_trades_id_seq', (SELECT COALESCE(MAX(id), 0) FROM school_trades) + 1)"))
        db.commit()
        print("✅ Sequences fixed")
        
        # Get district mapping
        print("\n📍 Loading districts...")
        district_ids = {}
        for province, districts in PROVINCE_DISTRICTS.items():
            for district in districts:
                result = db.execute(text("SELECT id FROM districts WHERE UPPER(name) = :name"), {"name": district})
                row = result.fetchone()
                if row:
                    district_ids[district] = row[0]
        print(f"✅ Found {len(district_ids)} districts")
        
        # Parse DOCX
        docx_path = '/app/LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC_2024-2025_lJhq.docx'
        schools_data = parse_all_tables(docx_path)
        print(f"\n✅ Parsed {len(schools_data)} schools from DOCX")
        
        # Process each school
        schools_created = 0
        schools_updated = 0
        trades_created = 0
        relationships_created = 0
        
        for school_data in schools_data:
            district = school_data['district']
            
            if district not in district_ids:
                print(f"⚠️  District not found: {district}")
                continue
            
            school_name = school_data['school_name']
            school_type = school_data['type']
            
            # Find or create school
            result = db.execute(text("""
                SELECT id FROM schools 
                WHERE name = :name AND district_id = :did
            """), {"name": school_name, "did": district_ids[district]})
            
            row = result.fetchone()
            if row:
                school_id = row[0]
                schools_updated += 1
            else:
                # Create school
                result = db.execute(text("""
                    INSERT INTO schools (district_id, name, school_type, is_active)
                    VALUES (:did, :name, :type, true)
                    RETURNING id
                """), {"did": district_ids[district], "name": school_name, "type": school_type})
                school_id = result.fetchone()[0]
                schools_created += 1
                db.commit()
            
            # Process trades
            for trade_full in school_data['trades']:
                level = extract_level(trade_full)
                trade_name = clean_trade_name(trade_full)
                
                if len(trade_name) < 3:
                    continue
                
                # Get or create trade
                result = db.execute(text("SELECT id FROM trades WHERE name = :name"), {"name": trade_name})
                row = result.fetchone()
                
                if row:
                    trade_id = row[0]
                else:
                    try:
                        result = db.execute(text("""
                            INSERT INTO trades (name, code)
                            VALUES (:name, :code)
                            RETURNING id
                        """), {"name": trade_name, "code": trade_name[:10].upper().replace(' ', '_')})
                        trade_id = result.fetchone()[0]
                        trades_created += 1
                        db.commit()
                    except:
                        db.rollback()
                        result = db.execute(text("SELECT id FROM trades WHERE name = :name"), {"name": trade_name})
                        row = result.fetchone()
                        if not row:
                            continue
                        trade_id = row[0]
                
                # Add relationship
                try:
                    result = db.execute(text("""
                        SELECT id FROM school_trades 
                        WHERE school_id = :sid AND trade_id = :tid
                    """), {"sid": school_id, "tid": trade_id})
                    
                    if not result.fetchone():
                        db.execute(text("""
                            INSERT INTO school_trades (school_id, trade_id, levels_offered, is_active)
                            VALUES (:sid, :tid, :levels::json, true)
                        """), {"sid": school_id, "tid": trade_id, "levels": json.dumps([level])})
                        relationships_created += 1
                        db.commit()
                except Exception as e:
                    db.rollback()
                    continue
        
        # Final stats
        print("\n" + "="*70)
        print("POPULATION COMPLETE")
        print("="*70)
        print(f"✅ Schools created: {schools_created}")
        print(f"✅ Schools updated: {schools_updated}")
        print(f"✅ Trades created: {trades_created}")
        print(f"✅ Relationships created: {relationships_created}")
        
        result = db.execute(text("SELECT COUNT(*) FROM schools"))
        print(f"\n📊 Total schools: {result.fetchone()[0]}")
        
        result = db.execute(text("SELECT COUNT(*) FROM trades"))
        print(f"📊 Total trades: {result.fetchone()[0]}")
        
        result = db.execute(text("SELECT COUNT(*) FROM school_trades"))
        print(f"📊 Total relationships: {result.fetchone()[0]}")
        
        result = db.execute(text("""
            SELECT COUNT(*) FROM schools s 
            WHERE EXISTS (SELECT 1 FROM school_trades st WHERE st.school_id = s.id)
        """))
        print(f"📊 Schools WITH trades: {result.fetchone()[0]}")
        
        result = db.execute(text("""
            SELECT COUNT(*) FROM schools s 
            WHERE NOT EXISTS (SELECT 1 FROM school_trades st WHERE st.school_id = s.id)
        """))
        print(f"⚠️  Schools WITHOUT trades: {result.fetchone()[0]}")
        
        print("\n✅ DATABASE FULLY POPULATED!")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    populate_database()
