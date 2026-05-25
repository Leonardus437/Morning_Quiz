#!/usr/bin/env python3
"""
Precise DOCX analyzer - understand exact structure
"""
from docx import Document

doc = Document('/app/schools.docx')

print(f"Total tables: {len(doc.tables)}")
print("\n" + "="*100)

# Analyze each table
for tidx, table in enumerate(doc.tables):
    print(f"\nTABLE {tidx}: {len(table.rows)} rows x {len(table.columns)} columns")
    print("-"*100)
    
    # Print first 5 rows to understand structure
    for ridx in range(min(5, len(table.rows))):
        cells = [c.text.strip() for c in table.rows[ridx].cells]
        print(f"\nRow {ridx} ({len(cells)} cells):")
        for cidx, cell in enumerate(cells):
            if cell:  # Only print non-empty cells
                print(f"  [{cidx}]: {cell[:80]}")
    
    if tidx >= 2:  # Only analyze first 3 tables in detail
        break

print("\n" + "="*100)
print("SEARCHING FOR MUHANGA DISTRICT")
print("="*100)

# Find Muhanga district schools
muhanga_schools = []
for tidx, table in enumerate(doc.tables):
    for ridx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if len(cells) > 1:
            district = cells[1] if len(cells) > 1 else ""
            school = cells[2] if len(cells) > 2 else ""
            
            if 'MUHANGA' in district.upper():
                print(f"\nTable {tidx}, Row {ridx}:")
                for i, cell in enumerate(cells[:8]):
                    print(f"  [{i}]: {cell[:100]}")
                muhanga_schools.append(school)

print(f"\nTotal schools found in MUHANGA: {len(muhanga_schools)}")
