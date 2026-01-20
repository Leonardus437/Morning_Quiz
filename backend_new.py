from fastapi import FastAPI, Depends, HTTPException, status, Request, File, UploadFile
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy import create_engine, Column, Integer, String, Text, Boolean, DateTime, JSON, ForeignKey, Float
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from datetime import datetime, timedelta
import hashlib
import secrets
from jose import JWTError, jwt
import os
import json
import io
import re

# Try to import optional dependencies for file processing
try:
    import PyPDF2
    import docx
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    import requests
    UPLOAD_ENABLED = True
except ImportError:
    UPLOAD_ENABLED = False
    print("⚠️ Optional dependencies not installed. File upload features disabled.")

# Database setup
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://quiz_user:quiz_pass123@db:5432/morning_quiz")
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Security
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-change-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 1440  # 24 hours

security = HTTPBearer()

app = FastAPI(title="TVET/TSS Quiz API - Offline First")

# Add CORS middleware for offline functionality
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Note: Some auxiliary endpoints were drafted in separate modules
# (e.g., backend/admin_excel_endpoint.py, backend/student_upload_endpoint.py)
# but not registered with the FastAPI app. The canonical implementation for
# /admin/upload-students-excel is provided below and wired directly to this app.

# Models
class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    password_hash = Column(String)
    role = Column(String, default="student")
    full_name = Column(String)
    department = Column(String)  # For students
    level = Column(String)  # For students
    departments = Column(JSON)  # For teachers (multiple departments)
    is_class_teacher = Column(Boolean, default=False)  # For teachers
    created_at = Column(DateTime, default=datetime.utcnow)

class Question(Base):
    __tablename__ = "questions"
    id = Column(Integer, primary_key=True, index=True)
    question_text = Column(Text)
    question_type = Column(String)
    options = Column(JSON)
    correct_answer = Column(String)
    points = Column(Integer, default=1)
    department = Column(String)
    level = Column(String)
    lesson_id = Column(Integer, ForeignKey("lessons.id"))
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class Quiz(Base):
    __tablename__ = "quizzes"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    description = Column(Text)
    scheduled_time = Column(DateTime)
    duration_minutes = Column(Integer, default=30)
    question_time_seconds = Column(Integer, default=60)
    countdown_started_at = Column(DateTime)
    is_active = Column(Boolean, default=False)
    department = Column(String)
    level = Column(String)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class QuizQuestion(Base):
    __tablename__ = "quiz_questions"
    id = Column(Integer, primary_key=True, index=True)
    quiz_id = Column(Integer, ForeignKey("quizzes.id"))
    question_id = Column(Integer, ForeignKey("questions.id"))
    question_order = Column(Integer)

class QuizAttempt(Base):
    __tablename__ = "quiz_attempts"
    id = Column(Integer, primary_key=True, index=True)
    quiz_id = Column(Integer, ForeignKey("quizzes.id"))
    user_id = Column(Integer, ForeignKey("users.id"))
    score = Column(Integer, default=0)
    total_questions = Column(Integer)
    answers = Column(JSON)
    started_at = Column(DateTime, default=datetime.utcnow)
    completed_at = Column(DateTime)

class StudentAnswer(Base):
    __tablename__ = "student_answers"
    id = Column(Integer, primary_key=True, index=True)
    attempt_id = Column(Integer, ForeignKey("quiz_attempts.id"))
    question_id = Column(Integer, ForeignKey("questions.id"))
    student_answer = Column(String)
    is_correct = Column(Boolean)

class Lesson(Base):
    __tablename__ = "lessons"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    code = Column(String, unique=True)
    description = Column(Text)
    department = Column(String)
    level = Column(String)
    classification = Column(String)  # Core, Specific, General
    is_active = Column(Boolean, default=True)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class ClassTeacher(Base):
    __tablename__ = "class_teachers"
    id = Column(Integer, primary_key=True, index=True)
    teacher_id = Column(Integer, ForeignKey("users.id"))
    department = Column(String)
    level = Column(String)
    assigned_by = Column(Integer, ForeignKey("users.id"))
    assigned_at = Column(DateTime, default=datetime.utcnow)

class Schedule(Base):
    __tablename__ = "schedules"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    description = Column(Text)
    scheduled_date = Column(DateTime)
    departments = Column(JSON)
    levels = Column(JSON)
    file_data = Column(Text)
    file_name = Column(String)
    file_type = Column(String)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class Announcement(Base):
    __tablename__ = "announcements"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    content = Column(Text)
    priority = Column(String, default="normal")
    departments = Column(JSON)
    levels = Column(JSON)
    is_active = Column(Boolean, default=True)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class TeacherLesson(Base):
    __tablename__ = "teacher_lessons"
    id = Column(Integer, primary_key=True, index=True)
    teacher_id = Column(Integer, ForeignKey("users.id"))
    lesson_id = Column(Integer, ForeignKey("lessons.id"))
    assigned_by = Column(Integer, ForeignKey("users.id"))
    assigned_at = Column(DateTime, default=datetime.utcnow)

class Notification(Base):
    __tablename__ = "notifications"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    title = Column(String)
    message = Column(Text)
    type = Column(String, default="info")
    is_read = Column(Boolean, default=False)
    created_at = Column(DateTime, default=datetime.utcnow)

class QuizSchedule(Base):
    __tablename__ = "quiz_schedules"
    id = Column(Integer, primary_key=True, index=True)
    quiz_id = Column(Integer, ForeignKey("quizzes.id"))
    scheduled_date = Column(DateTime)
    start_time = Column(String)
    end_time = Column(String)
    target_departments = Column(JSON)
    target_levels = Column(JSON)
    is_active = Column(Boolean, default=True)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class QuestionBank(Base):
    __tablename__ = "question_banks"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String)
    description = Column(Text)
    file_path = Column(String)
    total_questions = Column(Integer, default=0)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class SessionPlan(Base):
    __tablename__ = "session_plans"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    lesson_id = Column(Integer, ForeignKey("lessons.id"))
    department = Column(String)
    level = Column(String)
    duration_minutes = Column(Integer, default=90)
    learning_objectives = Column(JSON)
    teaching_methods = Column(JSON)
    resources_required = Column(JSON)
    assessment_methods = Column(JSON)
    content_outline = Column(Text)
    activities = Column(JSON)
    homework_assignment = Column(Text)
    reflection_notes = Column(Text)
    generated_content = Column(Text)
    template_used = Column(String, default="RTB_Standard")
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class SchemeOfWork(Base):
    __tablename__ = "schemes_of_work"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    lesson_id = Column(Integer, ForeignKey("lessons.id"))
    department = Column(String)
    level = Column(String)
    academic_year = Column(String)
    term = Column(String)
    total_weeks = Column(Integer, default=12)
    weekly_breakdown = Column(JSON)
    learning_outcomes = Column(JSON)
    assessment_schedule = Column(JSON)
    resources_list = Column(JSON)
    generated_content = Column(Text)
    template_used = Column(String, default="RTB_Standard")
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class QuizResults(Base):
    __tablename__ = "quiz_results"
    id = Column(Integer, primary_key=True, index=True)
    quiz_id = Column(Integer, ForeignKey("quizzes.id"))
    user_id = Column(Integer, ForeignKey("users.id"))
    score = Column(Integer, default=0)
    total_questions = Column(Integer)
    submitted_at = Column(DateTime, default=datetime.utcnow)

class LessonPlanTemplate(Base):
    __tablename__ = "lesson_plan_templates"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String)
    template_type = Column(String)
    template_content = Column(Text)
    is_rtb_official = Column(Boolean, default=False)
    department = Column(String)
    level = Column(String)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.utcnow)

class AIGenerationLog(Base):
    __tablename__ = "ai_generation_logs"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    generation_type = Column(String)
    input_parameters = Column(JSON)
    generated_content = Column(Text)
    processing_time = Column(Float)
    success = Column(Boolean, default=True)
    error_message = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)

# Pydantic models
class UserLogin(BaseModel):
    username: str
    password: str

class UserCreate(BaseModel):
    username: str
    password: str
    full_name: str
    role: str = "student"
    department: Optional[str] = None
    level: Optional[str] = None
    departments: Optional[List[str]] = None

class UserResponse(BaseModel):
    id: int
    username: str
    role: str
    full_name: str
    department: Optional[str] = None
    level: Optional[str] = None
    departments: Optional[List[str]] = None

class Token(BaseModel):
    access_token: str
    token_type: str
    user: Optional[Dict] = None

class QuestionCreate(BaseModel):
    question_text: str
    question_type: str
    options: Optional[Dict] = None
    correct_answer: str
    points: int = 1
    department: str
    level: str
    lesson_id: Optional[int] = None

class QuestionResponse(BaseModel):
    id: int
    question_text: str
    question_type: str
    options: Optional[Any] = None  # Can be list or dict
    correct_answer: str
    points: int = 1
    department: str
    level: str
    lesson_id: Optional[int] = None
    created_by: int
    created_at: datetime

class QuizCreate(BaseModel):
    title: str
    description: Optional[str] = None
    scheduled_time: Optional[datetime] = None
    duration_minutes: int = 30
    question_time_seconds: int = 60
    department: str
    level: str
    question_ids: List[int] = []

class QuizResponse(QuizCreate):
    id: int
    is_active: bool
    created_by: int
    created_at: datetime

class UserRegister(BaseModel):
    username: str
    password: str
    full_name: str
    role: str = "student"
    department: str = None
    level: str = None
    departments: List[str] = []
    is_class_teacher: bool = False
    class_department: str = None
    class_level: str = None

class BulkQuestionCreate(BaseModel):
    questions: List[QuestionCreate]

class QuizAnswer(BaseModel):
    question_id: int
    answer: str

class QuizSubmission(BaseModel):
    quiz_id: int
    answers: List[QuizAnswer]

class ScheduleCreate(BaseModel):
    title: str
    description: str = ""
    scheduled_date: datetime
    departments: List[str]
    levels: List[str]

class AnnouncementCreate(BaseModel):
    title: str
    content: str

# Dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Use bcrypt for secure password hashing
try:
    import bcrypt
    BCRYPT_AVAILABLE = True
except ImportError:
    BCRYPT_AVAILABLE = False

def hash_password_simple(password: str) -> str:
    if BCRYPT_AVAILABLE:
        # Use bcrypt for new passwords
        salt = bcrypt.gensalt()
        return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')
    else:
        # Fallback to SHA256 with salt
        salt = secrets.token_hex(16)
        password_hash = hashlib.sha256((password + salt).encode()).hexdigest()
        return f"{salt}:{password_hash}"

def verify_password_simple(password: str, hashed: str) -> bool:
    try:
        # Check if it's bcrypt format (starts with $2b$)
        if hashed.startswith('$2b$') and BCRYPT_AVAILABLE:
            return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
        # Check if it's SHA256 format (salt:hash)
        elif ':' in hashed:
            salt, password_hash = hashed.split(':', 1)
            computed_hash = hashlib.sha256((password + salt).encode()).hexdigest()
            return computed_hash == password_hash
        else:
            return False
    except Exception as e:
        print(f"Password verification error: {e}")
        return False

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(request: Request, db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    token = request.headers.get("authorization")
    if not token:
        raise credentials_exception
    try:
        token = token.split(" ")[1]  # Remove "Bearer " prefix
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = db.query(User).filter(User.username == username).first()
    if user is None:
        raise credentials_exception
    return user

# Authentication routes


@app.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate, db: Session = Depends(get_db)):
    # Check if user already exists
    existing_user = db.query(User).filter(User.username == user_data.username).first()
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Username already registered",
        )
    
    # Create new user
    hashed_password = hash_password_simple(user_data.password)
    db_user = User(
        username=user_data.username,
        password_hash=hashed_password,
        full_name=user_data.full_name,
        role=user_data.role,
        department=user_data.department,
        level=user_data.level,
        departments=user_data.departments
    )
    
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    
    # Create token
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": db_user.username, "role": db_user.role}, expires_delta=access_token_expires
    )
    
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/auth/me", response_model=UserResponse)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return current_user

# Question routes
@app.post("/questions", response_model=QuestionResponse)
async def create_question(
    question: QuestionCreate, 
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_question = Question(**question.dict(), created_by=current_user.id)
    db.add(db_question)
    db.commit()
    db.refresh(db_question)
    return db_question

@app.get("/questions", response_model=List[QuestionResponse])
async def get_questions(
    department: Optional[str] = None,
    level: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Question)
    
    if department:
        query = query.filter(Question.department == department)
    if level:
        query = query.filter(Question.level == level)
        
    return query.all()

# Quiz routes
@app.post("/quizzes", response_model=QuizResponse)
async def create_quiz(
    quiz: QuizCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_quiz = Quiz(**quiz.dict(), created_by=current_user.id)
    db.add(db_quiz)
    db.commit()
    db.refresh(db_quiz)
    return db_quiz



# Lesson routes
@app.get("/lessons")
async def get_lessons(
    department: Optional[str] = None,
    level: Optional[str] = None,
    active_only: bool = True,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Lesson)
    
    if department:
        query = query.filter(Lesson.department == department)
    if level:
        query = query.filter(Lesson.level == level)
    if active_only:
        query = query.filter(Lesson.is_active == True)
        
    return query.all()

# Schedule routes
@app.post("/schedules", response_model=Dict)
async def create_schedule(
    schedule: ScheduleCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_schedule = Schedule(**schedule.dict(), created_by=current_user.id)
    db.add(db_schedule)
    db.commit()
    db.refresh(db_schedule)
    return {"id": db_schedule.id}

# Announcement routes
@app.post("/announcements", response_model=Dict)
async def create_announcement(
    announcement: AnnouncementCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_announcement = Announcement(**announcement.dict(), created_by=current_user.id)
    db.add(db_announcement)
    db.commit()
    db.refresh(db_announcement)
    return {"id": db_announcement.id}

# Health check route
@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# Create tables
Base.metadata.create_all(bind=engine)

# Additional Pydantic models
class LessonCreate(BaseModel):
    title: str
    code: str
    description: str = ""
    department: str
    level: str
    classification: str  # Core, Specific, General

class TeacherLessonAssign(BaseModel):
    teacher_id: int
    lesson_id: int

class NotificationCreate(BaseModel):
    user_id: int
    title: str
    message: str
    type: str = "info"

class QuizScheduleCreate(BaseModel):
    quiz_id: int
    scheduled_date: str
    start_time: str
    end_time: str
    target_departments: List[str]
    target_levels: List[str]

class SDMSStudentSync(BaseModel):
    sdms_url: str
    api_key: str = None

class ClassTeacherAssign(BaseModel):
    teacher_id: int
    department: str
    level: str

class StudentUpload(BaseModel):
    students: List[dict]  # [{"username": "...", "full_name": "...", "password": "..."}]

class StudentFileUpload(BaseModel):
    students: List[dict]
    count: int

class TeacherPasswordReset(BaseModel):
    new_password: str

class TeacherUpdate(BaseModel):
    full_name: str = None
    password: str = None
    departments: List[str] = None

class AnnouncementCreateFull(BaseModel):
    title: str
    content: str
    priority: str = "normal"
    departments: List[str]
    levels: List[str]

class UserRegister(BaseModel):
    username: str
    password: str
    full_name: str
    role: str = "student"
    department: str = None
    level: str = None
    departments: List[str] = []
    is_class_teacher: bool = False
    class_department: str = None
    class_level: str = None

class BulkQuestionCreate(BaseModel):
    questions: List[QuestionCreate]

class QuizAnswer(BaseModel):
    question_id: int
    answer: str

class QuizSubmission(BaseModel):
    quiz_id: int
    answers: List[QuizAnswer]

# Auth functions
def verify_password(plain_password, hashed_password):
    return verify_password_simple(plain_password, hashed_password)

def get_password_hash(password):
    return hash_password_simple(password)

def create_access_token_new(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security), db: Session = Depends(get_db)):
    try:
        if not credentials or not credentials.credentials:
            raise HTTPException(status_code=401, detail="Authentication token required")
        
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise HTTPException(status_code=401, detail="Invalid token format")
    except JWTError as e:
        error_msg = str(e).lower()
        if "expired" in error_msg:
            raise HTTPException(status_code=401, detail="Token has expired. Please login again.")
        else:
            raise HTTPException(status_code=401, detail="Invalid token. Please login again.")
    except Exception as e:
        raise HTTPException(status_code=401, detail="Authentication failed. Please login again.")
    
    user = db.query(User).filter(User.username == username).first()
    if user is None:
        raise HTTPException(status_code=401, detail="User not found. Please login again.")
    return user

# Routes
@app.post("/auth/login")
def login(user_data: UserLogin, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == user_data.username).first()
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    
    if not verify_password(user_data.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid username or password")
    
    access_token = create_access_token_new(data={"sub": user.username, "role": user.role})
    
    # Create user dict with explicit None handling
    user_dict = {
        "id": user.id,
        "username": user.username,
        "role": user.role,
        "full_name": user.full_name or "",
        "department": user.department,
        "level": user.level,
        "departments": user.departments or [],
        "is_class_teacher": bool(getattr(user, 'is_class_teacher', False))
    }
    
    print(f"Login successful for user: {user.username}, role: {user.role}")
    print(f"User dict: {user_dict}")
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": user_dict
    }



@app.get("/questions")
def get_questions(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Teachers can only see questions they created
    if current_user.role == "teacher":
        return db.query(Question).filter(Question.created_by == current_user.id).all()
    else:  # admin - can see all
        return db.query(Question).all()

@app.post("/questions/bulk")
def create_bulk_questions(questions: BulkQuestionCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    if not questions.questions:
        raise HTTPException(status_code=400, detail="No questions provided")
    
    created_questions = []
    errors = []
    
    for i, question in enumerate(questions.questions):
        try:
            # Validate lesson assignment for teachers
            if current_user.role == "teacher":
                assignment = db.query(TeacherLesson).filter(
                    TeacherLesson.teacher_id == current_user.id,
                    TeacherLesson.lesson_id == question.lesson_id
                ).first()
                if not assignment:
                    errors.append(f"Question {i+1}: Not authorized for this lesson")
                    continue
            
            # Validate lesson exists and matches department/level
            lesson = db.query(Lesson).filter(
                Lesson.id == question.lesson_id,
                Lesson.department == question.department,
                Lesson.level == question.level,
                Lesson.is_active == True
            ).first()
            if not lesson:
                errors.append(f"Question {i+1}: Invalid lesson for {question.department} - {question.level}")
                continue
            
            # Validate MCQ options
            if question.question_type == 'mcq':
                valid_options = [opt for opt in question.options if opt.strip()]
                if len(valid_options) < 2:
                    errors.append(f"Question {i+1}: MCQ needs at least 2 options")
                    continue
                if question.correct_answer not in valid_options:
                    errors.append(f"Question {i+1}: Correct answer must be one of the options")
                    continue
            
            db_question = Question(
                question_text=question.question_text,
                question_type=question.question_type,
                options=question.options if question.question_type != 'short_answer' else [],
                correct_answer=question.correct_answer,
                points=question.points,
                department=question.department,
                level=question.level,
                lesson_id=question.lesson_id,
                created_by=current_user.id
            )
            db.add(db_question)
            created_questions.append(db_question)
            
        except Exception as e:
            errors.append(f"Question {i+1}: {str(e)}")
    
    if created_questions:
        db.commit()
    
    result_message = f"Successfully created {len(created_questions)} questions"
    if errors:
        result_message += f". {len(errors)} questions had errors."
    
    return {
        "message": result_message,
        "count": len(created_questions),
        "errors": errors[:5] if errors else []  # Limit error messages
    }

@app.post("/questions")
def create_question(question: QuestionCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Validate lesson assignment for teachers
    if current_user.role == "teacher":
        assignment = db.query(TeacherLesson).filter(
            TeacherLesson.teacher_id == current_user.id,
            TeacherLesson.lesson_id == question.lesson_id
        ).first()
        if not assignment:
            raise HTTPException(status_code=403, detail="You can only create questions for lessons assigned to you by DOS")
    
    # Validate lesson exists and matches department/level
    lesson = db.query(Lesson).filter(
        Lesson.id == question.lesson_id,
        Lesson.department == question.department,
        Lesson.level == question.level,
        Lesson.is_active == True
    ).first()
    if not lesson:
        raise HTTPException(status_code=400, detail=f"Invalid lesson selection for {question.department} - {question.level}")
    
    db_question = Question(
        question_text=question.question_text,
        question_type=question.question_type,
        options=question.options if question.question_type != 'short_answer' else [],
        correct_answer=question.correct_answer,
        points=question.points,
        department=question.department,
        level=question.level,
        lesson_id=question.lesson_id,
        created_by=current_user.id
    )
    db.add(db_question)
    db.commit()
    db.refresh(db_question)
    return db_question

@app.get("/quizzes")
def get_quizzes(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role == "admin":
        quizzes = db.query(Quiz).all()
    elif current_user.role == "teacher":
        # Teachers can only see quizzes they created
        quizzes = db.query(Quiz).filter(Quiz.created_by == current_user.id).all()
    else:  # student - STRICT filtering by department AND level
        if not current_user.department or not current_user.level:
            print(f"❌ Student {current_user.username} missing dept/level: {current_user.department}/{current_user.level}")
            raise HTTPException(status_code=400, detail="Student must have department and level assigned")
        
        print(f"🔍 Filtering quizzes for student {current_user.username}: {current_user.department} - {current_user.level}")
        
        # Get ALL active quizzes first for debugging
        all_quizzes = db.query(Quiz).filter(Quiz.is_active == True).all()
        print(f"📊 Total active quizzes: {len(all_quizzes)}")
        for q in all_quizzes:
            print(f"  Quiz {q.id}: {q.title} | {q.department} - {q.level}")
        
        # Apply strict filtering
        quizzes = db.query(Quiz).filter(
            Quiz.is_active == True,
            Quiz.department == current_user.department,
            Quiz.level == current_user.level
        ).all()
        
        print(f"✅ Student {current_user.username} ({current_user.department} - {current_user.level}) should see {len(quizzes)} quizzes")
        for q in quizzes:
            print(f"  ✓ Quiz {q.id}: {q.title}")
    
    # Add timing information for students
    if current_user.role == "student":
        for quiz in quizzes:
            if quiz.countdown_started_at:
                elapsed_time = (datetime.utcnow() - quiz.countdown_started_at).total_seconds()
                total_time = quiz.duration_minutes * 60
                quiz.time_remaining = max(0, total_time - elapsed_time)
                quiz.is_expired = elapsed_time > total_time
            else:
                quiz.time_remaining = quiz.duration_minutes * 60
                quiz.is_expired = False
    
    return quizzes

@app.post("/quizzes")
def create_quiz(quiz: QuizCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Admin can create quizzes for any department, teachers need department validation
    if current_user.role == "teacher":
        user_departments = current_user.departments or []
        if user_departments and quiz.department not in user_departments:
            raise HTTPException(status_code=403, detail="You can only create quizzes for your assigned departments")
    
    db_quiz = Quiz(
        title=quiz.title,
        description=quiz.description,
        scheduled_time=quiz.scheduled_time,
        duration_minutes=quiz.duration_minutes,
        question_time_seconds=quiz.question_time_seconds,
        department=quiz.department,
        level=quiz.level,
        created_by=current_user.id
    )
    db.add(db_quiz)
    db.commit()
    db.refresh(db_quiz)
    
    # Add questions to quiz (filter by department and level)
    for i, question_id in enumerate(quiz.question_ids):
        question = db.query(Question).filter(
            Question.id == question_id,
            Question.department == quiz.department,
            Question.level == quiz.level
        ).first()
        if question:
            quiz_question = QuizQuestion(
                quiz_id=db_quiz.id,
                question_id=question_id,
                question_order=i + 1
            )
            db.add(quiz_question)
    
    db.commit()
    return db_quiz

@app.get("/quizzes/{quiz_id}/questions")
def get_quiz_questions(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Get quiz details first
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    # For students: STRICT validation - they can only access quizzes for their exact department and level
    if current_user.role == "student":
        print(f"🔒 Validating quiz access for {current_user.username}:")
        print(f"   Student: {current_user.department} - {current_user.level}")
        print(f"   Quiz: {quiz.department} - {quiz.level}")
        
        if quiz.department != current_user.department or quiz.level != current_user.level:
            print(f"❌ Access denied: Department/Level mismatch")
            raise HTTPException(status_code=403, detail=f"Access denied. This quiz is for {quiz.department} - {quiz.level} students only.")
        
        if not quiz.is_active:
            print(f"❌ Access denied: Quiz not active")
            raise HTTPException(status_code=403, detail="Quiz is not active")
        
        print(f"✅ Access granted for quiz {quiz_id}")
    
    # Check if quiz time has expired
    if quiz.countdown_started_at:
        total_quiz_time = quiz.duration_minutes * 60  # Convert to seconds
        elapsed_time = (datetime.utcnow() - quiz.countdown_started_at).total_seconds()
        
        if elapsed_time > total_quiz_time:
            raise HTTPException(status_code=400, detail="Quiz time is over! The quiz has ended.")
        elif elapsed_time > 0:
            # Quiz is running, show time warning
            remaining_time = total_quiz_time - elapsed_time
            if remaining_time < 300:  # Less than 5 minutes
                raise HTTPException(status_code=400, detail=f"Time is going up, hurry up! Only {int(remaining_time/60)} minutes {int(remaining_time%60)} seconds left!")
    
    # Check if user already attempted this quiz
    attempt = db.query(QuizAttempt).filter(
        QuizAttempt.quiz_id == quiz_id,
        QuizAttempt.user_id == current_user.id
    ).first()
    
    if attempt and current_user.role != "admin":
        raise HTTPException(status_code=400, detail="Quiz already attempted")
    
    questions = db.query(Question).join(QuizQuestion).filter(
        QuizQuestion.quiz_id == quiz_id
    ).order_by(QuizQuestion.question_order).all()
    
    # Remove correct answers for students
    if current_user.role not in ["admin", "teacher"]:
        for question in questions:
            question.correct_answer = None
    
    return questions

@app.post("/quizzes/submit")
def submit_quiz(submission: QuizSubmission, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Get quiz first to validate access
    quiz = db.query(Quiz).filter(Quiz.id == submission.quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    # For students: validate they can only submit quizzes for their department and level
    if current_user.role == "student":
        if quiz.department != current_user.department or quiz.level != current_user.level:
            print(f"❌ Submit denied: {current_user.username} tried to submit quiz for {quiz.department}-{quiz.level} but is {current_user.department}-{current_user.level}")
            raise HTTPException(status_code=403, detail="You can only submit quizzes for your class")
    
    # Check if already attempted
    existing_attempt = db.query(QuizAttempt).filter(
        QuizAttempt.quiz_id == submission.quiz_id,
        QuizAttempt.user_id == current_user.id
    ).first()
    
    if existing_attempt:
        raise HTTPException(status_code=400, detail="Quiz already attempted")
    
    # Get quiz questions
    questions = db.query(Question).join(QuizQuestion).filter(
        QuizQuestion.quiz_id == submission.quiz_id
    ).all()
    
    # Calculate score based on marks allocated to each question
    score = 0
    total_marks = sum(q.points for q in questions)
    answers_dict = {ans.question_id: ans.answer for ans in submission.answers}
    
    for question in questions:
        user_answer = answers_dict.get(question.id, "")
        if question.question_type == "short_answer":
            if user_answer.lower().strip() == question.correct_answer.lower().strip():
                score += question.points
        else:
            if user_answer == question.correct_answer:
                score += question.points
    
    # Save attempt
    attempt = QuizAttempt(
        quiz_id=submission.quiz_id,
        user_id=current_user.id,
        score=score,
        total_questions=total_marks,
        answers=[{"question_id": ans.question_id, "answer": ans.answer} for ans in submission.answers],
        completed_at=datetime.utcnow()
    )
    db.add(attempt)
    db.flush()
    
    # Save individual answers
    for question in questions:
        user_answer = answers_dict.get(question.id, "")
        is_correct = False
        if question.question_type == "short_answer":
            is_correct = user_answer.lower().strip() == question.correct_answer.lower().strip()
        else:
            is_correct = user_answer == question.correct_answer
        
        student_answer = StudentAnswer(
            attempt_id=attempt.id,
            question_id=question.id,
            student_answer=user_answer,
            is_correct=is_correct
        )
        db.add(student_answer)
    
    # Notify teacher
    if quiz and quiz.created_by:
        percentage = round((score / total_marks) * 100, 1) if total_marks > 0 else 0
        notification = Notification(
            user_id=quiz.created_by,
            title=f"📊 New Quiz Submission",
            message=f"{current_user.full_name or current_user.username} completed '{quiz.title}' - Score: {score}/{total_marks} marks ({percentage}%)",
            type="quiz_submission"
        )
        db.add(notification)
    
    db.commit()
    
    return {"score": score, "total_questions": total_marks}

@app.get("/leaderboard/{quiz_id}")
def get_leaderboard(quiz_id: int, db: Session = Depends(get_db)):
    results = db.query(QuizAttempt, User).join(User).filter(
        QuizAttempt.quiz_id == quiz_id
    ).order_by(QuizAttempt.score.desc()).limit(10).all()
    
    return [
        {
            "username": user.username,
            "full_name": user.full_name,
            "score": attempt.score,
            "total_questions": attempt.total_questions,
            "completed_at": attempt.completed_at
        }
        for attempt, user in results
    ]

@app.get("/results/{quiz_id}")
def get_quiz_results(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    results = db.query(QuizAttempt, User).join(User).filter(
        QuizAttempt.quiz_id == quiz_id
    ).order_by(QuizAttempt.score.desc()).all()
    
    return [
        {
            "username": user.username,
            "full_name": user.full_name,
            "score": attempt.score,
            "total_questions": attempt.total_questions,
            "answers": attempt.answers,
            "started_at": attempt.started_at,
            "completed_at": attempt.completed_at
        }
        for attempt, user in results
    ]

@app.put("/quizzes/{quiz_id}/activate")
def activate_quiz(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz.is_active = True
    db.commit()
    return {"message": "Quiz activated and broadcasted to all students"}

@app.put("/quizzes/{quiz_id}/broadcast")
def broadcast_quiz(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    quiz.is_active = True
    quiz.scheduled_time = datetime.utcnow()
    quiz.countdown_started_at = datetime.utcnow()
    
    # Notify students in target department/level
    students = db.query(User).filter(
        User.role == "student",
        User.department == quiz.department,
        User.level == quiz.level
    ).all()
    
    for student in students:
        notification = Notification(
            user_id=student.id,
            title="🎯 Quiz Started - Hurry Up!",
            message=f"Quiz '{quiz.title}' has started NOW! Time per question: {quiz.question_time_seconds}s. Total duration: {quiz.duration_minutes} minutes. Join immediately!",
            type="quiz_started"
        )
        db.add(notification)
    
    db.commit()
    return {"message": "Quiz countdown started and broadcasted to all students immediately"}

@app.get("/quizzes/{quiz_id}/export/pdf")
def export_quiz_pdf(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Get quiz and results
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    results = db.query(QuizAttempt, User).join(User).filter(
        QuizAttempt.quiz_id == quiz_id
    ).order_by(QuizAttempt.score.desc()).all()
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], alignment=1, spaceAfter=30)
    story.append(Paragraph(f"Quiz Results: {quiz.title}", title_style))
    story.append(Spacer(1, 12))
    
    # Quiz info
    info_style = styles['Normal']
    story.append(Paragraph(f"<b>Description:</b> {quiz.description}", info_style))
    story.append(Paragraph(f"<b>Duration:</b> {quiz.duration_minutes} minutes", info_style))
    story.append(Paragraph(f"<b>Total Submissions:</b> {len(results)}", info_style))
    story.append(Spacer(1, 20))
    
    # Results table
    if results:
        data = [['Rank', 'Student Name', 'Username', 'Score', 'Percentage', 'Completed At']]
        for i, (attempt, user) in enumerate(results, 1):
            percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
            data.append([
                str(i),
                user.full_name or user.username,
                user.username,
                f"{attempt.score}/{attempt.total_questions}",
                f"{percentage}%",
                attempt.completed_at.strftime('%Y-%m-%d %H:%M') if attempt.completed_at else 'N/A'
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
    else:
        story.append(Paragraph("No submissions yet.", styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=quiz_{quiz_id}_results.pdf"}
    )

@app.get("/quizzes/{quiz_id}/export/excel")
def export_quiz_excel(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Get quiz and results
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    results = db.query(QuizAttempt, User).join(User).filter(
        QuizAttempt.quiz_id == quiz_id
    ).order_by(QuizAttempt.score.desc()).all()
    
    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Quiz Results"
    
    # Headers
    headers = ['Rank', 'Student Name', 'Username', 'Score', 'Total Questions', 'Percentage', 'Completed At']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    # Data
    for i, (attempt, user) in enumerate(results, 1):
        percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
        row_data = [
            i,
            user.full_name or user.username,
            user.username,
            attempt.score,
            attempt.total_questions,
            percentage,
            attempt.completed_at.strftime('%Y-%m-%d %H:%M') if attempt.completed_at else 'N/A'
        ]
        
        for col, value in enumerate(row_data, 1):
            cell = ws.cell(row=i+1, column=col, value=value)
            cell.alignment = Alignment(horizontal="center")
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=quiz_{quiz_id}_results.xlsx"}
    )

# DOS (Admin) specific endpoints
@app.post("/admin/register-teacher")
def register_teacher(user_data: UserRegister, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Check if username already exists
    existing_user = db.query(User).filter(User.username == user_data.username).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Username already exists")
    
    # Validate teacher data
    if not user_data.departments or len(user_data.departments) == 0:
        raise HTTPException(status_code=400, detail="Teacher must be assigned to at least one department")
    
    # Validate class teacher assignment
    if user_data.is_class_teacher:
        if not user_data.class_department or not user_data.class_level:
            raise HTTPException(status_code=400, detail="Class department and level required for class teachers")
    
    # Create teacher account
    hashed_password = hash_password_simple(user_data.password)
    teacher = User(
        username=user_data.username,
        password_hash=hashed_password,
        role="teacher",
        full_name=user_data.full_name,
        departments=user_data.departments,
        is_class_teacher=user_data.is_class_teacher
    )
    db.add(teacher)
    db.commit()
    db.refresh(teacher)
    
    # Auto-assign as class teacher if specified
    if user_data.is_class_teacher and user_data.class_department and user_data.class_level:
        class_teacher = ClassTeacher(
            teacher_id=teacher.id,
            department=user_data.class_department,
            level=user_data.class_level,
            assigned_by=current_user.id
        )
        db.add(class_teacher)
        
        # Create notification for new class teacher
        notification = Notification(
            user_id=teacher.id,
            title="🏫 Class Teacher Assignment",
            message=f"You have been assigned as class teacher for {user_data.class_department} - {user_data.class_level}. You can now upload student lists for this class.",
            type="class_teacher_assignment"
        )
        db.add(notification)
        
        db.commit()
    
    return {
        "message": "Teacher registered successfully",
        "teacher": {
            "id": teacher.id,
            "username": teacher.username,
            "full_name": teacher.full_name,
            "departments": teacher.departments,
            "is_class_teacher": teacher.is_class_teacher,
            "class_assignment": f"{user_data.class_department} - {user_data.class_level}" if user_data.is_class_teacher else None,
            "created_at": teacher.created_at
        }
    }

@app.get("/schedules")
def get_schedules(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role == "admin":
        return db.query(Schedule).order_by(Schedule.scheduled_date.desc()).all()
    elif current_user.role == "teacher":
        user_departments = current_user.departments or []
        schedules = db.query(Schedule).all()
        # Filter schedules that match teacher's departments
        filtered = []
        for schedule in schedules:
            schedule_depts = schedule.departments or []
            if any(dept in user_departments for dept in schedule_depts):
                filtered.append(schedule)
        return filtered
    else:
        # Students see schedules for their department/level
        schedules = db.query(Schedule).all()
        filtered = []
        for schedule in schedules:
            schedule_depts = schedule.departments or []
            schedule_levels = schedule.levels or []
            if (current_user.department in schedule_depts and 
                current_user.level in schedule_levels):
                filtered.append(schedule)
        return filtered

@app.post("/schedules/upload")
async def upload_schedule_file(
    file: UploadFile = File(...),
    title: str = "",
    description: str = "",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    if not file.filename.endswith(('.pdf', '.doc', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and Word files allowed")
    
    try:
        import base64
        
        # Read file and encode as base64
        content = await file.read()
        file_data = base64.b64encode(content).decode('utf-8')
        
        # Determine file type
        if file.filename.endswith('.pdf'):
            file_type = 'application/pdf'
        elif file.filename.endswith('.docx'):
            file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            file_type = 'application/msword'
        
        # Create schedule
        db_schedule = Schedule(
            title=title or f"Weekly Timetable - {datetime.utcnow().strftime('%Y-%m-%d')}",
            description=description or "Morning Quiz Weekly Schedule",
            scheduled_date=datetime.utcnow(),
            departments=[],
            levels=[],
            file_data=file_data,
            file_name=file.filename,
            file_type=file_type,
            created_by=current_user.id
        )
        db.add(db_schedule)
        
        # Notify all teachers
        teachers = db.query(User).filter(User.role == "teacher").all()
        for teacher in teachers:
            notification = Notification(
                user_id=teacher.id,
                title="📅 New Weekly Timetable",
                message=f"DOS has uploaded a new weekly timetable: {db_schedule.title}",
                type="schedule_upload"
            )
            db.add(notification)
        
        db.commit()
        db.refresh(db_schedule)
        return db_schedule
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/schedules")
def create_schedule(schedule: ScheduleCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    db_schedule = Schedule(
        title=schedule.title,
        description=schedule.description,
        scheduled_date=schedule.scheduled_date,
        departments=schedule.departments,
        levels=schedule.levels,
        created_by=current_user.id
    )
    db.add(db_schedule)
    db.commit()
    db.refresh(db_schedule)
    return db_schedule

@app.get("/announcements")
def get_announcements(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role == "admin":
        return db.query(Announcement).filter(Announcement.is_active == True).order_by(Announcement.created_at.desc()).all()
    elif current_user.role == "teacher":
        user_departments = current_user.departments or []
        announcements = db.query(Announcement).filter(Announcement.is_active == True).all()
        # Filter announcements that match teacher's departments
        filtered = []
        for announcement in announcements:
            announcement_depts = announcement.departments or []
            if any(dept in user_departments for dept in announcement_depts):
                filtered.append(announcement)
        return filtered
    else:
        # Students see announcements for their department/level
        announcements = db.query(Announcement).filter(Announcement.is_active == True).all()
        filtered = []
        for announcement in announcements:
            announcement_depts = announcement.departments or []
            announcement_levels = announcement.levels or []
            if (current_user.department in announcement_depts and 
                current_user.level in announcement_levels):
                filtered.append(announcement)
        return filtered

@app.post("/announcements")
def create_announcement(announcement: AnnouncementCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    db_announcement = Announcement(
        title=announcement.title,
        content=announcement.content,
        priority=announcement.priority,
        departments=announcement.departments,
        levels=announcement.levels,
        created_by=current_user.id
    )
    db.add(db_announcement)
    db.commit()
    db.refresh(db_announcement)
    return db_announcement

@app.put("/announcements/{announcement_id}/deactivate")
def deactivate_announcement(announcement_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    announcement = db.query(Announcement).filter(Announcement.id == announcement_id).first()
    if not announcement:
        raise HTTPException(status_code=404, detail="Announcement not found")
    
    announcement.is_active = False
    db.commit()
    return {"message": "Announcement deactivated"}

# Lesson management endpoints
@app.get("/lessons")
def get_lessons(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role == "admin":
        return db.query(Lesson).filter(Lesson.is_active == True).order_by(Lesson.department, Lesson.level, Lesson.title).all()
    elif current_user.role == "teacher":
        # Teachers see only lessons assigned to them
        assigned_lessons = db.query(Lesson).join(TeacherLesson).filter(
            TeacherLesson.teacher_id == current_user.id,
            Lesson.is_active == True
        ).order_by(Lesson.department, Lesson.level, Lesson.title).all()
        return assigned_lessons
    else:
        # Students see lessons for their department/level
        return db.query(Lesson).filter(
            Lesson.department == current_user.department,
            Lesson.level == current_user.level,
            Lesson.is_active == True
        ).order_by(Lesson.title).all()

@app.get("/my-courses")
def get_my_courses(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Teacher access required")
    
    assignments = db.query(TeacherLesson, Lesson).join(Lesson).filter(
        TeacherLesson.teacher_id == current_user.id,
        Lesson.is_active == True
    ).order_by(Lesson.department, Lesson.level, Lesson.title).all()
    
    return [{
        "assignment_id": assignment.id,
        "lesson": {
            "id": lesson.id,
            "title": lesson.title,
            "code": lesson.code,
            "description": lesson.description,
            "department": lesson.department,
            "level": lesson.level,
            "classification": lesson.classification
        },
        "assigned_at": assignment.assigned_at
    } for assignment, lesson in assignments]

@app.post("/lessons")
def create_lesson(lesson: LessonCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Check if lesson code already exists
    existing_lesson = db.query(Lesson).filter(Lesson.code == lesson.code).first()
    if existing_lesson:
        raise HTTPException(status_code=400, detail="Lesson code already exists")
    
    db_lesson = Lesson(
        title=lesson.title,
        code=lesson.code,
        description=lesson.description,
        department=lesson.department,
        level=lesson.level,
        classification=lesson.classification,
        created_by=current_user.id
    )
    db.add(db_lesson)
    db.commit()
    db.refresh(db_lesson)
    return db_lesson

@app.put("/lessons/{lesson_id}/deactivate")
def deactivate_lesson(lesson_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    lesson = db.query(Lesson).filter(Lesson.id == lesson_id).first()
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    lesson.is_active = False
    db.commit()
    return {"message": "Lesson deactivated"}

# Teacher assignment endpoints
@app.get("/teachers")
def get_teachers(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    teachers = db.query(User).filter(User.role == "teacher").all()
    return [{
        "id": teacher.id,
        "username": teacher.username,
        "full_name": teacher.full_name,
        "departments": teacher.departments
    } for teacher in teachers]

@app.post("/teacher-lessons")
def assign_lesson_to_teacher(assignment: TeacherLessonAssign, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Check if teacher exists
    teacher = db.query(User).filter(User.id == assignment.teacher_id, User.role == "teacher").first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    
    # Check if lesson exists
    lesson = db.query(Lesson).filter(Lesson.id == assignment.lesson_id, Lesson.is_active == True).first()
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    # Check if assignment already exists
    existing = db.query(TeacherLesson).filter(
        TeacherLesson.teacher_id == assignment.teacher_id,
        TeacherLesson.lesson_id == assignment.lesson_id
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="Teacher already assigned to this lesson")
    
    # Create assignment
    teacher_lesson = TeacherLesson(
        teacher_id=assignment.teacher_id,
        lesson_id=assignment.lesson_id,
        assigned_by=current_user.id
    )
    db.add(teacher_lesson)
    
    # Create notification for teacher
    notification = Notification(
        user_id=assignment.teacher_id,
        title="📚 New Lesson Assignment",
        message=f"You have been assigned to teach '{lesson.title}' ({lesson.code}) for {lesson.department} - {lesson.level}. Classification: {lesson.classification}",
        type="lesson_assignment"
    )
    db.add(notification)
    
    db.commit()
    return {"message": "Lesson assigned to teacher successfully"}

@app.get("/teacher-lessons/{teacher_id}")
def get_teacher_lessons(teacher_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    assignments = db.query(TeacherLesson, Lesson).join(Lesson).filter(
        TeacherLesson.teacher_id == teacher_id,
        Lesson.is_active == True
    ).all()
    
    return [{
        "id": assignment.id,
        "lesson": {
            "id": lesson.id,
            "title": lesson.title,
            "code": lesson.code,
            "department": lesson.department,
            "level": lesson.level,
            "classification": lesson.classification
        },
        "assigned_at": assignment.assigned_at
    } for assignment, lesson in assignments]

@app.delete("/teacher-lessons/{assignment_id}")
def remove_teacher_lesson(assignment_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    assignment = db.query(TeacherLesson).filter(TeacherLesson.id == assignment_id).first()
    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")
    
    # Get teacher and lesson info for notification
    teacher = db.query(User).filter(User.id == assignment.teacher_id).first()
    lesson = db.query(Lesson).filter(Lesson.id == assignment.lesson_id).first()
    
    # Create notification for teacher
    if teacher and lesson:
        notification = Notification(
            user_id=assignment.teacher_id,
            title="📚 Lesson Assignment Removed",
            message=f"Your assignment to teach '{lesson.title}' ({lesson.code}) has been removed by DOS.",
            type="lesson_removed"
        )
        db.add(notification)
    
    db.delete(assignment)
    db.commit()
    return {"message": "Teacher lesson assignment removed"}

# Notification endpoints
@app.get("/notifications")
def get_notifications(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    notifications = db.query(Notification).filter(
        Notification.user_id == current_user.id
    ).order_by(Notification.created_at.desc()).all()
    return notifications

@app.put("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    notification = db.query(Notification).filter(
        Notification.id == notification_id,
        Notification.user_id == current_user.id
    ).first()
    if not notification:
        raise HTTPException(status_code=404, detail="Notification not found")
    
    notification.is_read = True
    db.commit()
    return {"message": "Notification marked as read"}

@app.post("/quiz-results/{quiz_id}/forward")
def forward_quiz_results(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Teacher access required")
    
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id, Quiz.created_by == current_user.id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found or not created by you")
    
    # Get DOS/admin users
    admins = db.query(User).filter(User.role == "admin").all()
    
    for admin in admins:
        notification = Notification(
            user_id=admin.id,
            title="Quiz Results Forwarded",
            message=f"Teacher {current_user.full_name} has forwarded results for quiz '{quiz.title}' ({quiz.department} {quiz.level}).",
            type="quiz_results"
        )
        db.add(notification)
    
    db.commit()
    return {"message": "Quiz results forwarded to DOS successfully"}

@app.get("/student-report/{quiz_id}")
def get_student_report(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "student":
        raise HTTPException(status_code=403, detail="Student access required")
    
    # Get quiz and student attempt
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    attempt = db.query(QuizAttempt).filter(
        QuizAttempt.quiz_id == quiz_id,
        QuizAttempt.user_id == current_user.id
    ).first()
    if not attempt:
        raise HTTPException(status_code=404, detail="No attempt found")
    
    # Get student answers
    answers = db.query(StudentAnswer).filter(StudentAnswer.attempt_id == attempt.id).all()
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], alignment=1, spaceAfter=20, textColor=colors.HexColor('#1e40af'))
    story.append(Paragraph(f"📊 Quiz Performance Report", title_style))
    story.append(Paragraph(f"{quiz.title}", styles['Heading2']))
    story.append(Spacer(1, 20))
    
    # Student info
    story.append(Paragraph(f"<b>Student:</b> {current_user.full_name or current_user.username}", styles['Normal']))
    story.append(Paragraph(f"<b>Department:</b> {current_user.department or 'N/A'} | <b>Level:</b> {current_user.level or 'N/A'}", styles['Normal']))
    story.append(Paragraph(f"<b>Date:</b> {attempt.completed_at.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Results summary
    percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
    grade = "A+" if percentage >= 90 else "A" if percentage >= 80 else "B" if percentage >= 70 else "C" if percentage >= 60 else "D" if percentage >= 50 else "F"
    
    summary_data = [
        ['Score', 'Grade', 'Correct', 'Wrong'],
        [f'{percentage}%', grade, str(attempt.score), str(attempt.total_questions - attempt.score)]
    ]
    summary_table = Table(summary_data, colWidths=[120, 120, 120, 120])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#eff6ff')),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Detailed answers
    story.append(Paragraph("<b>📝 Detailed Question Analysis</b>", styles['Heading2']))
    story.append(Spacer(1, 10))
    
    for idx, answer in enumerate(answers, 1):
        question = db.query(Question).filter(Question.id == answer.question_id).first()
        if not question:
            continue
        
        is_correct = answer.is_correct
        status = "✅ CORRECT" if is_correct else "❌ WRONG"
        status_color = colors.HexColor('#10b981') if is_correct else colors.HexColor('#ef4444')
        
        q_style = ParagraphStyle('Question', parent=styles['Normal'], fontSize=11, textColor=status_color, fontName='Helvetica-Bold')
        story.append(Paragraph(f"Question {idx}: {status}", q_style))
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"<b>Q:</b> {question.question_text}", styles['Normal']))
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"<b>Your Answer:</b> {answer.student_answer}", styles['Normal']))
        correct_style = ParagraphStyle('Correct', parent=styles['Normal'], textColor=colors.HexColor('#10b981'))
        story.append(Paragraph(f"<b>Correct Answer:</b> {question.correct_answer}", correct_style))
        story.append(Spacer(1, 10))
    
    # Study tips
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>💡 Study Tips & Recommendations</b>", styles['Heading2']))
    story.append(Spacer(1, 10))
    
    if percentage >= 90:
        tip = "🌟 Outstanding! Keep up the excellent work and help others learn."
    elif percentage >= 80:
        tip = "✨ Great job! Review the questions you missed to achieve perfection."
    elif percentage >= 70:
        tip = "👍 Good work! Focus on understanding the concepts you found challenging."
    elif percentage >= 60:
        tip = "📚 You're on the right track! Spend more time reviewing the material."
    else:
        tip = "💪 Don't give up! Review all questions carefully and practice more."
    
    story.append(Paragraph(tip, styles['Normal']))
    story.append(Spacer(1, 10))
    story.append(Paragraph("• Review all wrong answers and understand why they were incorrect", styles['Normal']))
    story.append(Paragraph("• Study the correct answers and related concepts", styles['Normal']))
    story.append(Paragraph("• Practice similar questions to reinforce your learning", styles['Normal']))
    story.append(Paragraph("• Ask your teacher for clarification on difficult topics", styles['Normal']))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<i>Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC</i>", styles['Italic']))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=student_report_{quiz_id}.pdf"}
    )

@app.get("/student-quiz-details/{quiz_id}")
def get_student_quiz_details(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "student":
        raise HTTPException(status_code=403, detail="Student access required")
    
    attempt = db.query(QuizAttempt).filter(
        QuizAttempt.quiz_id == quiz_id,
        QuizAttempt.user_id == current_user.id
    ).first()
    
    if not attempt:
        raise HTTPException(status_code=404, detail="No attempt found")
    
    answers = db.query(StudentAnswer).filter(StudentAnswer.attempt_id == attempt.id).all()
    
    details = []
    for answer in answers:
        question = db.query(Question).filter(Question.id == answer.question_id).first()
        if question:
            details.append({
                "question_text": question.question_text,
                "student_answer": answer.student_answer,
                "correct_answer": question.correct_answer,
                "is_correct": answer.is_correct,
                "options": question.options if question.question_type == 'mcq' else []
            })
    
    return {"details": details, "score": attempt.score, "total": attempt.total_questions}

@app.get("/student/progress")
def get_student_progress(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["student", "admin"]:
        raise HTTPException(status_code=403, detail="Student or Admin access required")
    
    # For admin, get all students' progress; for student, get own progress
    if current_user.role == "admin":
        attempts = db.query(QuizAttempt, Quiz, User).join(
            Quiz, QuizAttempt.quiz_id == Quiz.id
        ).join(
            User, QuizAttempt.user_id == User.id
        ).order_by(QuizAttempt.completed_at.desc()).limit(50).all()
    else:
        attempts = db.query(QuizAttempt, Quiz).join(
            Quiz, QuizAttempt.quiz_id == Quiz.id
        ).filter(
            QuizAttempt.user_id == current_user.id
        ).order_by(QuizAttempt.completed_at.desc()).all()
    
    progress_data = []
    total_score = 0
    total_possible = 0
    
    for attempt, quiz, *user_data in attempts:
        user = user_data[0] if user_data else current_user
        percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
        
        # Grade calculation
        if percentage >= 90: grade = "A+"
        elif percentage >= 80: grade = "A"
        elif percentage >= 70: grade = "B"
        elif percentage >= 60: grade = "C"
        elif percentage >= 50: grade = "D"
        else: grade = "F"
        
        progress_data.append({
            "quiz_id": quiz.id,
            "quiz_title": quiz.title,
            "department": quiz.department,
            "level": quiz.level,
            "score": attempt.score,
            "total_questions": attempt.total_questions,
            "percentage": percentage,
            "grade": grade,
            "completed_at": attempt.completed_at,
            "student_name": user.full_name if current_user.role == "admin" else None
        })
        
        total_score += attempt.score
        total_possible += attempt.total_questions
    
    # Calculate overall performance
    overall_percentage = round((total_score / total_possible) * 100, 1) if total_possible > 0 else 0
    
    # Generate improvement tips
    tips = []
    if overall_percentage >= 90:
        tips = ["🌟 Excellent performance! Keep up the great work!", "💡 Consider helping classmates who need support", "📚 Challenge yourself with advanced topics"]
    elif overall_percentage >= 80:
        tips = ["✨ Great job! You're doing very well", "🎯 Focus on areas where you scored below 90%", "📖 Review incorrect answers to improve further"]
    elif overall_percentage >= 70:
        tips = ["👍 Good progress! Keep working hard", "📝 Spend more time reviewing before quizzes", "🤝 Consider forming study groups with classmates"]
    elif overall_percentage >= 60:
        tips = ["📚 You need more practice - don't give up!", "⏰ Allocate more study time daily", "🆘 Ask teachers for help on difficult topics"]
    else:
        tips = ["💪 Focus on fundamentals - you can improve!", "📖 Review all course materials thoroughly", "👨🏫 Schedule one-on-one sessions with teachers"]
    
    return {
        "recent_quizzes": progress_data,
        "overall_percentage": overall_percentage,
        "total_quizzes": len(progress_data),
        "improvement_tips": tips
    }

# Download endpoints for results
@app.get("/admin/results/download/excel")
def download_all_results_excel(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Admin or Teacher access required")
    
    try:
        # Get all quiz attempts with user and quiz info
        attempts = db.query(QuizAttempt, User, Quiz).join(User, QuizAttempt.user_id == User.id).join(Quiz, QuizAttempt.quiz_id == Quiz.id).all()
        
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Quiz Results"
        
        # Headers
        headers = ["Quiz Title", "Student Name", "Username", "Department", "Level", 
                  "Score", "Total Questions", "Percentage", "Completed At"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Data
        if attempts:
            for row, (attempt, user, quiz) in enumerate(attempts, 2):
                percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
                ws.cell(row=row, column=1, value=quiz.title)
                ws.cell(row=row, column=2, value=user.full_name or user.username)
                ws.cell(row=row, column=3, value=user.username)
                ws.cell(row=row, column=4, value=user.department or "N/A")
                ws.cell(row=row, column=5, value=user.level or "N/A")
                ws.cell(row=row, column=6, value=attempt.score)
                ws.cell(row=row, column=7, value=attempt.total_questions)
                ws.cell(row=row, column=8, value=f"{percentage}%")
                ws.cell(row=row, column=9, value=attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else 'N/A')
        else:
            ws.cell(row=2, column=1, value="No quiz attempts found")
        
        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=all_quiz_results.xlsx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@app.get("/admin/results/download/pdf")
def download_all_results_pdf(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Admin or Teacher access required")
    
    try:
        # Get all quiz attempts with user and quiz info
        attempts = db.query(QuizAttempt, User, Quiz).join(User, QuizAttempt.user_id == User.id).join(Quiz, QuizAttempt.quiz_id == Quiz.id).all()
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], alignment=1, spaceAfter=30)
        story.append(Paragraph("All Quiz Results Report", title_style))
        story.append(Spacer(1, 12))
        
        # Summary
        total_attempts = len(attempts)
        avg_score = sum(a[0].score for a in attempts) / len(attempts) if attempts else 0
        story.append(Paragraph(f"<b>Total Attempts:</b> {total_attempts}", styles['Normal']))
        story.append(Paragraph(f"<b>Average Score:</b> {avg_score:.1f}", styles['Normal']))
        story.append(Paragraph(f"<b>Generated:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC", styles['Normal']))
        story.append(Spacer(1, 20))
        
        if attempts:
            # Table data
            data = [['Quiz', 'Student', 'Score', 'Percentage', 'Date']]
            for attempt, user, quiz in attempts:
                percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
                data.append([
                    quiz.title[:20] + '...' if len(quiz.title) > 20 else quiz.title,
                    user.full_name or user.username,
                    f"{attempt.score}/{attempt.total_questions}",
                    f"{percentage}%",
                    attempt.completed_at.strftime('%m/%d/%Y') if attempt.completed_at else 'N/A'
                ])
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No quiz attempts found.", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=all_quiz_results.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# SDMS Integration
@app.post("/admin/sync-students")
def sync_students_from_sdms(sync_data: SDMSStudentSync, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Special case: Clear all students
    if sync_data.sdms_url == "CLEAR_ALL_STUDENTS" and sync_data.api_key == "CLEAR_MODE":
        try:
            students = db.query(User).filter(User.role == "student").all()
            student_count = len(students)
            
            # Get admin user to reassign questions
            admin = db.query(User).filter(User.role == "admin").first()
            if not admin:
                raise HTTPException(status_code=500, detail="No admin user found")
            
            # Update questions created by students to be owned by admin
            for student in students:
                db.query(Question).filter(Question.created_by == student.id).update({"created_by": admin.id})
                db.query(Quiz).filter(Quiz.created_by == student.id).update({"created_by": admin.id})
                
                # Delete student_answers first (child table)
                attempt_ids = [a.id for a in db.query(QuizAttempt).filter(QuizAttempt.user_id == student.id).all()]
                if attempt_ids:
                    db.query(StudentAnswer).filter(StudentAnswer.attempt_id.in_(attempt_ids)).delete(synchronize_session=False)
                
                # Then delete quiz_attempts (parent table)
                db.query(QuizAttempt).filter(QuizAttempt.user_id == student.id).delete()
                db.query(Notification).filter(Notification.user_id == student.id).delete()
            
            # Now delete students
            db.query(User).filter(User.role == "student").delete()
            
            db.commit()
            return {"message": f"Successfully cleared {student_count} students from the system", "count": student_count}
        except Exception as e:
            db.rollback()
            raise HTTPException(status_code=500, detail=f"Failed to clear students: {str(e)}")
    
    # Normal SDMS sync functionality
    try:
        headers = {}
        if sync_data.api_key:
            headers["Authorization"] = f"Bearer {sync_data.api_key}"
        
        response = requests.get(f"{sync_data.sdms_url}/api/students", headers=headers, timeout=30)
        if response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to connect to SDMS")
        
        students_data = response.json()
        created_count = 0
        updated_count = 0
        
        for student_data in students_data:
            existing_student = db.query(User).filter(User.username == student_data.get("student_id")).first()
            
            if existing_student:
                existing_student.full_name = student_data.get("full_name")
                existing_student.department = student_data.get("department")
                existing_student.level = student_data.get("level")
                updated_count += 1
            else:
                new_student = User(
                    username=student_data.get("student_id"),
                    password_hash=get_password_hash("student123"),
                    role="student",
                    full_name=student_data.get("full_name"),
                    department=student_data.get("department"),
                    level=student_data.get("level")
                )
                db.add(new_student)
                created_count += 1
        
        db.commit()
        return {"message": f"Synced students: {created_count} created, {updated_count} updated"}
        
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"SDMS connection error: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Sync failed: {str(e)}")

@app.get("/admin/students")
def get_students(department: str = None, level: str = None, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    query = db.query(User).filter(User.role == "student")
    
    if department:
        query = query.filter(User.department == department)
    if level:
        query = query.filter(User.level == level)
    
    students = query.all()
    return [{
        "id": student.id,
        "username": student.username,
        "full_name": student.full_name,
        "department": student.department,
        "level": student.level,
        "created_at": student.created_at
    } for student in students]

@app.post("/upload-questions")
async def upload_questions(file: UploadFile = File(...), current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    if not UPLOAD_ENABLED:
        raise HTTPException(status_code=501, detail="Document upload not available - missing dependencies")
    
    try:
        content = await file.read()
        questions_text = ""
        
        # Extract text based on file type
        if file.filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                questions_text += page.extract_text() + "\n"
        elif file.filename.endswith(('.doc', '.docx')):
            doc = docx.Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                questions_text += paragraph.text + "\n"
        elif file.filename.endswith('.txt'):
            questions_text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please use .doc, .docx, .pdf, or .txt")
        
        if not questions_text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the document")
        
        # Debug: Log extracted text
        print(f"📄 Extracted text preview: {questions_text[:200]}...")
        
        # Use intelligent parsing
        parsed_questions = parse_intelligent_questions(questions_text)
        
        # Debug: Log parsed results
        print(f"🤖 AI Parser found {len(parsed_questions)} questions")
        for i, q in enumerate(parsed_questions[:3]):
            print(f"  Q{i+1}: Type={q['question_type']}, Options={len(q['options'])}, Answer='{q['correct_answer']}'")
            if q['options']:
                print(f"       Options: {q['options'][:2]}...")
        
        if not parsed_questions:
            raise HTTPException(status_code=400, detail="No valid questions found. Please ensure questions are properly formatted.")
        
        return {"questions": parsed_questions, "count": len(parsed_questions)}
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")

def parse_intelligent_questions(content):
    """Parse questions with robust option and answer extraction"""
    questions = []
    
    # Split content into potential question blocks
    question_blocks = re.split(r'\n\s*(?=\d+[.)\s])', content)
    
    for block in question_blocks:
        if not block.strip():
            continue
            
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        if not lines:
            continue
            
        # Extract question text (first line, remove numbering)
        question_line = lines[0]
        question_text = re.sub(r'^\d+[.)\s]*', '', question_line).strip()
        
        if not question_text or len(question_text) < 5:
            continue
            
        # Extract options and answer from remaining lines
        options = []
        correct_answer = ''
        option_map = {}  # Map letters to option text
        
        for line in lines[1:]:
            # Try to match option patterns
            option_patterns = [
                r'^\s*([a-dA-D])[.)\s]+(.+?)\s*$',
                r'^\s*\(?([a-dA-D])\)?[:\s]+(.+?)\s*$',
                r'^\s*([a-dA-D])\s*[-–—]\s*(.+?)\s*$'
            ]
            
            option_found = False
            for pattern in option_patterns:
                match = re.match(pattern, line)
                if match:
                    letter = match.group(1).upper()
                    text = match.group(2).strip().rstrip('.,;:')
                    if text and len(text) > 1:
                        options.append(text)
                        option_map[letter] = text
                        option_found = True
                    break
            
            # If not an option, check for answer
            if not option_found:
                answer_patterns = [
                    r'(?:answer|correct|solution|ans)\s*[:.]?\s*([a-dA-D])\b',
                    r'(?:answer|correct|solution|ans)\s*[:.]?\s*(.+?)\s*$',
                    r'(?:correct\s+answer|the\s+answer)\s*(?:is)?\s*[:.]?\s*(.+?)\s*$',
                    r'^\s*([a-dA-D])\s*(?:is\s*)?(?:correct|right|answer)\b'
                ]
                
                for pattern in answer_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        answer_candidate = match.group(1).strip()
                        if answer_candidate:
                            # If it's a single letter, try to map to option text
                            if len(answer_candidate) == 1 and answer_candidate.upper() in option_map:
                                correct_answer = option_map[answer_candidate.upper()]
                            else:
                                correct_answer = answer_candidate
                        break
        
        # Determine question type and finalize
        question_type, final_options, final_answer = determine_question_type(question_text, options, correct_answer)
        
        if question_type and final_answer:  # Only add if we have a valid type and answer
            questions.append({
                'question_text': question_text,
                'question_type': question_type,
                'options': final_options,
                'correct_answer': final_answer
            })
    
    return questions

def determine_question_type(question_text, options, correct_answer):
    """Determine question type and ensure proper option/answer matching"""
    text_lower = question_text.lower()
    
    # True/False detection
    if (any(keyword in text_lower for keyword in ['true or false', 'true/false', 't/f', 'correct or incorrect']) or
        (correct_answer and any(word in correct_answer.lower() for word in ['true', 'false']))):
        
        tf_options = ['True', 'False']
        
        # Determine correct answer
        if correct_answer:
            answer_lower = correct_answer.lower()
            if any(word in answer_lower for word in ['true', 't', 'correct', 'yes']):
                final_answer = 'True'
            else:
                final_answer = 'False'
        else:
            final_answer = 'True'  # Default
            
        return 'true_false', tf_options, final_answer
    
    # Multiple choice detection (must have at least 2 valid options)
    elif len(options) >= 2:
        # Clean options
        clean_options = []
        for opt in options:
            cleaned = opt.strip().rstrip('.,;:')
            if cleaned and len(cleaned) > 1:
                clean_options.append(cleaned)
        
        if len(clean_options) >= 2:
            # Find matching answer
            final_answer = ''
            
            if correct_answer:
                answer_clean = correct_answer.strip()
                
                # Direct match with options
                for option in clean_options:
                    if answer_clean.lower() == option.lower():
                        final_answer = option
                        break
                
                # If no direct match, use the provided answer if it's reasonable
                if not final_answer and len(answer_clean) > 0:
                    # Check if it's a partial match
                    for option in clean_options:
                        if answer_clean.lower() in option.lower() or option.lower() in answer_clean.lower():
                            final_answer = option
                            break
                    
                    # If still no match, use first option as fallback
                    if not final_answer:
                        final_answer = clean_options[0]
            else:
                final_answer = clean_options[0]  # Default to first option
            
            return 'mcq', clean_options, final_answer
    
    # Short answer (fallback)
    final_answer = correct_answer.strip() if correct_answer else 'Sample answer'
    return 'short_answer', [], final_answer

# Initialize database
@app.on_event("startup")
def startup_event():
    # Create tables
    Base.metadata.create_all(bind=engine)
    
    # Add missing columns if they don't exist
    from sqlalchemy import text
    db = SessionLocal()
    try:
        # Check if is_class_teacher column exists
        db.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_class_teacher BOOLEAN DEFAULT FALSE"))
        db.commit()
    except Exception as e:
        print(f"Migration error: {e}")
        db.rollback()
    finally:
        db.close()
    
    # Create sample data if not exists
    db = SessionLocal()
    try:
        # Create admin user
        admin = db.query(User).filter(User.username == "admin").first()
        if not admin:
            admin = User(
                username="admin",
                password_hash=hash_password_simple("admin123"),
                role="admin",
                full_name="DOS Administrator",
                department=None,
                level=None,
                departments=["Software Development", "Computer System and Architecture", "Land Surveying", "Building Construction"],
                is_class_teacher=False
            )
            db.add(admin)
        else:
            # Always reset admin password and ensure all fields are set
            admin.password_hash = hash_password_simple("admin123")
            admin.role = "admin"
            admin.full_name = "DOS Administrator"
            admin.departments = ["Software Development", "Computer System and Architecture", "Land Surveying", "Building Construction"]
            admin.is_class_teacher = False
        
        # Reset ALL teacher passwords to ensure they can login with bcrypt
        teachers = db.query(User).filter(User.role == "teacher").all()
        for teacher in teachers:
            # Always reset password to bcrypt format
            teacher.password_hash = hash_password_simple("pass123")
            print(f"✅ Reset password for teacher: {teacher.username} -> pass123 (bcrypt)")
        
        # Create sample teachers if none exist
        if not teachers:
            sample_teachers = [
                {
                    "username": "teacher001",
                    "full_name": "John Smith",
                    "departments": ["Software Development", "Computer System and Architecture"]
                },
                {
                    "username": "teacher002", 
                    "full_name": "Mary Johnson",
                    "departments": ["Land Surveying", "Building Construction"]
                },
                {
                    "username": "teacher003",
                    "full_name": "David Wilson", 
                    "departments": ["Software Development"]
                }
            ]
            
            for teacher_data in sample_teachers:
                teacher = User(
                    username=teacher_data["username"],
                    password_hash=hash_password_simple("pass123"),
                    role="teacher",
                    full_name=teacher_data["full_name"],
                    departments=teacher_data["departments"],
                    is_class_teacher=False
                )
                db.add(teacher)
                print(f"✅ Created teacher: {teacher_data['username']} with password: pass123")
        
        # Create sample student
        student = db.query(User).filter(User.username == "student001").first()
        if not student:
            student = User(
                username="student001",
                password_hash=hash_password_simple("student123"),
                role="student",
                full_name="Alice Johnson",
                department="Software Development",
                level="Level 3",
                is_class_teacher=False
            )
            db.add(student)
        
        # Create sample lesson
        lesson = db.query(Lesson).filter(Lesson.code == "SD301").first()
        if not lesson:
            lesson = Lesson(
                title="Introduction to Programming",
                code="SD301",
                description="Basic programming concepts",
                department="Software Development",
                level="Level 3",
                classification="Core",
                created_by=1
            )
            db.add(lesson)
        
        db.commit()
        print("🚀 Database initialization complete!")
    except Exception as e:
        print(f"❌ Error creating sample data: {e}")
        db.rollback()
    finally:
        db.close()

@app.post("/admin/assign-class-teacher")
def assign_class_teacher(assignment: ClassTeacherAssign, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Check if teacher exists
    teacher = db.query(User).filter(User.id == assignment.teacher_id, User.role == "teacher").first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    
    # Check if assignment already exists
    existing = db.query(ClassTeacher).filter(
        ClassTeacher.teacher_id == assignment.teacher_id,
        ClassTeacher.department == assignment.department,
        ClassTeacher.level == assignment.level
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="Teacher already assigned to this class")
    
    # Create assignment
    class_teacher = ClassTeacher(
        teacher_id=assignment.teacher_id,
        department=assignment.department,
        level=assignment.level,
        assigned_by=current_user.id
    )
    db.add(class_teacher)
    
    # Mark teacher as class teacher
    teacher.is_class_teacher = True
    
    # Create notification for teacher
    notification = Notification(
        user_id=assignment.teacher_id,
        title="🏫 Class Teacher Assignment",
        message=f"You have been assigned as class teacher for {assignment.department} - {assignment.level}. You can now upload student lists for this class.",
        type="class_teacher_assignment"
    )
    db.add(notification)
    
    db.commit()
    
    return {"message": "Class teacher assigned successfully"}

@app.get("/admin/class-teachers")
def get_class_teachers(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    assignments = db.query(ClassTeacher, User).join(User, ClassTeacher.teacher_id == User.id).all()
    return [{
        "id": assignment.id,
        "teacher": {
            "id": teacher.id,
            "username": teacher.username,
            "full_name": teacher.full_name
        },
        "department": assignment.department,
        "level": assignment.level,
        "assigned_at": assignment.assigned_at
    } for assignment, teacher in assignments]

@app.post("/teacher/upload-students")
def upload_class_students(students_data: StudentUpload, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Teacher access required")
    
    created_count = 0
    updated_count = 0
    
    for student_data in students_data.students:
        if not all(k in student_data for k in ["username", "full_name", "department", "level"]):
            continue
        
        department = student_data["department"]
        level = student_data["level"]
        
        # Use original username with default password
        existing_student = db.query(User).filter(User.username == student_data["username"]).first()
        
        if existing_student:
            existing_student.full_name = student_data["full_name"]
            existing_student.department = department
            existing_student.level = level
            updated_count += 1
        else:
            new_student = User(
                username=student_data["username"],
                password_hash=get_password_hash("student123"),
                role="student",
                full_name=student_data["full_name"],
                department=department,
                level=level
            )
            db.add(new_student)
            created_count += 1
    
    # Notify admin about student upload
    if created_count > 0 or updated_count > 0:
        admins = db.query(User).filter(User.role == "admin").all()
        for admin in admins:
            notification = Notification(
                user_id=admin.id,
                title="📚 Students Uploaded",
                message=f"Teacher {current_user.full_name or current_user.username} uploaded {created_count + updated_count} students for {department} - {level}. Created: {created_count}, Updated: {updated_count}",
                type="student_upload"
            )
            db.add(notification)
    
    db.commit()
    return {"message": f"Students uploaded: {created_count} created, {updated_count} updated"}

@app.post("/teacher/upload-students-file")
async def upload_students_file(file: UploadFile = File(...), current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Teacher access required")
    
    if not UPLOAD_ENABLED:
        raise HTTPException(status_code=501, detail="Document upload not available - missing dependencies")
    
    try:
        content = await file.read()
        students_text = ""
        
        if file.filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                students_text += page.extract_text() + "\n"
        elif file.filename.endswith(('.doc', '.docx')):
            doc = docx.Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                students_text += paragraph.text + "\n"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        # Use default values - teacher can modify these in frontend
        default_dept = "Software Development"
        default_level = "Level 5"
        
        print(f"Teacher {current_user.username} processing student file with defaults: {default_dept} - {default_level}")
        
        # Parse student data from text - improved parsing
        lines = students_text.split('\n')
        students = []
        
        # Skip header lines and find actual student data
        for line in lines:
            line = line.strip()
            
            # Skip empty lines, headers, logos, titles
            if not line or len(line) < 3:
                continue
            if any(word in line.lower() for word in ['s/n', 'name', 'student', 'list', 'school', 'college', 'university', 'department', 'class', 'logo', 'header']):
                continue
            if line.startswith('_') or line.startswith('-') or line.startswith('='):
                continue
            
            # Look for patterns that indicate student data
            # Pattern 1: Number followed by name
            import re
            number_name_pattern = r'^(\d+)\s+(.+)$'
            match = re.match(number_name_pattern, line)
            
            if match:
                sn = match.group(1)
                name = match.group(2).strip()
                
                # Clean up name (remove extra spaces, dots, etc.)
                name = re.sub(r'\s+', ' ', name).strip()
                if len(name) > 2:  # Valid name should be more than 2 characters
                    base_username = f"student{sn.zfill(3)}"
                    students.append({
                        "username": base_username,
                        "full_name": name,
                        "department": default_dept,
                        "level": default_level,
                        "password": "student123"
                    })
                continue
            
            # Pattern 2: Try other separators
            parts = None
            if '\t' in line:
                parts = [p.strip() for p in line.split('\t') if p.strip()]
            elif '  ' in line:
                parts = [p.strip() for p in line.split('  ') if p.strip()]
            else:
                parts = line.split()
            
            if len(parts) >= 2 and parts[0].isdigit():
                sn = parts[0]
                name = ' '.join(parts[1:]).strip()
                
                # Clean and validate name
                name = re.sub(r'\s+', ' ', name).strip()
                if len(name) > 2:
                    base_username = f"student{sn.zfill(3)}"
                    students.append({
                        "username": base_username,
                        "full_name": name,
                        "department": default_dept,
                        "level": default_level,
                        "password": "student123"
                    })
        
        return {"students": students, "count": len(students)}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")


# Admin: Upload students via Excel (canonical route used by admin panel)
try:
    # When executed as a package (uvicorn backend.main:app)
    from backend.student_import import parse_student_excel
except Exception:
    try:
        # Fallback when backend is on sys.path without package context
        from student_import import parse_student_excel  # type: ignore
    except Exception:
        parse_student_excel = None  # Will be checked at runtime


@app.post("/admin/upload-students-excel")
async def admin_upload_students_excel(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if not file.filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")

    if parse_student_excel is None:
        raise HTTPException(status_code=500, detail="Student import utility not available")

    try:
        content = await file.read()
        # Parse Excel and get normalized student records
        result = parse_student_excel(io.BytesIO(content))

        created_count = 0
        updated_count = 0
        errors = []

        for student_data in result.get("students", []):
            try:
                username = student_data.get("username")
                if not username:
                    continue
                existing = db.query(User).filter(User.username == username).first()
                if existing:
                    existing.full_name = student_data.get("full_name")
                    existing.department = student_data.get("department")
                    existing.level = student_data.get("level")
                    updated_count += 1
                else:
                    new_student = User(
                        username=username,
                        password_hash=hash_password_simple(student_data.get("password", "student123")),
                        role="student",
                        full_name=student_data.get("full_name"),
                        department=student_data.get("department"),
                        level=student_data.get("level"),
                    )
                    db.add(new_student)
                    created_count += 1
            except Exception as e:
                try:
                    errors.append(f"{student_data.get('full_name', username)}: {str(e)}")
                except Exception:
                    errors.append(str(e))

        db.commit()

        return {
            "success": True,
            "message": f"Successfully imported {created_count + updated_count} students",
            "created": created_count,
            "updated": updated_count,
            "class_group": result.get("class_group"),
            "department": result.get("department"),
            "level": result.get("level"),
            "total": result.get("count", 0),
            "errors": errors,
        }
    except HTTPException:
        # Re-raise explicit HTTP errors
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Import failed: {str(e)}")

@app.get("/teacher/my-classes")
def get_my_classes(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "teacher":
        raise HTTPException(status_code=403, detail="Teacher access required")
    
    assignments = db.query(ClassTeacher).filter(ClassTeacher.teacher_id == current_user.id).all()
    return [{
        "id": assignment.id,
        "department": assignment.department,
        "level": assignment.level,
        "assigned_at": assignment.assigned_at
    } for assignment in assignments]

@app.post("/reset-admin")
def reset_admin_password(db: Session = Depends(get_db)):
    admin = db.query(User).filter(User.username == "admin").first()
    if not admin:
        admin = User(
            username="admin",
            password_hash=hash_password_simple("admin123"),
            role="admin",
            full_name="DOS Administrator",
            is_class_teacher=False
        )
        db.add(admin)
    else:
        admin.password_hash = hash_password_simple("admin123")
    
    db.commit()
    return {"message": "Admin password reset successfully"}

@app.post("/reset-teacher-password")
def reset_teacher_password(username: str, db: Session = Depends(get_db)):
    teacher = db.query(User).filter(User.username == username, User.role == "teacher").first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    
    # Reset to default password
    teacher.password_hash = hash_password_simple("pass123")
    db.commit()
    
    return {"message": f"Teacher {username} password reset to 'pass123' successfully"}

@app.post("/reset-all-teacher-passwords")
def reset_all_teacher_passwords(db: Session = Depends(get_db)):
    teachers = db.query(User).filter(User.role == "teacher").all()
    count = 0
    
    for teacher in teachers:
        teacher.password_hash = hash_password_simple("pass123")
        count += 1
    
    db.commit()
    return {"message": f"Reset passwords for {count} teachers to 'pass123' successfully"}

class TeacherUpdate(BaseModel):
    full_name: str = None
    password: str = None
    departments: List[str] = None

@app.put("/admin/teacher/{teacher_id}")
def update_teacher(teacher_id: int, teacher_data: TeacherUpdate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    teacher = db.query(User).filter(User.id == teacher_id, User.role == "teacher").first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    
    if teacher_data.full_name:
        teacher.full_name = teacher_data.full_name
    if teacher_data.password:
        teacher.password_hash = get_password_hash(teacher_data.password)
    if teacher_data.departments:
        teacher.departments = teacher_data.departments
    
    db.commit()
    return {"message": "Teacher updated successfully"}

@app.post("/admin/reset-teacher-password/{teacher_id}")
def reset_teacher_password_by_id(teacher_id: int, new_password: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    teacher = db.query(User).filter(User.id == teacher_id, User.role == "teacher").first()
    if not teacher:
        raise HTTPException(status_code=404, detail="Teacher not found")
    
    teacher.password_hash = get_password_hash(new_password)
    db.commit()
    
    return {"message": f"Password reset successfully for {teacher.username}"}

@app.delete("/admin/clear-all-students")
def clear_all_students(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    try:
        students = db.query(User).filter(User.role == "student").all()
        student_count = len(students)
        
        # Update questions/quizzes created by students to be owned by current admin
        for student in students:
            db.query(Question).filter(Question.created_by == student.id).update({"created_by": current_user.id})
            db.query(Quiz).filter(Quiz.created_by == student.id).update({"created_by": current_user.id})
            
            # Delete student_answers first (child table)
            attempt_ids = [a.id for a in db.query(QuizAttempt).filter(QuizAttempt.user_id == student.id).all()]
            if attempt_ids:
                db.query(StudentAnswer).filter(StudentAnswer.attempt_id.in_(attempt_ids)).delete(synchronize_session=False)
            
            # Then delete quiz_attempts (parent table)
            db.query(QuizAttempt).filter(QuizAttempt.user_id == student.id).delete()
            db.query(Notification).filter(Notification.user_id == student.id).delete()
        
        # Delete students
        db.query(User).filter(User.role == "student").delete()
        
        db.commit()
        return {"message": f"Cleared {student_count} students from the system", "count": student_count}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to clear students: {str(e)}")

@app.get("/schedules/{schedule_id}/download")
def download_schedule_file(schedule_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    schedule = db.query(Schedule).filter(Schedule.id == schedule_id).first()
    if not schedule or not schedule.file_data:
        raise HTTPException(status_code=404, detail="Schedule file not found")
    
    try:
        import base64
        
        # Decode base64 file data
        file_bytes = base64.b64decode(schedule.file_data)
        
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=schedule.file_type or "application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename={schedule.file_name}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@app.get("/health")
def health_check():
    """Basic health check endpoint for connection testing"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "service": "Morning Quiz API"
    }

@app.get("/test-admin")
def test_admin(db: Session = Depends(get_db)):
    """Test endpoint to check admin user"""
    user = db.query(User).filter(User.username == "admin").first()
    if user:
        return {
            "found": True,
            "id": user.id,
            "username": user.username,
            "role": user.role,
            "full_name": user.full_name,
            "departments": user.departments
        }
    return {"found": False}

@app.get("/auth/test")
def test_auth(current_user: User = Depends(get_current_user)):
    return {
        "message": "Authentication successful",
        "user": {
            "id": current_user.id,
            "username": current_user.username,
            "role": current_user.role,
            "full_name": current_user.full_name
        }
    }

@app.get("/quiz/{quiz_id}/students-results/excel")
def download_quiz_students_results_excel(quiz_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Get quiz details
    quiz = db.query(Quiz).filter(Quiz.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    
    # Check if teacher owns the quiz
    if current_user.role == "teacher" and quiz.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="You can only download results for your own quizzes")
    
    # Get all students and their results
    results = db.query(QuizAttempt, User).join(User, QuizAttempt.user_id == User.id).filter(
        QuizAttempt.quiz_id == quiz_id
    ).order_by(QuizAttempt.score.desc()).all()
    
    # Get all students in the same department/level (including those who didn't attempt)
    all_students = db.query(User).filter(
        User.role == "student",
        User.department == quiz.department,
        User.level == quiz.level
    ).order_by(User.full_name).all()
    
    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Student Results"
    
    # Headers
    headers = ['Rank', 'Student ID', 'Full Name', 'Department', 'Level', 'Score', 'Total', 'Percentage', 'Grade', 'Status', 'Completed At']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    # Create results lookup
    results_dict = {attempt.user_id: (attempt, user) for attempt, user in results}
    
    # Data rows
    row = 2
    rank = 1
    
    # First, add students who attempted the quiz (ranked)
    for attempt, user in results:
        percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
        
        # Calculate grade
        if percentage >= 90: grade = "A+"
        elif percentage >= 80: grade = "A"
        elif percentage >= 70: grade = "B"
        elif percentage >= 60: grade = "C"
        elif percentage >= 50: grade = "D"
        else: grade = "F"
        
        data = [
            rank,
            user.username,
            user.full_name or user.username,
            user.department,
            user.level,
            attempt.score,
            attempt.total_questions,
            f"{percentage}%",
            grade,
            "Completed",
            attempt.completed_at.strftime('%Y-%m-%d %H:%M:%S') if attempt.completed_at else 'N/A'
        ]
        
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row, column=col, value=value)
            cell.alignment = Alignment(horizontal="center")
            
            # Color coding based on grade
            if grade in ["A+", "A"]:
                cell.fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
            elif grade in ["B", "C"]:
                cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
            elif grade in ["D", "F"]:
                cell.fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
        
        row += 1
        rank += 1
    
    # Then, add students who didn't attempt (unranked)
    for student in all_students:
        if student.id not in results_dict:
            data = [
                "-",
                student.username,
                student.full_name or student.username,
                student.department,
                student.level,
                0,
                quiz.duration_minutes if hasattr(quiz, 'total_questions') else "-",
                "0%",
                "F",
                "Not Attempted",
                "-"
            ]
            
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.alignment = Alignment(horizontal="center")
                cell.fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
            
            row += 1
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 20)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Add summary at the bottom
    summary_row = row + 2
    ws.cell(row=summary_row, column=1, value="SUMMARY").font = Font(bold=True)
    ws.cell(row=summary_row + 1, column=1, value=f"Quiz: {quiz.title}")
    ws.cell(row=summary_row + 2, column=1, value=f"Department: {quiz.department}")
    ws.cell(row=summary_row + 3, column=1, value=f"Level: {quiz.level}")
    ws.cell(row=summary_row + 4, column=1, value=f"Total Students: {len(all_students)}")
    ws.cell(row=summary_row + 5, column=1, value=f"Attempted: {len(results)}")
    ws.cell(row=summary_row + 6, column=1, value=f"Not Attempted: {len(all_students) - len(results)}")
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=quiz_{quiz_id}_students_results.xlsx"}
    )

@app.put("/questions/{question_id}")
def update_question(question_id: int, question: QuestionCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Get existing question
    db_question = db.query(Question).filter(Question.id == question_id).first()
    if not db_question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    # Check ownership for teachers
    if current_user.role == "teacher" and db_question.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="You can only edit your own questions")
    
    # Validate lesson assignment for teachers
    if current_user.role == "teacher":
        assignment = db.query(TeacherLesson).filter(
            TeacherLesson.teacher_id == current_user.id,
            TeacherLesson.lesson_id == question.lesson_id
        ).first()
        if not assignment:
            raise HTTPException(status_code=403, detail="You can only create questions for lessons assigned to you")
    
    # Validate lesson exists
    lesson = db.query(Lesson).filter(
        Lesson.id == question.lesson_id,
        Lesson.department == question.department,
        Lesson.level == question.level,
        Lesson.is_active == True
    ).first()
    if not lesson:
        raise HTTPException(status_code=400, detail=f"Invalid lesson selection for {question.department} - {question.level}")
    
    # Update question
    db_question.question_text = question.question_text
    db_question.question_type = question.question_type
    db_question.options = question.options if question.question_type != 'short_answer' else []
    db_question.correct_answer = question.correct_answer
    db_question.points = question.points
    db_question.department = question.department
    db_question.level = question.level
    db_question.lesson_id = question.lesson_id
    
    db.commit()
    db.refresh(db_question)
    return db_question

@app.delete("/questions/{question_id}")
def delete_question(question_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Get existing question
    db_question = db.query(Question).filter(Question.id == question_id).first()
    if not db_question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    # Check ownership for teachers
    if current_user.role == "teacher" and db_question.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="You can only delete your own questions")
    
    # Check if question is used in any quiz
    quiz_usage = db.query(QuizQuestion).filter(QuizQuestion.question_id == question_id).first()
    if quiz_usage:
        raise HTTPException(status_code=400, detail="Cannot delete question that is used in a quiz")
    
    db.delete(db_question)
    db.commit()
    return {"message": "Question deleted successfully"}

@app.post("/admin/generate-student-credentials/{department}/{level}")
def generate_student_credentials_pdf(department: str, level: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Get students for the specified department and level
    students = db.query(User).filter(
        User.role == "student",
        User.department == department,
        User.level == level
    ).order_by(User.full_name).all()
    
    if not students:
        raise HTTPException(status_code=404, detail=f"No students found for {department} - {level}")
    
    # Create professional PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=20,
        alignment=1,  # Center
        textColor=colors.HexColor('#1f2937'),
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=15,
        alignment=1,
        textColor=colors.HexColor('#374151'),
        fontName='Helvetica-Bold'
    )
    
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        textColor=colors.HexColor('#1f2937'),
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        textColor=colors.HexColor('#374151'),
        fontName='Helvetica'
    )
    
    # Header with school branding
    story.append(Paragraph("🎓 MORNING QUIZ SYSTEM", title_style))
    story.append(Paragraph("Student Login Credentials", subtitle_style))
    story.append(Spacer(1, 20))
    
    # Class information
    class_info_style = ParagraphStyle(
        'ClassInfo',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8,
        alignment=1,
        textColor=colors.HexColor('#059669'),
        fontName='Helvetica-Bold'
    )
    
    story.append(Paragraph(f"📚 Department: {department}", class_info_style))
    story.append(Paragraph(f"🎯 Level: {level}", class_info_style))
    story.append(Paragraph(f"👥 Total Students: {len(students)}", class_info_style))
    story.append(Spacer(1, 25))
    
    # Instructions for class teacher
    story.append(Paragraph("📋 Instructions for Class Teacher", header_style))
    
    instructions = [
        "• Distribute these credentials to students individually",
        "• Ensure students keep their login details secure",
        "• Students should change passwords after first login (optional)",
        "• Report any login issues to the IT department immediately",
        "• Keep a copy of this document for your records"
    ]
    
    for instruction in instructions:
        story.append(Paragraph(instruction, body_style))
    
    story.append(Spacer(1, 20))
    
    # Student login instructions
    story.append(Paragraph("🔐 Student Login Instructions", header_style))
    
    login_steps = [
        "1. Open your web browser (Chrome, Firefox, Safari, etc.)",
        "2. Navigate to the Morning Quiz System website",
        "3. Click on 'Student Login' or go directly to the student portal",
        "4. Enter your Student ID as username",
        "5. Enter the default password: student123",
        "6. Click 'Login' to access your dashboard",
        "7. You can now participate in quizzes and view your results"
    ]
    
    for step in login_steps:
        story.append(Paragraph(step, body_style))
    
    story.append(Spacer(1, 25))
    
    # Student credentials table
    story.append(Paragraph("👥 Student Login Credentials", header_style))
    story.append(Spacer(1, 10))
    
    # Create table data
    table_data = [['#', 'Student ID', 'Full Name', 'Default Password']]
    
    for i, student in enumerate(students, 1):
        table_data.append([
            str(i),
            student.username,
            student.full_name or student.username,
            'student123'
        ])
    
    # Create and style the table
    table = Table(table_data, colWidths=[0.8*inch, 1.8*inch, 3*inch, 1.5*inch])
    table.setStyle(TableStyle([
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        
        # Data rows styling
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f9fafb')),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d1d5db')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f9fafb'), colors.HexColor('#ffffff')]),
    ]))
    
    story.append(table)
    story.append(Spacer(1, 30))
    
    # Security notice
    story.append(Paragraph("🔒 Security Guidelines", header_style))
    
    security_notes = [
        "⚠️ Keep login credentials confidential and secure",
        "🚫 Do not share your username and password with others",
        "💻 Always log out after using the system",
        "📱 Use only school-approved devices for accessing the system",
        "🆘 Report any suspicious activity to your class teacher immediately"
    ]
    
    for note in security_notes:
        story.append(Paragraph(note, body_style))
    
    story.append(Spacer(1, 30))
    
    # Footer with generation info
    footer_style = ParagraphStyle(
        'FooterStyle',
        parent=styles['Normal'],
        fontSize=9,
        alignment=1,
        textColor=colors.HexColor('#6b7280'),
        fontName='Helvetica-Oblique'
    )
    
    story.append(Paragraph(f"Generated on: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}", footer_style))
    story.append(Paragraph(f"Generated by: {current_user.full_name or current_user.username} (DOS)", footer_style))
    story.append(Paragraph("Morning Quiz System - Professional Student Management", footer_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=Student_Credentials_{department.replace(' ', '_')}_{level.replace(' ', '_')}.pdf"}
    )

# Health check for offline functionality
@app.get("/offline-status")
def offline_status():
    """Check if system is running offline"""
    return {
        "status": "online",
        "offline_capable": True,
        "timestamp": datetime.utcnow().isoformat(),
        "message": "Morning Quiz System - Offline First Architecture"
    }

@app.get("/admin/quiz-reports/{report_type}")
def generate_quiz_report(report_type: str, department: str = None, level: str = None, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    # Calculate date ranges
    now = datetime.utcnow()
    if report_type == "daily":
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + timedelta(days=1)
        title = f"Daily Quiz Report - {start_date.strftime('%Y-%m-%d')}"
    elif report_type == "weekly":
        days_since_monday = now.weekday()
        start_date = (now - timedelta(days=days_since_monday)).replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + timedelta(days=7)
        title = f"Weekly Quiz Report - {start_date.strftime('%Y-%m-%d')} to {(end_date-timedelta(days=1)).strftime('%Y-%m-%d')}"
    elif report_type == "monthly":
        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        if now.month == 12:
            end_date = start_date.replace(year=now.year + 1, month=1)
        else:
            end_date = start_date.replace(month=now.month + 1)
        title = f"Monthly Quiz Report - {start_date.strftime('%B %Y')}"
    else:
        raise HTTPException(status_code=400, detail="Invalid report type. Use: daily, weekly, or monthly")
    
    # Build query
    query = db.query(QuizAttempt, Quiz, User).join(
        Quiz, QuizAttempt.quiz_id == Quiz.id
    ).join(
        User, QuizAttempt.user_id == User.id
    ).filter(
        QuizAttempt.completed_at >= start_date,
        QuizAttempt.completed_at < end_date
    )
    
    if department:
        query = query.filter(Quiz.department == department)
    if level:
        query = query.filter(Quiz.level == level)
    
    attempts = query.order_by(QuizAttempt.completed_at.desc()).all()
    
    # Generate PDF report
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20, alignment=1, textColor=colors.HexColor('#1f2937'))
    story.append(Paragraph(title, title_style))
    
    if department or level:
        filter_text = f"Filtered by: {department or 'All Departments'} - {level or 'All Levels'}"
        story.append(Paragraph(filter_text, styles['Normal']))
    
    story.append(Spacer(1, 20))
    
    # Summary statistics
    total_attempts = len(attempts)
    if attempts:
        avg_score = sum(a[0].score for a in attempts) / len(attempts)
        avg_percentage = sum((a[0].score / a[0].total_questions) * 100 for a in attempts) / len(attempts)
        unique_students = len(set(a[2].id for a in attempts))
        unique_quizzes = len(set(a[1].id for a in attempts))
    else:
        avg_score = avg_percentage = unique_students = unique_quizzes = 0
    
    summary_data = [
        ['Metric', 'Value'],
        ['Total Quiz Attempts', str(total_attempts)],
        ['Unique Students', str(unique_students)],
        ['Unique Quizzes', str(unique_quizzes)],
        ['Average Score', f"{avg_score:.1f}"],
        ['Average Percentage', f"{avg_percentage:.1f}%"]
    ]
    
    summary_table = Table(summary_data, colWidths=[2*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Detailed results
    if attempts:
        story.append(Paragraph("Detailed Quiz Results", styles['Heading2']))
        story.append(Spacer(1, 10))
        
        data = [['Student', 'Quiz', 'Department', 'Level', 'Score', 'Percentage', 'Date']]
        for attempt, quiz, user in attempts:
            percentage = round((attempt.score / attempt.total_questions) * 100, 1) if attempt.total_questions > 0 else 0
            data.append([
                user.full_name or user.username,
                quiz.title[:25] + '...' if len(quiz.title) > 25 else quiz.title,
                quiz.department,
                quiz.level,
                f"{attempt.score}/{attempt.total_questions}",
                f"{percentage}%",
                attempt.completed_at.strftime('%m/%d %H:%M')
            ])
        
        results_table = Table(data, colWidths=[1.2*inch, 1.5*inch, 1*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.9*inch])
        results_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(results_table)
    else:
        story.append(Paragraph("No quiz attempts found for this period.", styles['Normal']))
    
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Generated: {now.strftime('%Y-%m-%d %H:%M:%S')} UTC by {current_user.full_name}", styles['Italic']))
    
    doc.build(story)
    buffer.seek(0)
    
    filename = f"{report_type}_quiz_report_{now.strftime('%Y%m%d')}"
    if department:
        filename += f"_{department.replace(' ', '_')}"
    if level:
        filename += f"_{level.replace(' ', '_')}"
    filename += ".pdf"
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

class QuizReportFilter(BaseModel):
    report_type: str  # daily, weekly, monthly
    department: Optional[str] = None
    level: Optional[str] = None

class SessionPlanCreate(BaseModel):
    title: str
    lesson_id: int
    department: str
    level: str
    duration_minutes: int = 90
    learning_objectives: List[str]
    teaching_methods: List[str]
    resources_required: List[str]
    assessment_methods: List[str]
    content_outline: str = ""
    activities: List[dict] = []
    homework_assignment: str = ""

class SchemeOfWorkCreate(BaseModel):
    title: str
    lesson_id: int
    department: str
    level: str
    academic_year: str
    term: str
    total_weeks: int = 12
    learning_outcomes: List[str]
    assessment_schedule: List[dict] = []
    resources_list: List[str] = []

class AIGenerateRequest(BaseModel):
    generation_type: str  # session_plan, scheme_of_work
    lesson_title: str
    department: str
    level: str
    duration_minutes: Optional[int] = 90
    academic_year: Optional[str] = None
    term: Optional[str] = None
    total_weeks: Optional[int] = 12
    additional_requirements: Optional[str] = ""

# Import performance reports module
try:
    from performance_reports import get_student_performance, generate_department_report
except ImportError:
    def get_student_performance(db, user_id):
        return {"totalQuizzes": 0, "averageScore": 0, "bestScore": 0, "improvement": 0, "quizHistory": []}
    
    def generate_department_report(db, department, level, report_type, date):
        return {"department": department, "level": level, "statistics": {"total_students": 0, "total_quizzes": 0, "average_score": 0, "highest_score": 0, "lowest_score": 0}, "results": []}

# Student performance endpoint
@app.get("/student/performance")
def get_student_performance_data(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "student":
        raise HTTPException(status_code=403, detail="Student access required")
    
    return get_student_performance(db, current_user.id)

# Department reports endpoints
@app.get("/admin/reports/department")
def download_department_report_pdf(department: str, level: str, reportType: str, date: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    report_data = generate_department_report(db, department, level, reportType, date)
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20, alignment=1)
    story.append(Paragraph(f"Department Report - {department} {level}", title_style))
    story.append(Paragraph(f"{reportType.title()} Report for {report_data['period']}", styles['Heading2']))
    story.append(Spacer(1, 20))
    
    # Statistics
    stats = report_data['statistics']
    stats_data = [
        ['Metric', 'Value'],
        ['Total Students', str(stats['total_students'])],
        ['Total Quizzes', str(stats['total_quizzes'])],
        ['Average Score', f"{stats['average_score']}%"],
        ['Highest Score', f"{stats['highest_score']}%"],
        ['Lowest Score', f"{stats['lowest_score']}%"]
    ]
    
    stats_table = Table(stats_data)
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 20))
    
    # Results
    if report_data['results']:
        story.append(Paragraph("Quiz Results", styles['Heading2']))
        results_data = [['Student', 'Quiz', 'Score', 'Date']]
        for result in report_data['results']:
            results_data.append([
                result['full_name'],
                result['quiz_title'][:30],
                f"{result['score']}%",
                result['submitted_at'][:10]
            ])
        
        results_table = Table(results_data)
        results_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(results_table)
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={department}_{level}_{reportType}_report.pdf"}
    )

@app.get("/admin/reports/department/excel")
def download_department_report_excel(department: str, level: str, reportType: str, date: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="DOS access required")
    
    report_data = generate_department_report(db, department, level, reportType, date)
    
    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Department Report"
    
    # Headers
    headers = ['Student Name', 'Quiz Title', 'Score', 'Total Questions', 'Percentage', 'Date']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
    
    # Data
    for row, result in enumerate(report_data['results'], 2):
        ws.cell(row=row, column=1, value=result['full_name'])
        ws.cell(row=row, column=2, value=result['quiz_title'])
        ws.cell(row=row, column=3, value=result['score'])
        ws.cell(row=row, column=4, value=result['total_questions'])
        ws.cell(row=row, column=5, value=f"{result['score']}%")
        ws.cell(row=row, column=6, value=result['submitted_at'][:10])
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={department}_{level}_{reportType}_report.xlsx"}
    )

# AI Generation Functions
def generate_session_plan_ai(lesson_title: str, department: str, level: str, duration_minutes: int = 90, additional_requirements: str = ""):
    """Generate RTB-compliant session plan using exact official template structure"""
    
    from rtb_template_parser import parse_rtb_session_plan_template, generate_rtb_module_code
    
    # Get RTB official template structure
    template = parse_rtb_session_plan_template()
    
    # Generate module code
    module_code = generate_rtb_module_code(department, level, 1)
    
    # Fill in the template with AI-generated content
    session_plan = {
        "header": {
            "institution_logo": "RTB_LOGO",
            "title": "RWANDA TRAINING BOARD (RTB)",
            "subtitle": "SESSION PLAN",
            "form_number": "RTB/TVET/SP/001"
        },
        "basic_info": {
            "institution_name": "RTB Affiliated TVET Institution",
            "department": department,
            "level": level,
            "instructor_name": "[To be filled by instructor]",
            "lesson_title": lesson_title,
            "module_code": module_code,
            "date": datetime.utcnow().strftime("%d/%m/%Y"),
            "duration": f"{duration_minutes} minutes",
            "class_size": "[Number of students]"
        },
        "learning_objectives": {
            "title": "LEARNING OBJECTIVES",
            "description": "By the end of this session, students will be able to:",
            "objectives": generate_rtb_learning_objectives(lesson_title, department, level)
        },
        "teaching_methods": {
            "title": "TEACHING METHODS",
            "methods": generate_rtb_teaching_methods(department, level)
        },
        "resources_required": {
            "title": "RESOURCES REQUIRED",
            "resources": generate_rtb_resources(lesson_title, department, level)
        },
        "lesson_structure": {
            "title": "LESSON STRUCTURE",
            "phases": generate_rtb_lesson_structure(lesson_title, duration_minutes)
        },
        "assessment": {
            "title": "ASSESSMENT",
            "formative": generate_rtb_formative_assessment(department, level),
            "summative": generate_rtb_summative_assessment(department, level)
        },
        "homework": {
            "title": "HOMEWORK/ASSIGNMENT",
            "assignment": generate_rtb_homework(lesson_title, department, level)
        },
        "reflection": {
            "title": "TEACHER REFLECTION",
            "what_went_well": "[To be completed after lesson delivery]",
            "areas_for_improvement": "[To be completed after lesson delivery]",
            "student_engagement": "[To be completed after lesson delivery]",
            "objectives_achieved": "[To be completed after lesson delivery]",
            "next_lesson_preparation": "[To be completed after lesson delivery]"
        },
        "signatures": {
            "instructor_signature": "Instructor Signature: _________________ Date: _______",
            "hod_signature": "HOD Signature: _________________ Date: _______"
        }
    }
    
    return session_plan

def generate_scheme_of_work_ai(lesson_title: str, department: str, level: str, academic_year: str, term: str, total_weeks: int = 12, additional_requirements: str = ""):
    """Generate RTB-compliant scheme of work using exact official template structure"""
    
    from rtb_template_parser import parse_rtb_scheme_of_work_template, generate_rtb_module_code
    
    # Get RTB official template structure
    template = parse_rtb_scheme_of_work_template()
    
    # Generate module code
    module_code = generate_rtb_module_code(department, level, 1)
    
    scheme_of_work = {
        "header": {
            "institution_logo": "RTB_LOGO",
            "title": "RWANDA TRAINING BOARD (RTB)",
            "subtitle": "SCHEME OF WORK",
            "form_number": "RTB/TVET/SOW/001"
        },
        "basic_info": {
            "institution_name": "RTB Affiliated TVET Institution",
            "department": department,
            "level": level,
            "subject": lesson_title,
            "module_code": module_code,
            "instructor_name": "[To be filled by instructor]",
            "academic_year": academic_year,
            "term": term,
            "total_weeks": total_weeks,
            "total_hours": total_weeks * 4  # Assuming 4 hours per week
        },
        "course_overview": {
            "title": "COURSE OVERVIEW",
            "description": generate_rtb_course_overview(lesson_title, department, level),
            "prerequisites": generate_rtb_prerequisites(lesson_title, department, level),
            "learning_outcomes": generate_rtb_course_outcomes(lesson_title, department, level)
        },
        "weekly_breakdown": {
            "title": "WEEKLY BREAKDOWN",
            "weeks": generate_rtb_weekly_breakdown(lesson_title, department, level, total_weeks)
        },
        "assessment_schedule": {
            "title": "ASSESSMENT SCHEDULE",
            "assessments": generate_rtb_assessment_schedule(total_weeks)
        },
        "resources": {
            "title": "REQUIRED RESOURCES",
            "textbooks": generate_rtb_textbooks(lesson_title, department, level),
            "equipment": generate_rtb_equipment(lesson_title, department, level),
            "materials": generate_rtb_materials(lesson_title, department, level),
            "digital_resources": generate_rtb_digital_resources(lesson_title, department, level)
        },
        "evaluation_criteria": {
            "title": "EVALUATION CRITERIA",
            "knowledge_understanding": "25% - Demonstration of theoretical knowledge and conceptual understanding",
            "practical_skills": "35% - Application of skills in practical situations and problem-solving",
            "analysis_evaluation": "20% - Critical thinking, analysis and evaluation of solutions",
            "communication_collaboration": "20% - Professional communication and teamwork skills"
        },
        "signatures": {
            "instructor_signature": "Instructor Signature: _________________ Date: _______",
            "hod_signature": "HOD Signature: _________________ Date: _______",
            "dos_signature": "DOS Signature: _________________ Date: _______"
        }
    }
    
    return scheme_of_work

def generate_learning_objectives(lesson_title: str, department: str, level: str):
    """Generate specific learning objectives based on RTB standards"""
    
    objectives_map = {
        "Software Development": {
            "Level 3": [
                "Understand fundamental programming concepts and syntax",
                "Apply problem-solving techniques in software development",
                "Demonstrate proficiency in debugging and testing code",
                "Create functional software applications using best practices"
            ],
            "Level 4": [
                "Design and implement complex software solutions",
                "Apply advanced programming paradigms and patterns",
                "Evaluate and optimize software performance",
                "Collaborate effectively in software development teams"
            ],
            "Level 5": [
                "Architect scalable and maintainable software systems",
                "Lead software development projects and teams",
                "Implement industry-standard security and quality measures",
                "Innovate and adapt to emerging technologies"
            ]
        },
        "Computer System and Architecture": {
            "Level 3": [
                "Understand computer hardware components and their functions",
                "Analyze system performance and bottlenecks",
                "Configure and maintain computer systems",
                "Troubleshoot hardware and software issues"
            ],
            "Level 4": [
                "Design efficient computer system architectures",
                "Implement system optimization strategies",
                "Evaluate emerging hardware technologies",
                "Manage complex IT infrastructure"
            ],
            "Level 5": [
                "Architect enterprise-level computing solutions",
                "Lead system integration and migration projects",
                "Develop strategic IT infrastructure plans",
                "Research and implement cutting-edge technologies"
            ]
        }
    }
    
    base_objectives = objectives_map.get(department, {}).get(level, [
        "Understand core concepts and principles",
        "Apply theoretical knowledge to practical situations",
        "Demonstrate competency in relevant skills",
        "Evaluate and reflect on learning outcomes"
    ])
    
    # Customize objectives based on lesson title
    customized_objectives = []
    for obj in base_objectives:
        if "programming" in lesson_title.lower():
            obj = obj.replace("concepts", "programming concepts")
        elif "database" in lesson_title.lower():
            obj = obj.replace("concepts", "database concepts")
        elif "network" in lesson_title.lower():
            obj = obj.replace("concepts", "networking concepts")
        customized_objectives.append(obj)
    
    return customized_objectives

def generate_teaching_methods(department: str, level: str):
    """Generate appropriate teaching methods for RTB standards"""
    
    methods = [
        "Interactive lectures with multimedia presentations",
        "Hands-on practical exercises and demonstrations",
        "Group discussions and collaborative learning",
        "Case study analysis and problem-solving sessions",
        "Individual and group project work",
        "Peer review and feedback sessions",
        "Industry guest speaker presentations",
        "Field visits and real-world observations"
    ]
    
    if department == "Software Development":
        methods.extend([
            "Code review and pair programming sessions",
            "Agile development methodology workshops",
            "Version control system training",
            "Software testing and debugging exercises"
        ])
    elif department == "Computer System and Architecture":
        methods.extend([
            "Hardware assembly and disassembly labs",
            "System performance monitoring exercises",
            "Network configuration workshops",
            "Troubleshooting simulation scenarios"
        ])
    
    return methods[:6]  # Return top 6 methods

def generate_resources(lesson_title: str, department: str, level: str):
    """Generate required resources for the lesson"""
    
    basic_resources = [
        "Whiteboard and markers",
        "Projector and screen",
        "Computers/laptops for students",
        "Internet connectivity",
        "Printed handouts and worksheets",
        "Reference textbooks and materials"
    ]
    
    if department == "Software Development":
        basic_resources.extend([
            "Integrated Development Environment (IDE)",
            "Programming language documentation",
            "Code repositories and version control tools",
            "Software development frameworks"
        ])
    elif department == "Computer System and Architecture":
        basic_resources.extend([
            "Computer hardware components",
            "System diagnostic tools",
            "Network cables and equipment",
            "Multimeters and testing equipment"
        ])
    
    return basic_resources

def generate_lesson_structure(lesson_title: str, duration_minutes: int):
    """Generate detailed lesson structure with timing"""
    
    # Calculate time allocation
    intro_time = max(5, duration_minutes // 18)  # ~5-10 minutes
    main_time = duration_minutes - intro_time - 15  # Main content
    conclusion_time = 10
    assessment_time = 5
    
    structure = [
        {
            "phase": "Introduction and Warm-up",
            "duration": f"{intro_time} minutes",
            "activities": [
                "Greet students and take attendance",
                "Review previous lesson concepts",
                "Introduce today's learning objectives",
                "Motivate students with real-world examples"
            ]
        },
        {
            "phase": "Main Content Delivery",
            "duration": f"{main_time} minutes",
            "activities": [
                "Present core concepts with examples",
                "Demonstrate practical applications",
                "Facilitate hands-on exercises",
                "Guide group discussions and activities",
                "Provide individual support as needed"
            ]
        },
        {
            "phase": "Conclusion and Summary",
            "duration": f"{conclusion_time} minutes",
            "activities": [
                "Summarize key learning points",
                "Address student questions and concerns",
                "Preview next lesson topics",
                "Assign homework and additional reading"
            ]
        },
        {
            "phase": "Assessment and Feedback",
            "duration": f"{assessment_time} minutes",
            "activities": [
                "Conduct quick formative assessment",
                "Gather student feedback on lesson",
                "Note areas for improvement"
            ]
        }
    ]
    
    return structure

def generate_assessment_methods(department: str, level: str):
    """Generate appropriate assessment methods"""
    
    methods = [
        "Formative assessment through questioning",
        "Practical exercises and demonstrations",
        "Written quizzes and tests",
        "Project-based assessments",
        "Peer evaluation and self-assessment",
        "Portfolio development and review"
    ]
    
    if department == "Software Development":
        methods.extend([
            "Code review and quality assessment",
            "Software project presentations",
            "Algorithm design challenges"
        ])
    elif department == "Computer System and Architecture":
        methods.extend([
            "Hardware troubleshooting scenarios",
            "System configuration tasks",
            "Performance optimization challenges"
        ])
    
    return methods[:5]

def generate_homework(lesson_title: str, department: str, level: str):
    """Generate relevant homework assignment"""
    
    homework_templates = {
        "Software Development": [
            "Complete the programming exercises from today's lesson",
            "Research and write a summary on advanced {topic} techniques",
            "Develop a small application demonstrating {concept}",
            "Review and practice the coding examples provided"
        ],
        "Computer System and Architecture": [
            "Research the latest developments in {topic} technology",
            "Complete the system configuration worksheet",
            "Analyze a real-world case study related to {concept}",
            "Prepare for next week's practical assessment"
        ]
    }
    
    templates = homework_templates.get(department, [
        "Review today's lesson materials and take notes",
        "Complete the assigned reading from the textbook",
        "Prepare questions for next lesson discussion",
        "Practice the skills demonstrated in class"
    ])
    
    # Customize based on lesson title
    topic = lesson_title.lower().replace("introduction to ", "").replace("advanced ", "")
    concept = topic.split()[0] if topic.split() else "the topic"
    
    selected_template = templates[0]
    homework = selected_template.format(topic=topic, concept=concept)
    
    return homework

def generate_reflection_template():
    """Generate reflection notes template"""
    
    return {
        "what_went_well": "Areas where the lesson was successful and engaging",
        "areas_for_improvement": "Aspects that could be enhanced in future lessons",
        "student_engagement": "Level of student participation and interest",
        "learning_objectives_met": "Assessment of whether objectives were achieved",
        "timing_and_pacing": "Evaluation of lesson timing and content pacing",
        "resources_effectiveness": "How well the resources supported learning",
        "next_lesson_adjustments": "Modifications needed for subsequent lessons"
    }

def generate_course_overview(lesson_title: str, department: str, level: str):
    """Generate comprehensive course overview"""
    
    overview_templates = {
        "Software Development": f"This course provides comprehensive training in {lesson_title}, covering both theoretical foundations and practical applications. Students will develop industry-relevant skills through hands-on projects, collaborative learning, and real-world problem-solving scenarios. The curriculum aligns with RTB standards and prepares students for professional software development careers.",
        "Computer System and Architecture": f"This course explores {lesson_title} with emphasis on practical implementation and industry best practices. Students will gain deep understanding of system components, performance optimization, and troubleshooting techniques. The program combines theoretical knowledge with extensive hands-on experience to prepare graduates for technical leadership roles."
    }
    
    return overview_templates.get(department, f"This comprehensive course in {lesson_title} provides students with essential knowledge and practical skills required for success in the {department} field. The curriculum is designed to meet RTB standards and industry requirements.")

def generate_learning_outcomes(lesson_title: str, department: str, level: str, total_weeks: int):
    """Generate comprehensive learning outcomes for scheme of work"""
    
    outcomes = [
        f"Demonstrate comprehensive understanding of {lesson_title} principles and concepts",
        f"Apply {lesson_title} knowledge to solve real-world problems and challenges",
        f"Analyze and evaluate different approaches and methodologies in {lesson_title}",
        f"Create and implement solutions using {lesson_title} best practices",
        f"Collaborate effectively in team-based {lesson_title} projects",
        f"Communicate technical concepts clearly to diverse audiences",
        f"Demonstrate professional ethics and responsibility in {lesson_title} practice",
        f"Adapt to emerging trends and technologies in {lesson_title}"
    ]
    
    return outcomes

def generate_weekly_breakdown(lesson_title: str, department: str, level: str, total_weeks: int):
    """Generate detailed weekly breakdown for scheme of work"""
    
    weekly_plan = []
    
    # Week 1: Introduction and Foundations
    weekly_plan.append({
        "week": 1,
        "topic": f"Introduction to {lesson_title}",
        "objectives": ["Understand course overview and expectations", "Explore fundamental concepts"],
        "activities": ["Course orientation", "Basic concepts introduction", "Initial assessment"],
        "assessment": "Diagnostic assessment"
    })
    
    # Weeks 2-4: Core Concepts
    for week in range(2, 5):
        weekly_plan.append({
            "week": week,
            "topic": f"Core Concepts in {lesson_title} - Part {week-1}",
            "objectives": [f"Master fundamental principles {week-1}", "Apply basic techniques"],
            "activities": ["Theoretical lessons", "Practical exercises", "Group discussions"],
            "assessment": "Formative assessment" if week < 4 else "Quiz 1"
        })
    
    # Weeks 5-8: Intermediate Topics
    for week in range(5, 9):
        weekly_plan.append({
            "week": week,
            "topic": f"Intermediate {lesson_title} Applications",
            "objectives": ["Develop intermediate skills", "Solve complex problems"],
            "activities": ["Advanced exercises", "Case studies", "Project work"],
            "assessment": "Project milestone" if week == 6 else "Practical assessment" if week == 8 else "Continuous assessment"
        })
    
    # Weeks 9-11: Advanced Topics and Projects
    for week in range(9, min(12, total_weeks)):
        weekly_plan.append({
            "week": week,
            "topic": f"Advanced {lesson_title} and Industry Applications",
            "objectives": ["Master advanced concepts", "Demonstrate expertise"],
            "activities": ["Advanced projects", "Industry case studies", "Presentations"],
            "assessment": "Major project" if week == 10 else "Continuous assessment"
        })
    
    # Final week: Review and Assessment
    if total_weeks >= 12:
        weekly_plan.append({
            "week": total_weeks,
            "topic": "Course Review and Final Assessment",
            "objectives": ["Consolidate learning", "Demonstrate competency"],
            "activities": ["Comprehensive review", "Final presentations", "Course evaluation"],
            "assessment": "Final examination"
        })
    
    return weekly_plan[:total_weeks]

def generate_assessment_schedule(total_weeks: int):
    """Generate comprehensive assessment schedule"""
    
    schedule = [
        {"week": 1, "type": "Diagnostic", "weight": "0%", "description": "Initial skills assessment"},
        {"week": 4, "type": "Quiz 1", "weight": "10%", "description": "Core concepts evaluation"},
        {"week": 6, "type": "Project Milestone", "weight": "15%", "description": "Mid-term project progress"},
        {"week": 8, "type": "Practical Assessment", "weight": "20%", "description": "Hands-on skills demonstration"},
        {"week": 10, "type": "Major Project", "weight": "25%", "description": "Comprehensive project submission"},
        {"week": total_weeks, "type": "Final Examination", "weight": "30%", "description": "Comprehensive final assessment"}
    ]
    
    return [item for item in schedule if item["week"] <= total_weeks]

def generate_comprehensive_resources(lesson_title: str, department: str, level: str):
    """Generate comprehensive resource list"""
    
    resources = {
        "textbooks": [
            f"Primary textbook: {lesson_title} - A Comprehensive Guide",
            f"Reference book: Advanced {lesson_title} Techniques",
            "RTB approved curriculum materials"
        ],
        "digital_resources": [
            "Online learning platform access",
            "Video tutorials and demonstrations",
            "Interactive simulations and tools",
            "Digital library resources"
        ],
        "equipment": [
            "Computers/laptops for each student",
            "Projector and presentation equipment",
            "Internet connectivity",
            "Specialized software and tools"
        ],
        "materials": [
            "Printed handouts and worksheets",
            "Assessment rubrics and guidelines",
            "Project templates and examples",
            "Reference cards and quick guides"
        ]
    }
    
    if department == "Software Development":
        resources["equipment"].extend([
            "Development environments (IDEs)",
            "Version control systems",
            "Testing frameworks and tools"
        ])
    elif department == "Computer System and Architecture":
        resources["equipment"].extend([
            "Hardware components for assembly",
            "Diagnostic and testing tools",
            "Network equipment and cables"
        ])
    
    return resources

def generate_evaluation_criteria(department: str, level: str):
    """Generate evaluation criteria for the course"""
    
    criteria = {
        "knowledge_understanding": {
            "weight": "25%",
            "description": "Demonstration of theoretical knowledge and conceptual understanding",
            "indicators": [
                "Accurate recall of key concepts",
                "Clear explanation of principles",
                "Understanding of relationships between concepts"
            ]
        },
        "practical_skills": {
            "weight": "35%",
            "description": "Application of skills in practical situations",
            "indicators": [
                "Competent use of tools and techniques",
                "Effective problem-solving approach",
                "Quality of practical outputs"
            ]
        },
        "analysis_evaluation": {
            "weight": "20%",
            "description": "Critical thinking and analytical abilities",
            "indicators": [
                "Ability to analyze complex problems",
                "Evaluation of different solutions",
                "Justification of decisions and approaches"
            ]
        },
        "communication_collaboration": {
            "weight": "20%",
            "description": "Professional communication and teamwork skills",
            "indicators": [
                "Clear and effective communication",
                "Collaborative working approach",
                "Professional presentation of work"
            ]
        }
    }
    
    return criteria

# RTB-Specific Generation Functions
def generate_rtb_learning_objectives(lesson_title: str, department: str, level: str):
    """Generate RTB-compliant learning objectives"""
    
    objectives_map = {
        "Software Development": {
            "Level 3": [
                "Understand and apply fundamental programming concepts and syntax",
                "Demonstrate problem-solving techniques in software development contexts",
                "Create simple functional software applications using industry standards",
                "Apply debugging and testing methodologies to ensure code quality"
            ],
            "Level 4": [
                "Design and implement intermediate-level software solutions",
                "Apply object-oriented programming principles and design patterns",
                "Evaluate and optimize software performance and efficiency",
                "Collaborate effectively in team-based software development projects"
            ],
            "Level 5": [
                "Architect scalable and maintainable enterprise software systems",
                "Lead software development projects using agile methodologies",
                "Implement advanced security measures and quality assurance practices",
                "Research and integrate emerging technologies into software solutions"
            ]
        },
        "Computer System and Architecture": {
            "Level 3": [
                "Identify and explain computer hardware components and their functions",
                "Analyze basic system performance metrics and bottlenecks",
                "Configure and maintain computer systems according to specifications",
                "Troubleshoot common hardware and software compatibility issues"
            ],
            "Level 4": [
                "Design efficient computer system architectures for specific requirements",
                "Implement system optimization strategies and performance tuning",
                "Evaluate emerging hardware technologies and their applications",
                "Manage complex IT infrastructure and network systems"
            ],
            "Level 5": [
                "Architect enterprise-level computing solutions and data centers",
                "Lead system integration and migration projects",
                "Develop strategic IT infrastructure plans and policies",
                "Research and implement cutting-edge computing technologies"
            ]
        }
    }
    
    base_objectives = objectives_map.get(department, {}).get(level, [
        "Understand core theoretical concepts and principles",
        "Apply knowledge to solve practical problems",
        "Demonstrate competency in relevant technical skills",
        "Evaluate and reflect on learning outcomes and applications"
    ])
    
    # Customize based on lesson title
    customized = []
    for obj in base_objectives:
        if "programming" in lesson_title.lower():
            obj = obj.replace("concepts", "programming concepts")
        elif "database" in lesson_title.lower():
            obj = obj.replace("concepts", "database design concepts")
        elif "network" in lesson_title.lower():
            obj = obj.replace("concepts", "networking concepts")
        customized.append(obj)
    
    return customized

def generate_rtb_teaching_methods(department: str, level: str):
    """Generate RTB-approved teaching methods"""
    
    methods = [
        "Interactive lectures with multimedia presentations",
        "Hands-on practical exercises and laboratory work",
        "Group discussions and collaborative problem-solving",
        "Case study analysis and real-world application scenarios",
        "Individual and team-based project work",
        "Peer review and constructive feedback sessions",
        "Industry expert guest lectures and demonstrations",
        "Field visits and workplace observation"
    ]
    
    if department == "Software Development":
        methods.extend([
            "Code review sessions and pair programming",
            "Agile development methodology workshops",
            "Version control system training and practice",
            "Software testing and quality assurance exercises"
        ])
    elif department == "Computer System and Architecture":
        methods.extend([
            "Hardware assembly and disassembly laboratories",
            "System performance monitoring and analysis",
            "Network configuration and troubleshooting workshops",
            "Simulation-based learning environments"
        ])
    
    return methods[:8]  # Return top 8 methods

def generate_rtb_resources(lesson_title: str, department: str, level: str):
    """Generate RTB-standard resource requirements"""
    
    resources = [
        "Whiteboard and markers for explanations",
        "Projector and screen for presentations",
        "Individual computers/laptops for students",
        "Reliable internet connectivity",
        "Printed handouts and reference materials",
        "RTB-approved textbooks and curriculum guides"
    ]
    
    if department == "Software Development":
        resources.extend([
            "Integrated Development Environment (IDE) software",
            "Programming language documentation and references",
            "Version control systems (Git/GitHub)",
            "Software development frameworks and libraries"
        ])
    elif department == "Computer System and Architecture":
        resources.extend([
            "Computer hardware components for hands-on practice",
            "System diagnostic and testing tools",
            "Network cables and connectivity equipment",
            "Multimeters and electronic testing devices"
        ])
    
    return resources

def generate_rtb_lesson_structure(lesson_title: str, duration_minutes: int):
    """Generate RTB-standard lesson structure with timing"""
    
    # RTB standard timing allocation
    intro_time = max(10, duration_minutes // 9)  # ~10-15 minutes
    development_time = duration_minutes - intro_time - 15  # Main content
    conclusion_time = 15  # Fixed 15 minutes
    
    phases = [
        {
            "phase": "Introduction (Opening)",
            "duration": f"{intro_time} minutes",
            "teacher_activities": [
                "Greet students and conduct attendance",
                "Review previous lesson key concepts",
                "Introduce today's learning objectives clearly",
                "Motivate students with real-world examples and applications"
            ],
            "student_activities": [
                "Respond to greetings and settle into learning mode",
                "Participate in review discussion",
                "Ask clarifying questions about objectives",
                "Share prior knowledge and experiences"
            ]
        },
        {
            "phase": "Development (Main Content)",
            "duration": f"{development_time} minutes",
            "teacher_activities": [
                "Present core concepts with clear explanations",
                "Demonstrate practical applications and examples",
                "Facilitate hands-on exercises and activities",
                "Guide group discussions and collaborative work",
                "Provide individual support and feedback"
            ],
            "student_activities": [
                "Actively listen and take notes",
                "Participate in practical exercises",
                "Engage in group discussions and activities",
                "Ask questions for clarification",
                "Practice new skills and concepts"
            ]
        },
        {
            "phase": "Conclusion (Closure)",
            "duration": f"{conclusion_time} minutes",
            "teacher_activities": [
                "Summarize key learning points and concepts",
                "Address student questions and concerns",
                "Preview next lesson topics and connections",
                "Assign homework and provide clear instructions"
            ],
            "student_activities": [
                "Participate in summary discussion",
                "Ask final questions for clarification",
                "Note homework assignments and deadlines",
                "Reflect on learning achievements"
            ]
        }
    ]
    
    return phases

def generate_rtb_formative_assessment(department: str, level: str):
    """Generate RTB-standard formative assessment methods"""
    
    return [
        "Questioning techniques during lesson delivery",
        "Observation of student participation and engagement",
        "Quick polls and exit tickets",
        "Peer assessment and feedback activities",
        "Short quizzes and knowledge checks"
    ]

def generate_rtb_summative_assessment(department: str, level: str):
    """Generate RTB-standard summative assessment methods"""
    
    return [
        "Written examination covering theoretical concepts",
        "Practical demonstration of skills and competencies",
        "Project-based assessment and portfolio review",
        "Oral presentation and viva voce examination"
    ]

def generate_rtb_homework(lesson_title: str, department: str, level: str):
    """Generate RTB-appropriate homework assignment"""
    
    homework_templates = {
        "Software Development": [
            "Complete the programming exercises provided in class handouts",
            "Research and prepare a summary on advanced {topic} techniques and applications",
            "Develop a small application demonstrating the {concept} learned today",
            "Review and practice the coding examples, ensuring understanding of each step"
        ],
        "Computer System and Architecture": [
            "Research the latest developments in {topic} technology and prepare a brief report",
            "Complete the system configuration worksheet provided",
            "Analyze the provided case study and prepare solutions for next class discussion",
            "Prepare for next week's practical assessment by reviewing relevant procedures"
        ]
    }
    
    templates = homework_templates.get(department, [
        "Review today's lesson materials and prepare comprehensive notes",
        "Complete the assigned reading from RTB-approved textbooks",
        "Prepare thoughtful questions for next lesson discussion",
        "Practice the skills and concepts demonstrated in class"
    ])
    
    # Customize based on lesson title
    topic = lesson_title.lower().replace("introduction to ", "").replace("advanced ", "")
    concept = topic.split()[0] if topic.split() else "the topic covered"
    
    selected_template = templates[0]
    homework = selected_template.format(topic=topic, concept=concept)
    
    return homework

# Scheme of Work specific functions
def generate_rtb_course_overview(lesson_title: str, department: str, level: str):
    """Generate RTB-compliant course overview"""
    
    overview_templates = {
        "Software Development": f"This module provides comprehensive training in {lesson_title}, covering both theoretical foundations and practical applications aligned with RTB competency standards. Students will develop industry-relevant skills through hands-on projects, collaborative learning, and real-world problem-solving scenarios. The curriculum is designed to meet RTB Level {level[-1]} requirements and prepare students for professional software development careers in Rwanda's growing ICT sector.",
        "Computer System and Architecture": f"This module explores {lesson_title} with emphasis on practical implementation and industry best practices according to RTB standards. Students will gain deep understanding of system components, performance optimization, and troubleshooting techniques. The program combines theoretical knowledge with extensive hands-on experience to prepare graduates for technical leadership roles in Rwanda's technology infrastructure development."
    }
    
    return overview_templates.get(department, f"This comprehensive module in {lesson_title} provides students with essential knowledge and practical skills required for success in the {department} field. The curriculum is designed to meet RTB competency standards and industry requirements in Rwanda.")

def generate_rtb_prerequisites(lesson_title: str, department: str, level: str):
    """Generate course prerequisites"""
    
    if level == "Level 3":
        return ["Successful completion of secondary education", "Basic computer literacy skills"]
    elif level == "Level 4":
        return ["Successful completion of Level 3 in related field", "Demonstrated competency in foundational concepts"]
    else:
        return ["Successful completion of Level 4 in related field", "Professional experience or internship recommended"]

def generate_rtb_course_outcomes(lesson_title: str, department: str, level: str):
    """Generate comprehensive course learning outcomes"""
    
    outcomes = [
        f"Demonstrate comprehensive understanding of {lesson_title} principles and applications",
        f"Apply {lesson_title} knowledge to solve complex real-world problems",
        f"Analyze and evaluate different approaches and methodologies in {lesson_title}",
        f"Create innovative solutions using {lesson_title} best practices and standards",
        f"Collaborate effectively in professional {lesson_title} project teams",
        f"Communicate technical concepts clearly to diverse stakeholders",
        f"Demonstrate professional ethics and responsibility in {lesson_title} practice",
        f"Adapt to emerging trends and technologies in {lesson_title} field"
    ]
    
    return outcomes

def generate_rtb_weekly_breakdown(lesson_title: str, department: str, level: str, total_weeks: int):
    """Generate detailed RTB-compliant weekly breakdown"""
    
    weekly_plan = []
    
    # Week 1: Course Introduction
    weekly_plan.append({
        "week": 1,
        "topic": f"Introduction to {lesson_title}",
        "learning_outcomes": [
            "Understand course structure and expectations",
            "Identify key concepts and terminology",
            "Recognize industry applications and career opportunities"
        ],
        "content": [
            "Course overview and learning objectives",
            "Historical context and current trends",
            "Industry standards and best practices",
            "Assessment methods and requirements"
        ],
        "activities": [
            "Course orientation and introductions",
            "Diagnostic assessment",
            "Industry overview presentation",
            "Learning contract establishment"
        ],
        "assessment": "Diagnostic assessment and participation",
        "resources": "Course handbook, industry reports"
    })
    
    # Weeks 2-4: Foundational Concepts
    for week in range(2, 5):
        weekly_plan.append({
            "week": week,
            "topic": f"Fundamental Concepts in {lesson_title} - Module {week-1}",
            "learning_outcomes": [
                f"Master core principles of module {week-1}",
                "Apply basic techniques and methodologies",
                "Demonstrate understanding through practical exercises"
            ],
            "content": [
                f"Theoretical foundations of module {week-1}",
                "Key principles and concepts",
                "Practical applications and examples",
                "Industry case studies"
            ],
            "activities": [
                "Interactive lectures and discussions",
                "Hands-on practical exercises",
                "Group problem-solving sessions",
                "Individual skill practice"
            ],
            "assessment": "Formative assessment" if week < 4 else "Quiz 1 (10%)",
            "resources": "Textbook chapters, online resources, lab materials"
        })
    
    # Weeks 5-8: Intermediate Applications
    for week in range(5, 9):
        weekly_plan.append({
            "week": week,
            "topic": f"Intermediate Applications in {lesson_title}",
            "learning_outcomes": [
                "Develop intermediate-level skills and competencies",
                "Solve complex problems using advanced techniques",
                "Integrate multiple concepts in practical scenarios"
            ],
            "content": [
                "Advanced concepts and methodologies",
                "Complex problem-solving techniques",
                "Integration of multiple systems/concepts",
                "Quality assurance and best practices"
            ],
            "activities": [
                "Advanced practical exercises",
                "Case study analysis and solutions",
                "Team-based project work",
                "Peer review and feedback sessions"
            ],
            "assessment": "Project milestone" if week == 6 else "Practical assessment (20%)" if week == 8 else "Continuous assessment",
            "resources": "Advanced materials, project guidelines, industry standards"
        })
    
    # Weeks 9-11: Advanced Topics and Integration
    for week in range(9, min(12, total_weeks)):
        weekly_plan.append({
            "week": week,
            "topic": f"Advanced {lesson_title} and Professional Practice",
            "learning_outcomes": [
                "Master advanced concepts and techniques",
                "Demonstrate professional-level competency",
                "Apply knowledge in real-world contexts"
            ],
            "content": [
                "Advanced professional techniques",
                "Industry standards and regulations",
                "Emerging trends and technologies",
                "Professional ethics and responsibility"
            ],
            "activities": [
                "Capstone project development",
                "Industry case study presentations",
                "Professional skills workshops",
                "Portfolio development"
            ],
            "assessment": "Major project (30%)" if week == 10 else "Continuous assessment",
            "resources": "Professional standards, industry publications, project resources"
        })
    
    # Final week: Review and Assessment
    if total_weeks >= 12:
        weekly_plan.append({
            "week": total_weeks,
            "topic": "Course Review and Final Assessment",
            "learning_outcomes": [
                "Consolidate and integrate all learning",
                "Demonstrate comprehensive competency",
                "Reflect on learning journey and achievements"
            ],
            "content": [
                "Comprehensive course review",
                "Integration of all modules",
                "Professional portfolio completion",
                "Career planning and next steps"
            ],
            "activities": [
                "Final project presentations",
                "Comprehensive review sessions",
                "Peer evaluation and feedback",
                "Course evaluation and reflection"
            ],
            "assessment": "Final examination (40%)",
            "resources": "All course materials, portfolio guidelines"
        })
    
    return weekly_plan[:total_weeks]

def generate_rtb_assessment_schedule(total_weeks: int):
    """Generate RTB-compliant assessment schedule"""
    
    schedule = [
        {
            "week": 1,
            "assessment_type": "Diagnostic Assessment",
            "weight": "0%",
            "description": "Initial skills and knowledge assessment",
            "method": "Written test and practical demonstration"
        },
        {
            "week": 4,
            "assessment_type": "Quiz 1",
            "weight": "10%",
            "description": "Foundational concepts evaluation",
            "method": "Written examination"
        },
        {
            "week": 6,
            "assessment_type": "Project Milestone",
            "weight": "15%",
            "description": "Mid-term project progress assessment",
            "method": "Project submission and presentation"
        },
        {
            "week": 8,
            "assessment_type": "Practical Assessment",
            "weight": "20%",
            "description": "Hands-on skills demonstration",
            "method": "Practical examination and portfolio"
        },
        {
            "week": 10,
            "assessment_type": "Major Project",
            "weight": "25%",
            "description": "Comprehensive project submission",
            "method": "Project report, presentation, and demonstration"
        },
        {
            "week": total_weeks,
            "assessment_type": "Final Examination",
            "weight": "30%",
            "description": "Comprehensive final assessment",
            "method": "Written and practical examination"
        }
    ]
    
    return [item for item in schedule if item["week"] <= total_weeks]

def generate_rtb_textbooks(lesson_title: str, department: str, level: str):
    """Generate RTB-approved textbook list"""
    
    return [
        f"RTB Official Curriculum Guide for {lesson_title}",
        f"{lesson_title}: A Comprehensive Textbook (RTB Approved)",
        f"Practical Guide to {lesson_title} - {level} Edition",
        "Industry Standards and Best Practices Manual"
    ]

def generate_rtb_equipment(lesson_title: str, department: str, level: str):
    """Generate required equipment list"""
    
    equipment = [
        "Individual computers/laptops for each student",
        "Projector and presentation screen",
        "Whiteboard and markers",
        "Internet connectivity and network access"
    ]
    
    if department == "Software Development":
        equipment.extend([
            "Software development tools and IDEs",
            "Version control systems",
            "Testing and debugging tools",
            "Database management systems"
        ])
    elif department == "Computer System and Architecture":
        equipment.extend([
            "Computer hardware components",
            "Diagnostic and testing equipment",
            "Network cables and connectors",
            "Multimeters and electronic tools"
        ])
    
    return equipment

def generate_rtb_materials(lesson_title: str, department: str, level: str):
    """Generate required materials list"""
    
    return [
        "Printed handouts and worksheets",
        "Assessment rubrics and guidelines",
        "Project templates and examples",
        "Reference cards and quick guides",
        "Safety guidelines and procedures"
    ]

def generate_rtb_digital_resources(lesson_title: str, department: str, level: str):
    """Generate digital resources list"""
    
    return [
        "RTB Learning Management System access",
        "Online video tutorials and demonstrations",
        "Interactive simulations and virtual labs",
        "Digital library and research databases",
        "Industry-standard software licenses",
        "Online collaboration and communication tools"
    ]

# Lesson Planning API Endpoints

@app.post("/ai/generate-session-plan")
def generate_session_plan_endpoint(request: AIGenerateRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    try:
        start_time = datetime.utcnow()
        
        # Generate session plan using AI
        session_plan_data = generate_session_plan_ai(
            lesson_title=request.lesson_title,
            department=request.department,
            level=request.level,
            duration_minutes=request.duration_minutes or 90,
            additional_requirements=request.additional_requirements or ""
        )
        
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        
        # Log the generation
        log_entry = AIGenerationLog(
            user_id=current_user.id,
            generation_type="session_plan",
            input_parameters=request.dict(),
            generated_content=json.dumps(session_plan_data),
            processing_time=processing_time,
            success=True
        )
        db.add(log_entry)
        db.commit()
        
        return {
            "success": True,
            "session_plan": session_plan_data,
            "processing_time": processing_time,
            "message": "Session plan generated successfully using RTB standards"
        }
        
    except Exception as e:
        # Log the error
        log_entry = AIGenerationLog(
            user_id=current_user.id,
            generation_type="session_plan",
            input_parameters=request.dict(),
            processing_time=0,
            success=False,
            error_message=str(e)
        )
        db.add(log_entry)
        db.commit()
        
        raise HTTPException(status_code=500, detail=f"Failed to generate session plan: {str(e)}")

@app.post("/ai/generate-scheme-of-work")
def generate_scheme_of_work_endpoint(request: AIGenerateRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    try:
        start_time = datetime.utcnow()
        
        # Generate scheme of work using AI
        scheme_data = generate_scheme_of_work_ai(
            lesson_title=request.lesson_title,
            department=request.department,
            level=request.level,
            academic_year=request.academic_year or "2024-2025",
            term=request.term or "Term 1",
            total_weeks=request.total_weeks or 12,
            additional_requirements=request.additional_requirements or ""
        )
        
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        
        # Log the generation
        log_entry = AIGenerationLog(
            user_id=current_user.id,
            generation_type="scheme_of_work",
            input_parameters=request.dict(),
            generated_content=json.dumps(scheme_data),
            processing_time=processing_time,
            success=True
        )
        db.add(log_entry)
        db.commit()
        
        return {
            "success": True,
            "scheme_of_work": scheme_data,
            "processing_time": processing_time,
            "message": "Scheme of work generated successfully using RTB standards"
        }
        
    except Exception as e:
        # Log the error
        log_entry = AIGenerationLog(
            user_id=current_user.id,
            generation_type="scheme_of_work",
            input_parameters=request.dict(),
            processing_time=0,
            success=False,
            error_message=str(e)
        )
        db.add(log_entry)
        db.commit()
        
        raise HTTPException(status_code=500, detail=f"Failed to generate scheme of work: {str(e)}")

@app.post("/session-plans")
def create_session_plan(plan: SessionPlanCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Validate lesson exists
    lesson = db.query(Lesson).filter(Lesson.id == plan.lesson_id).first()
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    # Create session plan
    db_plan = SessionPlan(
        title=plan.title,
        lesson_id=plan.lesson_id,
        department=plan.department,
        level=plan.level,
        duration_minutes=plan.duration_minutes,
        learning_objectives=plan.learning_objectives,
        teaching_methods=plan.teaching_methods,
        resources_required=plan.resources_required,
        assessment_methods=plan.assessment_methods,
        content_outline=plan.content_outline,
        activities=plan.activities,
        homework_assignment=plan.homework_assignment,
        created_by=current_user.id
    )
    
    db.add(db_plan)
    db.commit()
    db.refresh(db_plan)
    
    return db_plan

@app.post("/schemes-of-work")
def create_scheme_of_work(scheme: SchemeOfWorkCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    # Validate lesson exists
    lesson = db.query(Lesson).filter(Lesson.id == scheme.lesson_id).first()
    if not lesson:
        raise HTTPException(status_code=404, detail="Lesson not found")
    
    # Create scheme of work
    db_scheme = SchemeOfWork(
        title=scheme.title,
        lesson_id=scheme.lesson_id,
        department=scheme.department,
        level=scheme.level,
        academic_year=scheme.academic_year,
        term=scheme.term,
        total_weeks=scheme.total_weeks,
        learning_outcomes=scheme.learning_outcomes,
        assessment_schedule=scheme.assessment_schedule,
        resources_list=scheme.resources_list,
        created_by=current_user.id
    )
    
    db.add(db_scheme)
    db.commit()
    db.refresh(db_scheme)
    
    return db_scheme

@app.get("/session-plans")
def get_session_plans(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    if current_user.role == "admin":
        return db.query(SessionPlan).order_by(SessionPlan.created_at.desc()).all()
    else:
        return db.query(SessionPlan).filter(SessionPlan.created_by == current_user.id).order_by(SessionPlan.created_at.desc()).all()

@app.get("/schemes-of-work")
def get_schemes_of_work(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    if current_user.role == "admin":
        return db.query(SchemeOfWork).order_by(SchemeOfWork.created_at.desc()).all()
    else:
        return db.query(SchemeOfWork).filter(SchemeOfWork.created_by == current_user.id).order_by(SchemeOfWork.created_at.desc()).all()

@app.get("/session-plans/{plan_id}/export/pdf")
def export_session_plan_pdf(plan_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    plan = db.query(SessionPlan).filter(SessionPlan.id == plan_id).first()
    if not plan:
        raise HTTPException(status_code=404, detail="Session plan not found")
    
    # Check ownership for teachers
    if current_user.role == "teacher" and plan.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="You can only export your own session plans")
    
    try:
        # Create professional RTB-compliant PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # RTB Header
        header_style = ParagraphStyle('RTBHeader', parent=styles['Heading1'], fontSize=16, spaceAfter=20, alignment=1, textColor=colors.HexColor('#1f2937'))
        story.append(Paragraph("RWANDA TRAINING BOARD (RTB)", header_style))
        story.append(Paragraph("SESSION PLAN", styles['Heading2']))
        story.append(Spacer(1, 20))
        
        # Plan details table
        plan_details = [
            ['Institution:', 'RTB Affiliated Institution'],
            ['Department:', plan.department],
            ['Level:', plan.level],
            ['Lesson Title:', plan.title],
            ['Duration:', f"{plan.duration_minutes} minutes"],
            ['Date:', plan.created_at.strftime('%Y-%m-%d')],
            ['Instructor:', current_user.full_name or current_user.username]
        ]
        
        details_table = Table(plan_details, colWidths=[2*inch, 4*inch])
        details_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        story.append(details_table)
        story.append(Spacer(1, 20))
        
        # Learning Objectives
        story.append(Paragraph("LEARNING OBJECTIVES", styles['Heading3']))
        for i, obj in enumerate(plan.learning_objectives or [], 1):
            story.append(Paragraph(f"{i}. {obj}", styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Teaching Methods
        story.append(Paragraph("TEACHING METHODS", styles['Heading3']))
        for method in plan.teaching_methods or []:
            story.append(Paragraph(f"• {method}", styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Resources Required
        story.append(Paragraph("RESOURCES REQUIRED", styles['Heading3']))
        for resource in plan.resources_required or []:
            story.append(Paragraph(f"• {resource}", styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Content Outline
        if plan.content_outline:
            story.append(Paragraph("CONTENT OUTLINE", styles['Heading3']))
            story.append(Paragraph(plan.content_outline, styles['Normal']))
            story.append(Spacer(1, 15))
        
        # Assessment Methods
        story.append(Paragraph("ASSESSMENT METHODS", styles['Heading3']))
        for method in plan.assessment_methods or []:
            story.append(Paragraph(f"• {method}", styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Homework Assignment
        if plan.homework_assignment:
            story.append(Paragraph("HOMEWORK ASSIGNMENT", styles['Heading3']))
            story.append(Paragraph(plan.homework_assignment, styles['Normal']))
            story.append(Spacer(1, 15))
        
        # Reflection Section
        story.append(Paragraph("TEACHER REFLECTION", styles['Heading3']))
        reflection_table = [
            ['What went well:', ''],
            ['Areas for improvement:', ''],
            ['Student engagement level:', ''],
            ['Objectives achieved:', ''],
            ['Next lesson adjustments:', '']
        ]
        
        reflection_tbl = Table(reflection_table, colWidths=[2*inch, 4*inch], rowHeights=[0.8*inch]*5)
        reflection_tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        story.append(reflection_tbl)
        
        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, alignment=1, textColor=colors.grey)
        story.append(Paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} | RTB Session Plan Template", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Session_Plan_{plan.id}_{plan.title.replace(' ', '_')}.pdf"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

@app.get("/schemes-of-work/{scheme_id}/export/pdf")
def export_scheme_of_work_pdf(scheme_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if current_user.role not in ["admin", "teacher"]:
        raise HTTPException(status_code=403, detail="Teacher or Admin access required")
    
    scheme = db.query(SchemeOfWork).filter(SchemeOfWork.id == scheme_id).first()
    if not scheme:
        raise HTTPException(status_code=404, detail="Scheme of work not found")
    
    # Check ownership for teachers
    if current_user.role == "teacher" and scheme.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="You can only export your own schemes of work")
    
    try:
        # Create professional RTB-compliant PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # RTB Header
        header_style = ParagraphStyle('RTBHeader', parent=styles['Heading1'], fontSize=16, spaceAfter=20, alignment=1, textColor=colors.HexColor('#1f2937'))
        story.append(Paragraph("RWANDA TRAINING BOARD (RTB)", header_style))
        story.append(Paragraph("SCHEME OF WORK", styles['Heading2']))
        story.append(Spacer(1, 20))
        
        # Scheme details table
        scheme_details = [
            ['Institution:', 'RTB Affiliated Institution'],
            ['Department:', scheme.department],
            ['Level:', scheme.level],
            ['Subject:', scheme.title],
            ['Academic Year:', scheme.academic_year],
            ['Term:', scheme.term],
            ['Total Weeks:', str(scheme.total_weeks)],
            ['Instructor:', current_user.full_name or current_user.username]
        ]
        
        details_table = Table(scheme_details, colWidths=[2*inch, 4*inch])
        details_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        story.append(details_table)
        story.append(Spacer(1, 20))
        
        # Learning Outcomes
        story.append(Paragraph("LEARNING OUTCOMES", styles['Heading3']))
        for i, outcome in enumerate(scheme.learning_outcomes or [], 1):
            story.append(Paragraph(f"{i}. {outcome}", styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Weekly Breakdown
        if scheme.weekly_breakdown:
            story.append(Paragraph("WEEKLY BREAKDOWN", styles['Heading3']))
            
            # Create weekly breakdown table
            weekly_data = [['Week', 'Topic', 'Objectives', 'Assessment']]
            for week_info in scheme.weekly_breakdown:
                objectives_text = '; '.join(week_info.get('objectives', []))
                weekly_data.append([
                    str(week_info.get('week', '')),
                    week_info.get('topic', '')[:40] + '...' if len(week_info.get('topic', '')) > 40 else week_info.get('topic', ''),
                    objectives_text[:50] + '...' if len(objectives_text) > 50 else objectives_text,
                    week_info.get('assessment', '')
                ])
            
            weekly_table = Table(weekly_data, colWidths=[0.8*inch, 2.5*inch, 2.5*inch, 1.2*inch])
            weekly_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(weekly_table)
            story.append(Spacer(1, 15))
        
        # Assessment Schedule
        if scheme.assessment_schedule:
            story.append(Paragraph("ASSESSMENT SCHEDULE", styles['Heading3']))
            
            assessment_data = [['Week', 'Assessment Type', 'Weight', 'Description']]
            for assessment in scheme.assessment_schedule:
                assessment_data.append([
                    str(assessment.get('week', '')),
                    assessment.get('type', ''),
                    assessment.get('weight', ''),
                    assessment.get('description', '')[:40] + '...' if len(assessment.get('description', '')) > 40 else assessment.get('description', '')
                ])
            
            assessment_table = Table(assessment_data, colWidths=[0.8*inch, 1.5*inch, 1*inch, 3.7*inch])
            assessment_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(assessment_table)
            story.append(Spacer(1, 15))
        
        # Resources List
        if scheme.resources_list:
            story.append(Paragraph("REQUIRED RESOURCES", styles['Heading3']))
            for resource in scheme.resources_list:
                story.append(Paragraph(f"• {resource}", styles['Normal']))
        
        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, alignment=1, textColor=colors.grey)
        story.append(Paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} | RTB Scheme of Work Template", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Scheme_of_Work_{scheme.id}_{scheme.title.replace(' ', '_')}.pdf"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

# Create tables on startup
Base.metadata.create_all(bind=engine)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

@app.route('/rtb')
def rtb_interface():
    return send_from_directory('../frontend/static', 'rtb.html')
