#!/usr/bin/env python3
"""
ULTRA-PRECISE DOCX Parser - Extract EVERYTHING correctly
"""
from docx import Document
import re
import json

doc = Document('/app/schools.docx')

print("="*100)
print("ULTRA-PRECISE DOCX PARSING")
print("="*100)

all_schools = []
districts_found = set()

for table_idx, table in enumerate(doc.tables):
    print(f"\nProcessing Table {table_idx + 1}/{len(doc.tables)}...")
    
    for row_idx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        
        # Skip if not enough cells
        if len(cells) < 7:
            continue
        
        # Extract fields
        sn = cells[0].strip()
        district = cells[1].strip().upper()
        school_name = cells[2].strip()
        school_type = cells[3].strip()
        status = cells[4].strip() if len(cells) > 4 else ""
        boarding = cells[5].strip() if len(cells) > 5 else ""
        
        # Find accredited trades column (usually 6 or 7)
        accredited_col = 6
        non_accredited_col = 7
        
        # Adjust if table has extra columns
        if len(cells) > 9:
            accredited_col = 7
            non_accredited_col = 8
        
        accredited_trades = cells[accredited_col] if len(cells) > accredited_col else ""
        non_accredited_trades = cells[non_accredited_col] if len(cells) > non_accredited_col else ""
        
        # Skip header rows
        if school_name in ['School Name', ''] or district in ['DISTRICT', '']:
            continue
        
        # Skip if SN is not a number
        if not sn.isdigit():
            continue
        
        # Clean school type
        if school_type.upper() not in ['TSS', 'VTC']:
            if 'VTC' in school_name.upper():
                school_type = 'VTC'
            elif 'TSS' in school_name.upper():
                school_type = 'TSS'
            else:
                school_type = 'TSS'  # Default
        
        # Parse trades
        all_trades_text = accredited_trades + "\n" + non_accredited_trades
        trades = []
        
        for line in all_trades_text.split('\n'):
            line = line.strip()
            if not line or line.lower() in ['none', 'n/a', '']:
                continue
            
            # Split by common separators
            trade_parts = re.split(r'[|]', line)
            for trade_part in trade_parts:
                trade_part = trade_part.strip()
                
                # Skip non-trade text
                if len(trade_part) < 3:
                    continue
                if trade_part.lower() in ['none', 'n/a', 'government', 'aided', 'boarding', 'day', 'public', 'private']:
                    continue
                
                # Extract level
                level_match = re.search(r'[Ll](\d+[-\d]*)', trade_part)
                if level_match:
                    level = f"L{level_match.group(1)}"
                else:
                    level = "L3-5"  # Default
                
                # Clean trade name
                trade_clean = re.sub(r'^\d+\.\s*', '', trade_part)  # Remove numbering
                trade_clean = re.sub(r'\s*[Ll]\d+[-\d]*\s*$', '', trade_clean)  # Remove level suffix
                trade_clean = re.sub(r'\s+', ' ', trade_clean).strip()
                
                if trade_clean and len(trade_clean) > 2:
                    trades.append({
                        'name': trade_clean,
                        'level': level
                    })
        
        # Store school
        if school_name and district and trades:
            all_schools.append({
                'sn': sn,
                'district': district,
                'school_name': school_name,
                'school_type': school_type,
                'status': status,
                'boarding': boarding,
                'trades': trades
            })
            districts_found.add(district)

print(f"\n{'='*100}")
print(f"PARSING COMPLETE")
print(f"{'='*100}")
print(f"Total schools found: {len(all_schools)}")
print(f"Total districts found: {len(districts_found)}")

print(f"\nDistricts found:")
for d in sorted(districts_found):
    count = len([s for s in all_schools if s['district'] == d])
    print(f"  - {d}: {count} schools")

# Save to JSON
output = {
    'schools': all_schools,
    'districts': sorted(list(districts_found))
}

with open('/app/schools_precise.json', 'w') as f:
    json.dump(output, f, indent=2)

print(f"\nData saved to /app/schools_precise.json")

# Check specific cases
print(f"\n{'='*100}")
print("CHECKING SPECIFIC CASES")
print(f"{'='*100}")

# Check for Muhanga
muhanga_schools = [s for s in all_schools if 'MUHANGA' in s['district']]
print(f"\nMUHANGA district: {len(muhanga_schools)} schools")
if muhanga_schools:
    for s in muhanga_schools:
        print(f"  - {s['school_name']}")

# Check for schools with "None" trades
none_trades = [s for s in all_schools if any('none' in t['name'].lower() for t in s['trades'])]
print(f"\nSchools with 'None' in trades: {len(none_trades)}")
if none_trades:
    for s in none_trades[:5]:
        print(f"  - {s['school_name']}: {[t['name'] for t in s['trades']]}")

# Sample schools
print(f"\nSample schools:")
for s in all_schools[:3]:
    print(f"\n  {s['school_name']} ({s['district']})")
    print(f"    Type: {s['school_type']}")
    print(f"    Trades: {len(s['trades'])}")
    for t in s['trades'][:3]:
        print(f"      - {t['name']} ({t['level']})")
