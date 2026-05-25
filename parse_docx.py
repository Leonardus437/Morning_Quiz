#!/usr/bin/env python3
"""
Simplified repopulation - parse DOCX and generate SQL
"""
from docx import Document
import re
import json

# Province-District mapping
PROVINCE_DISTRICTS = {
    "Eastern Province": ["Bugesera", "Gatsibo", "Kayonza", "Kirehe", "Ngoma", "Nyagatare", "Rwamagana"],
    "Kigali City": ["Gasabo", "Kicukiro", "Nyarugenge"],
    "Northern Province": ["Burera", "Gakenke", "Gicumbi", "Musanze", "Rulindo"],
    "Southern Province": ["Gisagara", "Huye", "Kamonyi", "Muhanga", "Nyamagabe", "Nyanza", "Nyaruguru", "Ruhango"],
    "Western Province": ["Karongi", "Ngororero", "Nyabihu", "Nyamasheke", "Rubavu", "Rusizi", "Rutsiro"]
}

def get_province_for_district(district_name):
    district_upper = district_name.upper().strip()
    for province, districts in PROVINCE_DISTRICTS.items():
        for dist in districts:
            if dist.upper() == district_upper:
                return province
    return None

def clean_school_name(name):
    name = name.strip()
    name = re.sub(r'\s+', ' ', name)
    return name

def clean_trade_name(trade):
    trade = trade.strip()
    trade = re.sub(r'^\d+\.\s*', '', trade)
    trade = re.sub(r'\s*\([Ll]\d+[-\d]*\)\s*$', '', trade)
    trade = re.sub(r'\s+', ' ', trade).strip()
    return trade

def extract_level(trade_text):
    match = re.search(r'\(([Ll]\d+[-\d]*)\)', trade_text)
    if match:
        return match.group(1).upper()
    return "L3-5"

def generate_trade_code(trade_name):
    words = trade_name.upper().split()
    if len(words) == 1:
        code = words[0][:10]
    else:
        code = '_'.join([w[:8] for w in words[:2]])
    return code[:20]

def assign_category(trade_name):
    trade_lower = trade_name.lower()
    
    if any(x in trade_lower for x in ['computer', 'software', 'network', 'ict', 'information', 'cyber', 'data']):
        return 'ICT & Computing'
    elif any(x in trade_lower for x in ['building', 'construction', 'masonry', 'plumbing', 'carpentry workshop']):
        return 'Construction'
    elif any(x in trade_lower for x in ['electrical', 'electronic', 'electricity', 'telecommunication']):
        return 'Electrical & Electronics'
    elif any(x in trade_lower for x in ['hotel', 'tourism', 'culinary', 'food and beverage', 'hospitality']):
        return 'Hospitality & Tourism'
    elif any(x in trade_lower for x in ['fashion', 'tailoring', 'textile', 'garment']):
        return 'Fashion & Textiles'
    elif any(x in trade_lower for x in ['hairdressing', 'beauty', 'cosmetology', 'salon']):
        return 'Beauty & Wellness'
    elif any(x in trade_lower for x in ['welding', 'metal', 'fabrication']):
        return 'Metal Work'
    elif any(x in trade_lower for x in ['wood', 'carpentry', 'joinery', 'furniture']) and 'workshop' not in trade_lower:
        return 'Wood Work'
    elif any(x in trade_lower for x in ['agriculture', 'farming', 'agri']):
        return 'Agriculture'
    elif any(x in trade_lower for x in ['automobile', 'automotive', 'motor']):
        return 'Automotive'
    elif any(x in trade_lower for x in ['manufacturing']):
        return 'Manufacturing'
    elif any(x in trade_lower for x in ['survey']):
        return 'Surveying'
    elif any(x in trade_lower for x in ['food processing']):
        return 'Food Processing'
    else:
        return 'General'

print("Parsing DOCX file...")
doc = Document('/app/LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC_2024-2025_lJhq.docx')

schools_data = {}

for table_idx, table in enumerate(doc.tables):
    print(f"Processing table {table_idx + 1}/{len(doc.tables)}...")
    
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            continue
            
        cells = [cell.text.strip() for cell in row.cells]
        
        if len(cells) < 7:
            continue
        
        district = cells[1].strip()
        school_name = cells[2].strip()
        school_type = cells[3].strip() if len(cells) > 3 else "TSS"
        
        if not school_name or school_name in ['School Name', '']:
            continue
        if not district or district in ['District', '']:
            continue
            
        school_name = clean_school_name(school_name)
        
        if school_type.upper() not in ['TSS', 'VTC']:
            if 'VTC' in school_name.upper():
                school_type = 'VTC'
            else:
                school_type = 'TSS'
        
        province = get_province_for_district(district)
        if not province:
            continue
        
        accredited_trades = cells[6] if len(cells) > 6 else ""
        non_accredited_trades = cells[7] if len(cells) > 7 else ""
        
        all_trades_text = accredited_trades + "\n" + non_accredited_trades
        
        trades = []
        for line in all_trades_text.split('\n'):
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            trade_parts = re.split(r'[|,]', line)
            for trade_part in trade_parts:
                trade_part = trade_part.strip()
                if trade_part and len(trade_part) > 2:
                    level = extract_level(trade_part)
                    trade_clean = clean_trade_name(trade_part)
                    if trade_clean:
                        trades.append({'name': trade_clean, 'level': level})
        
        if school_name not in schools_data:
            schools_data[school_name] = {
                'district': district,
                'province': province,
                'type': school_type,
                'trades': trades
            }
        else:
            existing_trade_names = {t['name'] for t in schools_data[school_name]['trades']}
            for trade in trades:
                if trade['name'] not in existing_trade_names:
                    schools_data[school_name]['trades'].append(trade)

print(f"\nFound {len(schools_data)} unique schools")

# Output data as JSON for processing
output = {
    'provinces': PROVINCE_DISTRICTS,
    'schools': schools_data
}

with open('/app/schools_parsed.json', 'w') as f:
    json.dump(output, f, indent=2)

print("Data saved to /app/schools_parsed.json")

# Print some statistics
print(f"\nSchools by province:")
province_counts = {}
for school_name, data in schools_data.items():
    prov = data['province']
    province_counts[prov] = province_counts.get(prov, 0) + 1

for prov, count in sorted(province_counts.items()):
    print(f"  {prov}: {count} schools")

# Check for specific schools
print("\nChecking specific schools:")
test_schools = ['APEK', 'RRUKOMA', 'IGNACE', 'NYAMATA', 'JANJA']
for test in test_schools:
    found = [name for name in schools_data.keys() if test.upper() in name.upper()]
    if found:
        print(f"  ✓ {test}: {', '.join(found)}")
    else:
        print(f"  ✗ {test}: NOT FOUND")
