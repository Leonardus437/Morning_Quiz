#!/usr/bin/env python3
"""
Complete repopulation script - carefully parse DOCX and avoid duplicates
"""
import psycopg2
from docx import Document
import re
import json

# Database connection
conn = psycopg2.connect(
    dbname="morning_quiz",
    user="quiz_user",
    password="quiz_pass123",
    host="db",
    port="5432"
)
cur = conn.cursor()

# Province-District mapping for Rwanda
PROVINCE_DISTRICTS = {
    "Eastern Province": ["Bugesera", "Gatsibo", "Kayonza", "Kirehe", "Ngoma", "Nyagatare", "Rwamagana"],
    "Kigali City": ["Gasabo", "Kicukiro", "Nyarugenge"],
    "Northern Province": ["Burera", "Gakenke", "Gicumbi", "Musanze", "Rulindo"],
    "Southern Province": ["Gisagara", "Huye", "Kamonyi", "Muhanga", "Nyamagabe", "Nyanza", "Nyaruguru", "Ruhango"],
    "Western Province": ["Karongi", "Ngororero", "Nyabihu", "Nyamasheke", "Rubavu", "Rusizi", "Rutsiro"]
}

def get_province_for_district(district_name):
    """Find province for a given district"""
    district_upper = district_name.upper().strip()
    for province, districts in PROVINCE_DISTRICTS.items():
        for dist in districts:
            if dist.upper() == district_upper:
                return province
    return None

def clean_school_name(name):
    """Clean school name"""
    name = name.strip()
    # Remove extra spaces
    name = re.sub(r'\s+', ' ', name)
    return name

def clean_trade_name(trade):
    """Clean trade name - remove numbering and level suffixes"""
    trade = trade.strip()
    # Remove numbering like "1. ", "2. "
    trade = re.sub(r'^\d+\.\s*', '', trade)
    # Remove level suffixes like (L3-5), (L1), etc.
    trade = re.sub(r'\s*\([Ll]\d+[-\d]*\)\s*$', '', trade)
    trade = re.sub(r'\s+', ' ', trade).strip()
    return trade

def extract_level(trade_text):
    """Extract level from trade text"""
    match = re.search(r'\(([Ll]\d+[-\d]*)\)', trade_text)
    if match:
        return match.group(1).upper()
    return "L3-5"  # Default

def generate_trade_code(trade_name):
    """Generate unique trade code"""
    words = trade_name.upper().split()
    if len(words) == 1:
        code = words[0][:10]
    else:
        code = '_'.join([w[:8] for w in words[:2]])
    return code[:20]

def assign_category(trade_name):
    """Assign category based on trade name"""
    trade_lower = trade_name.lower()
    
    if any(x in trade_lower for x in ['computer', 'software', 'network', 'ict', 'information', 'cyber', 'data']):
        return 'ICT & Computing'
    elif any(x in trade_lower for x in ['building', 'construction', 'masonry', 'plumbing', 'carpentry workshop']):
        return 'Construction'
    elif any(x in trade_lower for x in ['electrical', 'electronic', 'electricity', 'telecommunication']):
        return 'Electrical & Electronics'
    elif any(x in trade_lower for x in ['hotel', 'tourism', 'culinary', 'food and beverage', 'hospitality']):
        return 'Hospitality & Tourism'
    elif any(x in trade_lower for x in ['fashion', 'tailoring', 'textile', 'garment']):
        return 'Fashion & Textiles'
    elif any(x in trade_lower for x in ['hairdressing', 'beauty', 'cosmetology', 'salon']):
        return 'Beauty & Wellness'
    elif any(x in trade_lower for x in ['welding', 'metal', 'fabrication']):
        return 'Metal Work'
    elif any(x in trade_lower for x in ['wood', 'carpentry', 'joinery', 'furniture']) and 'workshop' not in trade_lower:
        return 'Wood Work'
    elif any(x in trade_lower for x in ['agriculture', 'farming', 'agri']):
        return 'Agriculture'
    elif any(x in trade_lower for x in ['automobile', 'automotive', 'motor']):
        return 'Automotive'
    elif any(x in trade_lower for x in ['manufacturing']):
        return 'Manufacturing'
    elif any(x in trade_lower for x in ['survey']):
        return 'Surveying'
    elif any(x in trade_lower for x in ['food processing']):
        return 'Food Processing'
    else:
        return 'General'

print("Starting complete repopulation...")

# Step 1: Clear existing data
print("\n1. Clearing existing school data...")
cur.execute("TRUNCATE school_trades, schools, trades, districts, provinces RESTART IDENTITY CASCADE;")
conn.commit()
print("   ✓ Data cleared")

# Step 2: Create provinces
print("\n2. Creating provinces...")
province_ids = {}
for province_name in PROVINCE_DISTRICTS.keys():
    code = province_name.split()[0][:3].upper()
    cur.execute(
        "INSERT INTO provinces (name, code) VALUES (%s, %s) RETURNING id",
        (province_name, code)
    )
    province_ids[province_name] = cur.fetchone()[0]
conn.commit()
print(f"   ✓ Created {len(province_ids)} provinces")

# Step 3: Create districts
print("\n3. Creating districts...")
district_ids = {}
for province_name, districts in PROVINCE_DISTRICTS.items():
    for district_name in districts:
        code = district_name[:3].upper()
        cur.execute(
            "INSERT INTO districts (name, code, province_id) VALUES (%s, %s, %s) RETURNING id",
            (district_name, code, province_ids[province_name])
        )
        district_ids[district_name.upper()] = cur.fetchone()[0]
conn.commit()
print(f"   ✓ Created {len(district_ids)} districts")

# Step 4: Parse DOCX and extract all schools with trades
print("\n4. Parsing DOCX file...")
doc = Document('/app/LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC_2024-2025_lJhq.docx')

schools_data = {}  # {school_name: {district, type, trades: []}}

for table_idx, table in enumerate(doc.tables):
    print(f"   Processing table {table_idx + 1}/{len(doc.tables)}...")
    
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:  # Skip header
            continue
            
        cells = [cell.text.strip() for cell in row.cells]
        
        if len(cells) < 7:
            continue
        
        district = cells[1].strip()
        school_name = cells[2].strip()
        school_type = cells[3].strip() if len(cells) > 3 else "TSS"
        
        # Skip invalid rows
        if not school_name or school_name in ['School Name', '']:
            continue
        if not district or district in ['District', '']:
            continue
            
        # Clean school name
        school_name = clean_school_name(school_name)
        
        # Determine school type
        if school_type.upper() not in ['TSS', 'VTC']:
            if 'VTC' in school_name.upper():
                school_type = 'VTC'
            else:
                school_type = 'TSS'
        
        # Get province
        province = get_province_for_district(district)
        if not province:
            print(f"   ⚠ Unknown district: {district} for school {school_name}")
            continue
        
        # Extract trades from columns 6 and 7
        accredited_trades = cells[6] if len(cells) > 6 else ""
        non_accredited_trades = cells[7] if len(cells) > 7 else ""
        
        all_trades_text = accredited_trades + "\n" + non_accredited_trades
        
        # Parse trades
        trades = []
        for line in all_trades_text.split('\n'):
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            # Split by pipe or comma
            trade_parts = re.split(r'[|,]', line)
            for trade_part in trade_parts:
                trade_part = trade_part.strip()
                if trade_part and len(trade_part) > 2:
                    level = extract_level(trade_part)
                    trade_clean = clean_trade_name(trade_part)
                    if trade_clean:
                        trades.append({'name': trade_clean, 'level': level})
        
        # Store school data (avoid duplicates by using school name as key)
        if school_name not in schools_data:
            schools_data[school_name] = {
                'district': district,
                'province': province,
                'type': school_type,
                'trades': trades
            }
        else:
            # Merge trades if duplicate
            existing_trade_names = {t['name'] for t in schools_data[school_name]['trades']}
            for trade in trades:
                if trade['name'] not in existing_trade_names:
                    schools_data[school_name]['trades'].append(trade)

print(f"   ✓ Found {len(schools_data)} unique schools")

# Step 5: Insert schools
print("\n5. Inserting schools into database...")
school_ids = {}
for school_name, data in schools_data.items():
    district_key = data['district'].upper()
    if district_key not in district_ids:
        print(f"   ⚠ District not found: {district_key} for {school_name}")
        continue
    
    code = ''.join([c[0] for c in school_name.split()[:3]]).upper()[:10]
    
    cur.execute("""
        INSERT INTO schools (name, code, district_id, school_type, is_active)
        VALUES (%s, %s, %s, %s, true)
        RETURNING id
    """, (school_name, code, district_ids[district_key], data['type']))
    
    school_ids[school_name] = cur.fetchone()[0]

conn.commit()
print(f"   ✓ Inserted {len(school_ids)} schools")

# Step 6: Insert trades
print("\n6. Inserting trades...")
trade_ids = {}
all_trades = set()

for school_name, data in schools_data.items():
    for trade in data['trades']:
        all_trades.add(trade['name'])

for trade_name in sorted(all_trades):
    code = generate_trade_code(trade_name)
    category = assign_category(trade_name)
    
    try:
        cur.execute("""
            INSERT INTO trades (name, code, category)
            VALUES (%s, %s, %s)
            RETURNING id
        """, (trade_name, code, category))
        trade_ids[trade_name] = cur.fetchone()[0]
    except psycopg2.IntegrityError:
        conn.rollback()
        # Trade already exists, get its ID
        cur.execute("SELECT id FROM trades WHERE name = %s", (trade_name,))
        result = cur.fetchone()
        if result:
            trade_ids[trade_name] = result[0]

conn.commit()
print(f"   ✓ Inserted {len(trade_ids)} unique trades")

# Step 7: Create school-trade relationships
print("\n7. Creating school-trade relationships...")
relationship_count = 0

for school_name, data in schools_data.items():
    if school_name not in school_ids:
        continue
    
    school_id = school_ids[school_name]
    
    # Group trades by name to collect all levels
    trade_levels = {}
    for trade in data['trades']:
        trade_name = trade['name']
        level = trade['level']
        if trade_name not in trade_levels:
            trade_levels[trade_name] = set()
        trade_levels[trade_name].add(level)
    
    # Insert relationships
    for trade_name, levels in trade_levels.items():
        if trade_name not in trade_ids:
            continue
        
        trade_id = trade_ids[trade_name]
        levels_json = json.dumps(sorted(list(levels)))
        
        try:
            cur.execute("""
                INSERT INTO school_trades (school_id, trade_id, levels_offered, is_active)
                VALUES (%s, %s, %s, true)
            """, (school_id, trade_id, levels_json))
            relationship_count += 1
        except psycopg2.IntegrityError:
            conn.rollback()
            continue

conn.commit()
print(f"   ✓ Created {relationship_count} school-trade relationships")

# Step 8: Verification
print("\n8. Verification:")
cur.execute("SELECT COUNT(*) FROM provinces")
print(f"   Provinces: {cur.fetchone()[0]}")

cur.execute("SELECT COUNT(*) FROM districts")
print(f"   Districts: {cur.fetchone()[0]}")

cur.execute("SELECT COUNT(*) FROM schools")
print(f"   Schools: {cur.fetchone()[0]}")

cur.execute("SELECT COUNT(*) FROM trades")
print(f"   Trades: {cur.fetchone()[0]}")

cur.execute("SELECT COUNT(*) FROM school_trades")
print(f"   School-Trade Relationships: {cur.fetchone()[0]}")

# Check for duplicates
cur.execute("SELECT name, COUNT(*) FROM schools GROUP BY name HAVING COUNT(*) > 1")
duplicates = cur.fetchall()
if duplicates:
    print(f"\n   ⚠ WARNING: Found {len(duplicates)} duplicate schools:")
    for dup in duplicates:
        print(f"      - {dup[0]}: {dup[1]} times")
else:
    print("   ✓ No duplicate schools")

# Check specific schools
print("\n9. Checking specific schools:")
test_schools = ['APEK RRUKOMA', 'SAINT IGNACE', 'NYAMATA TSS', 'JANJA TSS']
for test_name in test_schools:
    cur.execute("SELECT name FROM schools WHERE name ILIKE %s", (f'%{test_name}%',))
    results = cur.fetchall()
    if results:
        print(f"   ✓ Found: {', '.join([r[0] for r in results])}")
    else:
        print(f"   ✗ NOT FOUND: {test_name}")

cur.close()
conn.close()

print("\n✅ Complete repopulation finished!")
