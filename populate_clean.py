#!/usr/bin/env python3
"""
Clean population from DOCX - Start from scratch
"""

import os
import re
import json
from docx import Document
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://quiz_user:quiz_pass123@db:5432/morning_quiz")

PROVINCE_DISTRICTS = {
    'Eastern Province': ['BUGESERA', 'GATSIBO', 'KAYONZA', 'KIREHE', 'NGOMA', 'RWAMAGANA', 'NYAGATARE'],
    'Kigali City': ['GASABO', 'KICUKIRO', 'NYARUGENGE'],
    'Northern Province': ['BURERA', 'GAKENKE', 'GICUMBI', 'MUSANZE', 'RULINDO'],
    'Southern Province': ['GISAGARA', 'HUYE', 'KAMONYI', 'MUHANGA', 'NYAMAGABE', 'NYANZA', 'NYARUGURU', 'RUHANGO'],
    'Western Province': ['KARONGI', 'NGORORERO', 'NYABIHU', 'NYAMASHEKE', 'RUBAVU', 'RUSIZI', 'RUTSIRO']
}

def extract_level(trade_str):
    """Extract level from trade string like 'Building Construction L3-5'"""
    match = re.search(r'L(\d+)(?:-(\d+))?', trade_str, re.IGNORECASE)
    if match:
        if match.group(2):
            return f"L{match.group(1)}-{match.group(2)}"
        return f"L{match.group(1)}"
    return 'L3-5'

def clean_trade_name(trade_str):
    """Remove level and numbering from trade name"""
    # Remove numbering like "1. "
    trade_str = re.sub(r'^\s*\d+\.\s*', '', trade_str)
    # Remove level suffix like "L3-5"
    trade_str = re.sub(r'\s*L\d+(?:-\d+)?\s*$', '', trade_str, flags=re.IGNORECASE)
    # Remove leading dots
    trade_str = re.sub(r'^\s*\.+\s*', '', trade_str)
    return trade_str.strip()

def parse_docx(docx_path):
    """Parse DOCX and extract all schools with trades"""
    doc = Document(docx_path)
    all_schools = []
    
    print(f"📄 Parsing DOCX with {len(doc.tables)} tables...")
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 5:
            continue
        
        # Find header row
        header_row_idx = None
        for idx in range(min(5, len(table.rows))):
            row_text = ' '.join([cell.text for cell in table.rows[idx].cells])
            if 'District' in row_text and 'School' in row_text:
                header_row_idx = idx
                break
        
        if header_row_idx is None:
            continue
        
        header_cells = [cell.text.strip() for cell in table.rows[header_row_idx].cells]
        
        # Find column indices
        district_col = school_col = type_col = trades_col = None
        for idx, header in enumerate(header_cells):
            if 'District' in header:
                district_col = idx
            elif 'School' in header:
                school_col = idx
            elif 'Type' in header:
                type_col = idx
            elif 'Accredited' in header and 'Not' not in header:
                trades_col = idx
        
        if district_col is None or school_col is None or trades_col is None:
            continue
        
        print(f"  Table {table_idx + 1}: Processing...")
        
        # Parse data rows
        for row_idx in range(header_row_idx + 1, len(table.rows)):
            row = table.rows[row_idx]
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) <= max(district_col, school_col, trades_col):
                continue
            
            district = cells[district_col].upper().strip()
            school_name = cells[school_col].replace('\n', ' ').strip()
            school_type = cells[type_col] if type_col and type_col < len(cells) else 'TSS'
            trades_text = cells[trades_col]
            
            # Validate
            if not district or not school_name or len(school_name) < 3:
                continue
            
            all_districts = [d for districts in PROVINCE_DISTRICTS.values() for d in districts]
            if district not in all_districts:
                continue
            
            # Parse trades - split by newline AND pipe
            trades = []
            if trades_text and trades_text not in ['None', '']:
                trade_items = re.split(r'[\n|]+', trades_text)
                for trade_item in trade_items:
                    trade_item = trade_item.strip()
                    # Remove numbering
                    trade_item = re.sub(r'^\d+\.\s*', '', trade_item)
                    if trade_item and len(trade_item) > 5 and trade_item not in ['None']:
                        trades.append(trade_item)
            
            if trades:
                all_schools.append({
                    'district': district,
                    'school_name': school_name,
                    'school_type': school_type,
                    'trades': trades
                })
    
    return all_schools

def populate_clean():
    """Populate database from scratch"""
    engine = create_engine(DATABASE_URL)
    Session = sessionmaker(bind=engine)
    db = Session()
    
    print("\n" + "="*70)
    print("CLEAN POPULATION FROM DOCX")
    print("="*70)
    
    try:
        # 1. Insert Provinces
        print("\n1️⃣ Creating Provinces...")
        province_ids = {}
        for province_name in PROVINCE_DISTRICTS.keys():
            result = db.execute(text("""
                INSERT INTO provinces (name, code) 
                VALUES (:name, :code) 
                RETURNING id
            """), {"name": province_name, "code": province_name[:3].upper()})
            province_ids[province_name] = result.fetchone()[0]
            db.commit()
        print(f"   ✅ Created {len(province_ids)} provinces")
        
        # 2. Insert Districts
        print("\n2️⃣ Creating Districts...")
        district_ids = {}
        for province_name, districts in PROVINCE_DISTRICTS.items():
            for district in districts:
                result = db.execute(text("""
                    INSERT INTO districts (province_id, name, code) 
                    VALUES (:pid, :name, :code) 
                    RETURNING id
                """), {
                    "pid": province_ids[province_name], 
                    "name": district.title(),  # Store as title case
                    "code": district[:3].upper()
                })
                district_ids[district] = result.fetchone()[0]  # Key is uppercase for matching
                db.commit()
        print(f"   ✅ Created {len(district_ids)} districts")
        
        # 3. Parse DOCX
        print("\n3️⃣ Parsing DOCX...")
        docx_path = '/app/LIST_OF_REACREDITTED_AND_NON_REACREDITED_TSS_AND_VTC_2024-2025_lJhq.docx'
        schools_data = parse_docx(docx_path)
        print(f"   ✅ Found {len(schools_data)} schools with trades")
        
        # 4. Insert Schools and Trades
        print("\n4️⃣ Creating Schools and Trades...")
        trade_ids = {}
        school_count = 0
        relationship_count = 0
        
        for school_data in schools_data:
            district = school_data['district']
            
            if district not in district_ids:
                print(f"   ⚠️  District not found: {district}")
                continue
            
            # Insert school
            result = db.execute(text("""
                INSERT INTO schools (district_id, name, school_type, is_active)
                VALUES (:did, :name, :type, true)
                RETURNING id
            """), {
                "did": district_ids[district],
                "name": school_data['school_name'],
                "type": school_data['school_type']
            })
            school_id = result.fetchone()[0]
            school_count += 1
            db.commit()
            
            # Insert trades for this school
            for trade_full in school_data['trades']:
                level = extract_level(trade_full)
                trade_name = clean_trade_name(trade_full)
                
                if len(trade_name) < 3:
                    continue
                
                # Get or create trade
                if trade_name not in trade_ids:
                    # Create unique code
                    base_code = trade_name[:10].upper().replace(' ', '_')
                    code = base_code
                    counter = 1
                    
                    # Check if code exists
                    while True:
                        result = db.execute(text("SELECT id FROM trades WHERE code = :code"), {"code": code})
                        existing = result.fetchone()
                        if not existing:
                            break
                        code = f"{base_code}_{counter}"
                        counter += 1
                    
                    result = db.execute(text("""
                        INSERT INTO trades (name, code)
                        VALUES (:name, :code)
                        RETURNING id
                    """), {
                        "name": trade_name,
                        "code": code
                    })
                    trade_ids[trade_name] = result.fetchone()[0]
                    db.commit()
                
                trade_id = trade_ids[trade_name]
                
                # Insert school-trade relationship
                db.execute(text("""
                    INSERT INTO school_trades (school_id, trade_id, levels_offered, is_active)
                    VALUES (:sid, :tid, CAST(:levels AS json), true)
                """), {
                    "sid": school_id,
                    "tid": trade_id,
                    "levels": json.dumps([level])
                })
                relationship_count += 1
                db.commit()
            
            if school_count % 20 == 0:
                print(f"   ... {school_count} schools processed")
        
        # Final stats
        print("\n" + "="*70)
        print("✅ POPULATION COMPLETE!")
        print("="*70)
        print(f"Provinces: {len(province_ids)}")
        print(f"Districts: {len(district_ids)}")
        print(f"Schools: {school_count}")
        print(f"Trades: {len(trade_ids)}")
        print(f"Relationships: {relationship_count}")
        
        # Verify
        result = db.execute(text("SELECT COUNT(*) FROM schools WHERE is_active = true"))
        print(f"\n📊 Active schools in database: {result.fetchone()[0]}")
        
        result = db.execute(text("""
            SELECT COUNT(*) FROM schools s 
            WHERE EXISTS (SELECT 1 FROM school_trades st WHERE st.school_id = s.id)
        """))
        print(f"📊 Schools with trades: {result.fetchone()[0]}")
        
        print("\n✅ Database ready!")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    populate_clean()
