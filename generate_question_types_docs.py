#!/usr/bin/env python3
"""
Question Types Documentation Generator
Generates PDF, DOC, and TXT files with all question types for Blockchain and Web3 module
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def create_question_content():
    """Create comprehensive question type content for Blockchain and Web3"""
    
    content = {
        'title': 'TVET Quiz System - Question Types Reference\nBlockchain and Web3 Development Module',
        'subtitle': 'Complete Guide for Students and Teachers',
        'date': datetime.now().strftime('%B %d, %Y'),
        'questions': [
            {
                'type': 'Multiple Choice Question (MCQ)',
                'description': 'Students select one correct answer from multiple options',
                'example': {
                    'question': 'What is the primary programming language used for Ethereum smart contracts?',
                    'options': [
                        'A) JavaScript',
                        'B) Python', 
                        'C) Solidity',
                        'D) Java'
                    ],
                    'answer': 'C) Solidity',
                    'explanation': 'Solidity is specifically designed for writing smart contracts on Ethereum blockchain.'
                }
            },
            {
                'type': 'True/False Question',
                'description': 'Students evaluate a statement as true or false',
                'example': {
                    'question': 'Smart contracts on Ethereum are immutable once deployed.',
                    'options': ['True', 'False'],
                    'answer': 'True',
                    'explanation': 'Once deployed, smart contracts cannot be modified, ensuring trust and transparency.'
                }
            },
            {
                'type': 'Short Answer Question',
                'description': 'Students provide a brief written response',
                'example': {
                    'question': 'Define what a blockchain is in your own words.',
                    'answer': 'A blockchain is a distributed ledger technology that maintains a continuously growing list of records (blocks) that are linked and secured using cryptography.',
                    'explanation': 'Accepts various correct definitions focusing on key concepts: distributed, ledger, blocks, cryptography.'
                }
            },
            {
                'type': 'Essay Question',
                'description': 'Students provide detailed written responses',
                'example': {
                    'question': 'Explain the advantages and disadvantages of using blockchain technology in supply chain management. Provide specific examples.',
                    'answer': 'Key points should include: transparency, traceability, immutability, cost considerations, scalability challenges, and real-world examples like Walmart\'s food tracking.',
                    'explanation': 'Graded based on depth of analysis, examples provided, and understanding of blockchain principles.'
                }
            },
            {
                'type': 'Multiple Select Question',
                'description': 'Students can select multiple correct answers',
                'example': {
                    'question': 'Which of the following are consensus mechanisms used in blockchain networks? (Select all that apply)',
                    'options': [
                        '☐ Proof of Work (PoW)',
                        '☐ Proof of Stake (PoS)',
                        '☐ HTTP Protocol',
                        '☐ Delegated Proof of Stake (DPoS)',
                        '☐ FTP Transfer'
                    ],
                    'answer': 'Proof of Work (PoW), Proof of Stake (PoS), Delegated Proof of Stake (DPoS)',
                    'explanation': 'Multiple consensus mechanisms exist. HTTP and FTP are network protocols, not consensus mechanisms.'
                }
            },
            {
                'type': 'Dropdown Select Question',
                'description': 'Students select from a dropdown menu',
                'example': {
                    'question': 'The gas limit in Ethereum transactions refers to:',
                    'options': [
                        'Select an option...',
                        'The maximum amount of gas units willing to spend',
                        'The current price of Ethereum',
                        'The number of confirmations needed',
                        'The wallet balance'
                    ],
                    'answer': 'The maximum amount of gas units willing to spend',
                    'explanation': 'Gas limit prevents infinite loops and controls transaction costs.'
                }
            },
            {
                'type': 'Fill in the Blanks',
                'description': 'Students complete sentences with missing words',
                'example': {
                    'question': 'In Solidity, the _____ keyword is used to declare state variables that cannot be changed after contract deployment, while _____ variables can be modified.',
                    'blanks': ['constant', 'public'],
                    'answer': 'constant, public (or other valid variable types)',
                    'explanation': 'Uses underscores (____) to indicate blanks. Students fill in appropriate Solidity keywords.'
                }
            },
            {
                'type': 'Matching Pairs',
                'description': 'Students match items from two columns',
                'example': {
                    'question': 'Match the blockchain terms with their definitions:',
                    'left_column': [
                        'Smart Contract',
                        'DApp',
                        'Token',
                        'Mining'
                    ],
                    'right_column': [
                        'Self-executing contract with terms directly written into code',
                        'Decentralized application running on blockchain',
                        'Digital asset representing value or utility',
                        'Process of validating transactions and adding blocks'
                    ],
                    'answer': 'Smart Contract → Self-executing contract, DApp → Decentralized application, Token → Digital asset, Mining → Process of validating transactions',
                    'explanation': 'Students drag and drop or draw lines to connect related items.'
                }
            },
            {
                'type': 'Drag & Drop Ordering',
                'description': 'Students arrange items in correct sequence',
                'example': {
                    'question': 'Arrange the steps of deploying a smart contract in the correct order:',
                    'items': [
                        'Write Solidity code',
                        'Compile the contract',
                        'Deploy to testnet for testing',
                        'Deploy to mainnet',
                        'Test contract functions'
                    ],
                    'answer': '1. Write Solidity code → 2. Compile the contract → 3. Deploy to testnet for testing → 4. Test contract functions → 5. Deploy to mainnet',
                    'explanation': 'Proper development workflow ensures contract security and functionality.'
                }
            },
            {
                'type': 'Linear Scale Rating',
                'description': 'Students rate on a scale of 1-10',
                'example': {
                    'question': 'On a scale of 1-10, how important is security auditing for smart contracts before mainnet deployment?',
                    'scale': 'Scale: 1 (Not Important) ←→ 10 (Extremely Important)',
                    'answer': '9-10 (Expected range)',
                    'explanation': 'Security auditing is crucial due to immutability and financial risks in smart contracts.'
                }
            },
            {
                'type': 'Code Writing Question',
                'description': 'Students write code in Solidity or other languages',
                'example': {
                    'question': 'Write a simple Solidity function that stores and retrieves a string value:',
                    'language': 'Solidity',
                    'answer': '''pragma solidity ^0.8.0;

contract SimpleStorage {
    string private storedData;
    
    function setData(string memory _data) public {
        storedData = _data;
    }
    
    function getData() public view returns (string memory) {
        return storedData;
    }
}''',
                    'explanation': 'Students must demonstrate understanding of Solidity syntax, state variables, and function visibility.'
                }
            },
            {
                'type': 'SQL Query Question',
                'description': 'Students write database queries (for blockchain data analysis)',
                'example': {
                    'question': 'Write a SQL query to find all transactions with value greater than 1 ETH from a blockchain transactions table:',
                    'answer': '''SELECT * FROM transactions 
WHERE value > 1000000000000000000 
ORDER BY block_number DESC;''',
                    'explanation': 'Query demonstrates understanding of blockchain data structure and Wei to ETH conversion (1 ETH = 10^18 Wei).'
                }
            }
        ]
    }
    
    return content

def generate_txt_file(content, filename):
    """Generate TXT file"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write(content['title'].center(80) + "\n")
        f.write(content['subtitle'].center(80) + "\n")
        f.write(f"Generated: {content['date']}".center(80) + "\n")
        f.write("=" * 80 + "\n\n")
        
        f.write("OVERVIEW\n")
        f.write("-" * 40 + "\n")
        f.write("This document contains all 12 question types supported by the TVET Quiz System,\n")
        f.write("specifically designed for Blockchain and Web3 Development education.\n")
        f.write("Each question type includes examples relevant to blockchain technology,\n")
        f.write("smart contracts, and decentralized applications.\n\n")
        
        for i, q in enumerate(content['questions'], 1):
            f.write(f"{i}. {q['type'].upper()}\n")
            f.write("=" * len(f"{i}. {q['type']}") + "\n\n")
            f.write(f"Description: {q['description']}\n\n")
            
            f.write("EXAMPLE:\n")
            f.write("-" * 20 + "\n")
            f.write(f"Question: {q['example']['question']}\n\n")
            
            if 'options' in q['example']:
                if isinstance(q['example']['options'], list):
                    for option in q['example']['options']:
                        f.write(f"  {option}\n")
                else:
                    f.write(f"Options: {q['example']['options']}\n")
                f.write("\n")
            
            if 'left_column' in q['example']:
                f.write("Left Column:\n")
                for item in q['example']['left_column']:
                    f.write(f"  • {item}\n")
                f.write("\nRight Column:\n")
                for item in q['example']['right_column']:
                    f.write(f"  • {item}\n")
                f.write("\n")
            
            if 'items' in q['example']:
                f.write("Items to arrange:\n")
                for item in q['example']['items']:
                    f.write(f"  • {item}\n")
                f.write("\n")
            
            if 'scale' in q['example']:
                f.write(f"{q['example']['scale']}\n\n")
            
            if 'language' in q['example']:
                f.write(f"Programming Language: {q['example']['language']}\n\n")
            
            f.write(f"Correct Answer: {q['example']['answer']}\n\n")
            f.write(f"Explanation: {q['example']['explanation']}\n\n")
            f.write("-" * 80 + "\n\n")

def generate_pdf_file(content, filename):
    """Generate PDF file"""
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=20,
        alignment=1  # Center
    )
    
    # Title page
    story.append(Paragraph(content['title'], title_style))
    story.append(Paragraph(content['subtitle'], subtitle_style))
    story.append(Paragraph(f"Generated: {content['date']}", styles['Normal']))
    story.append(Spacer(1, 0.5*inch))
    
    # Overview
    story.append(Paragraph("OVERVIEW", styles['Heading2']))
    overview_text = """This document contains all 12 question types supported by the TVET Quiz System, 
    specifically designed for Blockchain and Web3 Development education. Each question type includes 
    examples relevant to blockchain technology, smart contracts, and decentralized applications."""
    story.append(Paragraph(overview_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Questions
    for i, q in enumerate(content['questions'], 1):
        if i > 1:
            story.append(PageBreak())
        
        story.append(Paragraph(f"{i}. {q['type']}", styles['Heading2']))
        story.append(Paragraph(f"<b>Description:</b> {q['description']}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("<b>EXAMPLE:</b>", styles['Heading3']))
        story.append(Paragraph(f"<b>Question:</b> {q['example']['question']}", styles['Normal']))
        
        if 'options' in q['example']:
            if isinstance(q['example']['options'], list):
                for option in q['example']['options']:
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{option}", styles['Normal']))
            else:
                story.append(Paragraph(f"<b>Options:</b> {q['example']['options']}", styles['Normal']))
        
        if 'left_column' in q['example']:
            story.append(Paragraph("<b>Left Column:</b>", styles['Normal']))
            for item in q['example']['left_column']:
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;• {item}", styles['Normal']))
            story.append(Paragraph("<b>Right Column:</b>", styles['Normal']))
            for item in q['example']['right_column']:
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;• {item}", styles['Normal']))
        
        if 'items' in q['example']:
            story.append(Paragraph("<b>Items to arrange:</b>", styles['Normal']))
            for item in q['example']['items']:
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;• {item}", styles['Normal']))
        
        if 'scale' in q['example']:
            story.append(Paragraph(q['example']['scale'], styles['Normal']))
        
        if 'language' in q['example']:
            story.append(Paragraph(f"<b>Programming Language:</b> {q['example']['language']}", styles['Normal']))
        
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph(f"<b>Correct Answer:</b> {q['example']['answer']}", styles['Normal']))
        story.append(Paragraph(f"<b>Explanation:</b> {q['example']['explanation']}", styles['Normal']))
    
    doc.build(story)

def generate_doc_file(content, filename):
    """Generate DOC file"""
    doc = Document()
    
    # Title
    title = doc.add_heading(content['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(content['subtitle'], level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f"Generated: {content['date']}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Overview
    doc.add_heading('OVERVIEW', level=1)
    overview_text = """This document contains all 12 question types supported by the TVET Quiz System, 
    specifically designed for Blockchain and Web3 Development education. Each question type includes 
    examples relevant to blockchain technology, smart contracts, and decentralized applications."""
    doc.add_paragraph(overview_text)
    
    # Questions
    for i, q in enumerate(content['questions'], 1):
        doc.add_page_break()
        
        doc.add_heading(f"{i}. {q['type']}", level=1)
        
        desc_para = doc.add_paragraph()
        desc_para.add_run("Description: ").bold = True
        desc_para.add_run(q['description'])
        
        doc.add_heading('EXAMPLE:', level=2)
        
        question_para = doc.add_paragraph()
        question_para.add_run("Question: ").bold = True
        question_para.add_run(q['example']['question'])
        
        if 'options' in q['example']:
            if isinstance(q['example']['options'], list):
                for option in q['example']['options']:
                    doc.add_paragraph(f"    {option}", style='List Bullet')
            else:
                options_para = doc.add_paragraph()
                options_para.add_run("Options: ").bold = True
                options_para.add_run(q['example']['options'])
        
        if 'left_column' in q['example']:
            left_para = doc.add_paragraph()
            left_para.add_run("Left Column:").bold = True
            for item in q['example']['left_column']:
                doc.add_paragraph(f"• {item}", style='List Bullet')
            
            right_para = doc.add_paragraph()
            right_para.add_run("Right Column:").bold = True
            for item in q['example']['right_column']:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        
        if 'items' in q['example']:
            items_para = doc.add_paragraph()
            items_para.add_run("Items to arrange:").bold = True
            for item in q['example']['items']:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        
        if 'scale' in q['example']:
            doc.add_paragraph(q['example']['scale'])
        
        if 'language' in q['example']:
            lang_para = doc.add_paragraph()
            lang_para.add_run("Programming Language: ").bold = True
            lang_para.add_run(q['example']['language'])
        
        answer_para = doc.add_paragraph()
        answer_para.add_run("Correct Answer: ").bold = True
        answer_para.add_run(q['example']['answer'])
        
        explanation_para = doc.add_paragraph()
        explanation_para.add_run("Explanation: ").bold = True
        explanation_para.add_run(q['example']['explanation'])
    
    doc.save(filename)

def main():
    """Main function to generate all document formats"""
    print("Generating Question Types Documentation...")
    print("Context: Blockchain and Web3 Development Module")
    
    # Create content
    content = create_question_content()
    
    # Create output directory
    output_dir = "question_types_docs"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate files
    base_filename = "Blockchain_Web3_Question_Types"
    
    print("\nGenerating TXT file...")
    txt_file = os.path.join(output_dir, f"{base_filename}.txt")
    generate_txt_file(content, txt_file)
    print(f"TXT file created: {txt_file}")
    
    print("\nGenerating PDF file...")
    pdf_file = os.path.join(output_dir, f"{base_filename}.pdf")
    generate_pdf_file(content, pdf_file)
    print(f"PDF file created: {pdf_file}")
    
    print("\nGenerating DOC file...")
    doc_file = os.path.join(output_dir, f"{base_filename}.docx")
    generate_doc_file(content, doc_file)
    print(f"DOC file created: {doc_file}")
    
    print(f"\nAll documents generated successfully!")
    print(f"Output directory: {os.path.abspath(output_dir)}")
    print("\nFiles created:")
    print(f"   - {base_filename}.txt")
    print(f"   - {base_filename}.pdf") 
    print(f"   - {base_filename}.docx")
    
    print("\nThese documents contain:")
    print("   - All 12 question types with Blockchain/Web3 examples")
    print("   - Solidity code examples")
    print("   - Smart contract scenarios")
    print("   - DApp development questions")
    print("   - Cryptocurrency and token concepts")

if __name__ == "__main__":
    main()